from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import List, Optional, Any, Dict
from app.services.llm_service import llm_service
from app.core.logging_config import get_logger
from app.services.message_state_service import message_state_manager
from app.services.dataframe_state_manager import dataframe_state_manager
from app.services.dataset_service import dataset_manager
from io import BytesIO
from fastapi.responses import StreamingResponse
import base64
import json
import re
import zipfile
import os
import tempfile
import shutil

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Documentation download will not work.")

try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
    print("Warning: openpyxl not installed. Excel file generation will not work.")

logger = get_logger(__name__)


def create_excel_from_table_data(table_data: List[Dict[str, Any]], headers: List[str], file_path: str):
    """
    Create an Excel file from table data.
    
    Args:
        table_data: List of dictionaries containing table row data
        headers: List of column headers
        file_path: Path where Excel file should be saved
    """
    if not OPENPYXL_AVAILABLE:
        logger.error("openpyxl not available, cannot create Excel file")
        return
    
    wb = openpyxl.Workbook()
    ws = wb.active
    
    # Add headers
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Add data rows
    for row_idx, row_data in enumerate(table_data, start=2):
        for col_idx, header in enumerate(headers, start=1):
            value = row_data.get(header, '')
            
            # Format numeric values
            if isinstance(value, (int, float)) and not isinstance(value, bool):
                if isinstance(value, float):
                    ws.cell(row=row_idx, column=col_idx, value=value)
                    ws.cell(row=row_idx, column=col_idx).number_format = '0.0000'
                else:
                    ws.cell(row=row_idx, column=col_idx, value=value)
            else:
                ws.cell(row=row_idx, column=col_idx, value=str(value) if value else '')
            
            ws.cell(row=row_idx, column=col_idx).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    
    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column].width = adjusted_width
    
    wb.save(file_path)
    logger.info(f"Excel file created at {file_path}")


def _get_excel_icon_path() -> str:
    """
    Return the bundled microsoft-excel-logo.jpg path for use as the <v:imagedata>
    icon in the pure-Python OOXML embedding path.
    """
    current_dir = os.path.dirname(os.path.abspath(__file__))
    static_dir = os.path.abspath(os.path.join(current_dir, '..', 'static'))
    bundled_logo = os.path.join(static_dir, 'microsoft-excel-logo.jpg')
    if os.path.exists(bundled_logo):
        logger.info(f"Using bundled Excel logo: {bundled_logo}")
        return bundled_logo
    logger.warning("Bundled Excel logo not found; OLE icon image will be omitted")
    return None


def _get_excel_exe_path() -> str:
    """
    Return the path to excel.exe for use as IconFileName in the Windows COM
    AddOLEObject call.  Only PE files (.exe/.dll) are valid for that parameter;
    image files are silently ignored by Word and produce a blank white icon.
    Returns None on non-Windows or when Excel is not installed.
    """
    try:
        import winreg
        key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
                             r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\excel.exe")
        path, _ = winreg.QueryValueEx(key, "")
        winreg.CloseKey(key)
        if path and os.path.exists(path):
            logger.info(f"Found Excel executable for COM icon: {path}")
            return path
    except Exception:
        pass
    logger.info("Excel executable not found in registry; COM will use Word's default OLE icon")
    return None


def _embed_excel_ole_pure_python(docx_path: str, excel_files: List[str]) -> bool:
    """
    Cross-platform pure-Python OLE embedding for Linux/RI (Azure) environments
    where Microsoft Word / COM is not available.

    Directly manipulates the .docx ZIP to:
      1. Copy each .xlsx into word/embeddings/ as an embedded OLE part.
      2. Copy microsoft-excel-logo.jpg into word/media/ as the icon image.
      3. Patch [Content_Types].xml and word/_rels/document.xml.rels.
      4. Replace each placeholder paragraph in word/document.xml with a
         properly structured <w:object> / <o:OLEObject> block including the
         required VML shapetype definition so Word renders it as a clickable
         icon without any corruption warning.
    """
    import zipfile as _zf
    import shutil
    import re as _re

    RT_IMAGE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
    RT_PACKAGE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/package'
    CT_XLSX  = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    CT_JPG   = 'image/jpeg'
    CT_NS    = 'http://schemas.openxmlformats.org/package/2006/content-types'
    REL_NS   = 'http://schemas.openxmlformats.org/package/2006/relationships'

    icon_path = _get_excel_icon_path()

    tmp_path = docx_path + '.tmp'
    try:
        # ── 1. Read entire docx into memory, keep a backup ────────────────────
        shutil.copy2(docx_path, tmp_path)
        with _zf.ZipFile(tmp_path, 'r') as zin:
            names    = zin.namelist()
            file_map = {n: zin.read(n) for n in names}

        # ── 2. Parse relationships, find next free rId ────────────────────────
        rels_path = 'word/_rels/document.xml.rels'
        rels_text = file_map.get(rels_path, b'').decode('utf-8')

        existing_nums = [int(m) for m in _re.findall(r'Id="rId(\d+)"', rels_text)]
        next_rid = max(existing_nums, default=0) + 1

        media_count = sum(1 for n in names if n.startswith('word/media/'))
        embed_count = sum(1 for n in names if n.startswith('word/embeddings/'))

        # ── 3. Add icon image (shared across all OLE objects) ─────────────────
        icon_rid  = None
        icon_part = None
        if icon_path and os.path.exists(icon_path):
            media_count += 1
            icon_part = f'word/media/image{media_count}.jpg'
            with open(icon_path, 'rb') as f:
                file_map[icon_part] = f.read()
            icon_rid   = f'rId{next_rid}'
            next_rid  += 1
            rels_text  = rels_text.replace(
                '</Relationships>',
                f'<Relationship Id="{icon_rid}" '
                f'Type="{RT_IMAGE}" '
                f'Target="media/image{media_count}.jpg"/>'
                f'</Relationships>'
            )
            logger.info(f"Added icon image part: {icon_part} ({icon_rid})")

        # ── 4. Add each Excel embedding part + relationship ───────────────────
        ole_info = []  # (display_name, ole_rid, icon_rid, filename, embed_idx)
        for excel_path in excel_files:
            filename     = os.path.basename(excel_path)
            display_name = filename.replace('.xlsx', '').replace('.xls', '')
            embed_count += 1
            embed_part   = f'word/embeddings/Microsoft_Excel_Sheet{embed_count}.xlsx'
            with open(excel_path, 'rb') as f:
                file_map[embed_part] = f.read()
            ole_rid    = f'rId{next_rid}'
            next_rid  += 1
            rels_text  = rels_text.replace(
                '</Relationships>',
                f'<Relationship Id="{ole_rid}" '
                f'Type="{RT_PACKAGE}" '
                f'Target="embeddings/Microsoft_Excel_Sheet{embed_count}.xlsx"/>'
                f'</Relationships>'
            )
            ole_info.append((display_name, ole_rid, icon_rid, filename, embed_count))
            logger.info(f"Added OLE embedding part: {embed_part} ({ole_rid})")

        file_map[rels_path] = rels_text.encode('utf-8')

        # ── 5. Patch [Content_Types].xml ──────────────────────────────────────
        ct_text = file_map.get('[Content_Types].xml', b'').decode('utf-8')
        if CT_XLSX not in ct_text:
            ct_text = ct_text.replace(
                '</Types>',
                f'<Default Extension="xlsx" ContentType="{CT_XLSX}"/></Types>'
            )
        if CT_JPG not in ct_text and icon_part:
            ct_text = ct_text.replace(
                '</Types>',
                f'<Default Extension="jpg" ContentType="{CT_JPG}"/></Types>'
            )
        for _, _, _, _, embed_idx in ole_info:
            part_name = f'/word/embeddings/Microsoft_Excel_Sheet{embed_idx}.xlsx'
            if part_name not in ct_text:
                ct_text = ct_text.replace(
                    '</Types>',
                    f'<Override PartName="{part_name}" ContentType="{CT_XLSX}"/></Types>'
                )
        file_map['[Content_Types].xml'] = ct_text.encode('utf-8')

        # ── 6. Patch word/document.xml ────────────────────────────────────────
        doc_str = file_map.get('word/document.xml', b'').decode('utf-8')

        # Ensure required VML namespaces are declared on <w:document>
        for attr, val in [
            ('xmlns:v',  'urn:schemas-microsoft-com:vml'),
            ('xmlns:o',  'urn:schemas-microsoft-com:office:office'),
        ]:
            if attr not in doc_str:
                doc_str = doc_str.replace('<w:document ', f'<w:document {attr}="{val}" ', 1)

        # Replace each placeholder paragraph with a proper <w:object> block.
        # The structure below is based on a COM-generated .docx from Word.
        for display_name, ole_rid, img_rid, filename, _ in ole_info:
            shape_id = f'_x0000_i{abs(hash(ole_rid)) % 90000 + 1025}'
            object_id = f'_{abs(hash((filename, ole_rid))) % 2000000000 + 1000000000}'

            v_imagedata = (
                f'<v:imagedata r:id="{img_rid}" o:title="Microsoft Excel"/>'
                if img_rid else ''
            )

            shapetype_block = (
                '<v:shapetype id="_x0000_t75" coordsize="21600,21600" '
                'o:spt="75" o:preferrelative="t" '
                'path="m@4@5l@4@11@9@11@9@5xe" filled="f" stroked="f">'
                '<v:stroke joinstyle="miter"/>'
                '<v:formulas>'
                '<v:f eqn="if lineDrawn pixelLineWidth 0"/>'
                '<v:f eqn="sum @0 1 0"/>'
                '<v:f eqn="sum 0 0 @1"/>'
                '<v:f eqn="prod @2 1 2"/>'
                '<v:f eqn="prod @3 21600 pixelWidth"/>'
                '<v:f eqn="prod @3 21600 pixelHeight"/>'
                '<v:f eqn="sum @0 0 1"/>'
                '<v:f eqn="prod @6 1 2"/>'
                '<v:f eqn="prod @7 21600 pixelWidth"/>'
                '<v:f eqn="sum @8 21600 0"/>'
                '<v:f eqn="prod @7 21600 pixelHeight"/>'
                '<v:f eqn="sum @10 21600 0"/>'
                '</v:formulas>'
                '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>'
                '<o:lock v:ext="edit" aspectratio="t"/>'
                '</v:shapetype>'
            )

            ole_block = (
                f'<w:object w:dxaOrig="1534" w:dyaOrig="997">'
                f'{shapetype_block}'
                f'<v:shape id="{shape_id}" type="#_x0000_t75" '
                f'style="width:76.5pt;height:49.5pt" o:ole="">'
                f'{v_imagedata}'
                f'</v:shape>'
                f'<o:OLEObject Type="Embed" ProgID="Excel.Sheet.12" '
                f'ShapeID="{shape_id}" DrawAspect="Icon" ObjectID="{object_id}" '
                f'r:id="{ole_rid}"/>'
                f'</w:object>'
            )

            replacement_para = (
                f'<w:p>'
                f'<w:pPr><w:jc w:val="left"/>'
                f'<w:rPr><w:noProof/></w:rPr></w:pPr>'
                f'<w:r><w:rPr><w:noProof/></w:rPr>{ole_block}</w:r>'
                f'</w:p>'
            )

            escaped = _re.escape(display_name)
            pattern = (
                r'<w:p[ >](?:(?!<w:p[ >]).)*?'
                r'<w:t[^>]*>' + escaped + r'</w:t>'
                r'(?:(?!<w:p[ >]).)*?</w:p>'
            )
            new_doc_str, n_subs = _re.subn(
                pattern, lambda m, rp=replacement_para: rp,
                doc_str, count=1, flags=_re.DOTALL
            )
            if n_subs:
                doc_str = new_doc_str
                logger.info(f"Replaced placeholder paragraph for: {display_name}")
            else:
                logger.warning(f"Placeholder not found for: {display_name} - using inline fallback")
                doc_str = doc_str.replace(
                    f'<w:t>{display_name}</w:t>',
                    f'<w:t>{display_name}</w:t><w:r>{ole_block}</w:r>',
                    1
                )

        file_map['word/document.xml'] = doc_str.encode('utf-8')

        # ── 7. Repack the docx ────────────────────────────────────────────────
        with _zf.ZipFile(docx_path, 'w', _zf.ZIP_DEFLATED) as zout:
            for name, data in file_map.items():
                zout.writestr(name, data)

        os.remove(tmp_path)
        logger.info("✅ Pure-Python OLE embedding complete")
        return True

    except Exception as e:
        import traceback
        logger.error(f"Pure-Python OLE embedding failed: {e}")
        logger.error(traceback.format_exc())
        if os.path.exists(tmp_path):
            shutil.copy2(tmp_path, docx_path)
            os.remove(tmp_path)
        return False


def post_process_docx_with_ole_objects(docx_path: str, excel_files: List[str]) -> bool:
    """
    Embed Excel files as OLE objects (clickable icons) inside the DOCX.

    Strategy:
      - On Windows with pywin32 + Microsoft Word installed: use COM automation
        exclusively. COM produces a fully native OLE object with the real Excel
        icon and reliable double-click behaviour. The pure-Python path is NOT
        run on Windows because Word rejects a docx that has already been
        modified by the pure-Python path.
      - On Linux / RI (Azure) where COM is unavailable: use the pure-Python
        OOXML path which directly manipulates the .docx ZIP. The resulting icon
        shows the bundled microsoft-excel-logo.jpg and double-clicking opens the
        embedded Excel file in whatever spreadsheet app the user has installed.

    Args:
        docx_path: Absolute path to the .docx file to modify in-place.
        excel_files: List of absolute paths to .xlsx files to embed.

    Returns:
        True if embedding succeeded by either path, False otherwise.
    """
    if not excel_files:
        return True

    # Set OLE_EMBED_MODE=pure_python in your .env to force the Linux/RI code
    # path on Windows for local testing. Remove or unset it for normal use.
    if os.environ.get('OLE_EMBED_MODE', '').lower() == 'pure_python':
        logger.info("OLE_EMBED_MODE=pure_python - skipping COM, using pure-Python path (test mode)")
        return _embed_excel_ole_pure_python(docx_path, excel_files)

    # ── Try Windows COM first ─────────────────────────────────────────────────
    word_doc = None
    try:
        import win32com.client as win32
        import pythoncom

        logger.info("Windows + pywin32 detected - using COM automation for OLE embedding...")
        pythoncom.CoInitialize()

        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        word_doc = word.Documents.Open(docx_path)

        try:
            word_doc.Final = False
            if word_doc.ProtectionType != -1:
                word_doc.Unprotect()
        except Exception as e:
            logger.warning(f"Could not modify document protection: {e}")

        excel_exe_path = _get_excel_exe_path()

        for excel_path in excel_files:
            filename     = os.path.basename(excel_path)
            display_name = filename.replace('.xlsx', '').replace('.xls', '')
            try:
                word_range = word_doc.Content
                word_range.Find.ClearFormatting()
                word_range.Find.Text = display_name

                if word_range.Find.Execute():
                    try:
                        cell_range = word_range.Cells(1).Range
                        cell_range.Delete()
                        insert_range = cell_range
                    except Exception:
                        insert_range = word_range.Paragraphs(1).Range
                        insert_range.Delete()

                    if excel_exe_path:
                        ole_obj = word_doc.InlineShapes.AddOLEObject(
                            FileName=excel_path,
                            IconFileName=excel_exe_path,
                            IconIndex=0,
                            DisplayAsIcon=True,
                            Range=insert_range
                        )
                    else:
                        ole_obj = word_doc.InlineShapes.AddOLEObject(
                            FileName=excel_path,
                            DisplayAsIcon=True,
                            Range=insert_range
                        )

                    try:
                        ole_obj.Range.Paragraphs(1).Alignment = 0
                        text_range = ole_obj.Range
                        text_range.Collapse(0)
                        text_range.Text = f"  {filename}"
                        text_range.Font.Bold = True
                        text_range.Font.Size = 11
                        text_range.Font.ColorIndex = 5
                    except Exception as text_error:
                        logger.warning(f"Could not add filename text: {text_error}")
                else:
                    logger.warning(f"COM: could not find placeholder for: {filename}")

            except Exception as e:
                logger.error(f"COM: failed to embed OLE for {filename}: {e}")
                continue

        try:
            word_doc.SaveAs2(FileName=docx_path, FileFormat=16, AddToRecentFiles=False)
        except Exception:
            try:
                word_doc.Save()
            except Exception:
                pass

        word_doc.Close(SaveChanges=False)
        word.Quit()
        pythoncom.CoUninitialize()
        logger.info("✅ COM automation OLE embedding complete")
        return True

    except ImportError:
        logger.info("pywin32 not available - falling back to pure-Python OLE embedding (Linux/RI mode)")
    except Exception as e:
        try:
            if word_doc:
                word_doc.Close(SaveChanges=False)
            if 'word' in locals():
                word.Quit()  # noqa: F821
            if 'pythoncom' in locals():
                pythoncom.CoUninitialize()  # noqa: F821
        except Exception:
            pass
        logger.warning(f"COM automation failed: {e} - falling back to pure-Python OLE embedding")

    # ── Fallback: pure-Python OOXML path (Linux / RI / no Word installed) ────
    return _embed_excel_ole_pure_python(docx_path, excel_files)


def embed_excel_as_linked_object(doc, excel_file_path: str, display_text: str = None):
    """
    Add a reference to an Excel file that will be embedded later via post-processing.
    
    Note: True OLE embedding requires COM automation with Word installed.
    This function creates a styled reference that will be replaced with actual
    OLE objects during post-processing if win32com is available.
    
    Args:
        doc: Document object
        excel_file_path: Path to the Excel file
        display_text: Text to display (default: filename without extension)
    """
    filename = os.path.basename(excel_file_path)
    if display_text is None:
        display_text = filename.replace('.xlsx', '').replace('.xls', '')
    
    # Create a simple left-aligned reference (will be replaced by OLE object in post-processing)
    # Just add the display name text that will be found and replaced
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add display text that will be searched for during OLE embedding
    run = para.add_run(display_text)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 204)
    
    logger.info(f"Added Excel file reference: {filename}")
    doc.add_paragraph()
    return True


documentation_router = APIRouter()


class DataSummaryRequest(BaseModel):
    columns: List[str]
    data_dictionary: Optional[str] = None
    model_objective: Optional[str] = None


class DataSummaryResponse(BaseModel):
    success: bool
    summary: Optional[str] = None
    error: Optional[str] = None


class DataQualityMetrics(BaseModel):
    emptyColumns: int
    constantColumns: int
    sparseColumns: int
    formattingIssues: int
    emptyColumnNames: List[str]
    constantColumnNames: List[str]
    sparseColumnNames: List[str]
    formattingIssueColumnNames: List[str]


class DataQualitySummaryRequest(BaseModel):
    metrics: DataQualityMetrics
    recommendations: List[str]
    totalRows: int
    totalColumns: int


class DataQualitySummaryResponse(BaseModel):
    success: bool
    summary: Optional[str] = None
    error: Optional[str] = None


class TargetDefinitionRequest(BaseModel):
    target_variable: str
    data_dictionary: Optional[str] = None
    columns: List[str]
    problem_statement: Optional[str] = None


class TargetDefinitionResponse(BaseModel):
    success: bool
    definition: Optional[str] = None
    error: Optional[str] = None


class ModelObjectiveRequest(BaseModel):
    project_description: Optional[str] = None
    problem_statement: Optional[str] = None
    data_summary: Optional[str] = None
    target_variable_name: Optional[str] = None
    target_definition: Optional[str] = None


class ModelObjectiveResponse(BaseModel):
    success: bool
    objective: Optional[str] = None
    error: Optional[str] = None


class MonotonicitySummaryRequest(BaseModel):
    models: List[Dict[str, Any]]  # List of model data with monotonicity metrics


class MonotonicitySummaryResponse(BaseModel):
    success: bool
    writeup: Optional[str] = None
    error: Optional[str] = None


class EventRateRequest(BaseModel):
    dataset_id: str
    target_variable: str


class EventRateResponse(BaseModel):
    success: bool
    event_count: Optional[int] = None
    total_count: Optional[int] = None
    percentage: Optional[float] = None
    error: Optional[str] = None


class SamplingPlanRequest(BaseModel):
    dataset_id: str
    target_variable: str


class SamplingPlanResponse(BaseModel):
    success: bool
    has_split: Optional[bool] = None
    train: Optional[Dict[str, Any]] = None
    hold: Optional[Dict[str, Any]] = None
    writeup: Optional[str] = None
    error: Optional[str] = None


class SamplingPlanWriteupRequest(BaseModel):
    sampling_plan: Dict[str, Any]


class SamplingPlanWriteupResponse(BaseModel):
    success: bool
    writeup: Optional[str] = None
    error: Optional[str] = None


class ModelValidationWriteupRequest(BaseModel):
    model_validation: Dict[str, Any]
    data_summary: Optional[str] = None


class ModelValidationWriteupResponse(BaseModel):
    success: bool
    writeup: Optional[str] = None
    error: Optional[str] = None


class SegmentationUnderstandingRequest(BaseModel):
    data_summary: str  # Data Summary from 1. OBJECTIVE
    segments: List[Dict[str, Any]]  # All segment data including rules, totals, event rates, distributions
    segment_sizes: List[int]  # Segment sizes data
    segment_proportions: List[float]  # Segment proportions data
    event_rates: List[float]  # Event rates for each segment
    iv_report: Optional[Dict[str, Any]] = None  # IV report data if available


class SegmentationUnderstandingResponse(BaseModel):
    success: bool
    understanding: Optional[str] = None
    error: Optional[str] = None


class QualityCheckPlanRequest(BaseModel):
    dataset_id: str


class QualityCheckPlanResponse(BaseModel):
    success: bool
    plan: Optional[Dict[str, Any]] = None
    error: Optional[str] = None


class ColumnStatsRequest(BaseModel):
    dataset_id: str


class ColumnStatsResponse(BaseModel):
    success: bool
    stats: Optional[List[Dict[str, Any]]] = None
    error: Optional[str] = None


class QualityChangesWriteupRequest(BaseModel):
    quality_check_plan: Dict[str, Any]  # Analysis plan table data
    column_stats: List[Dict[str, Any]]  # Column stats table data


class QualityChangesWriteupResponse(BaseModel):
    success: bool
    writeup: Optional[str] = None
    error: Optional[str] = None


class DataInsightsRequest(BaseModel):
    dataset_id: str


class DataInsightsResponse(BaseModel):
    success: bool
    insights: Optional[Dict[str, Any]] = None
    error: Optional[str] = None


class TransformedVariablesRequest(BaseModel):
    dataset_id: str


class TransformedVariablesResponse(BaseModel):
    success: bool
    transformed_variables: Optional[List[Dict[str, Any]]] = None
    error: Optional[str] = None


class VariableAnalysisRequest(BaseModel):
    dataset_id: str


class VariableAnalysisResponse(BaseModel):
    success: bool
    variable_analysis: Optional[Dict[str, Any]] = None
    variable_statistics: Optional[List[Dict[str, Any]]] = None
    error: Optional[str] = None


class FeatureEngineeringWriteupRequest(BaseModel):
    transformed_variables: List[Dict[str, Any]]  # List of transformed variables with new_variable_name, var_type, variable_definition, transformation_methods


class FeatureEngineeringWriteupResponse(BaseModel):
    success: bool
    writeup: Optional[str] = None
    error: Optional[str] = None


class DecileProgressionWriteupRequest(BaseModel):
    model_name: str
    deciles: List[Dict[str, Any]]  # Decile table data
    monotonicity_score: float  # Percentage (0-100)
    violations: List[Dict[str, Any]]  # List of violations with fromDecile, toDecile, drop


class DecileProgressionWriteupResponse(BaseModel):
    success: bool
    writeup: Optional[str] = None
    error: Optional[str] = None


@documentation_router.post("/documentation/generate-data-summary", response_model=DataSummaryResponse)
async def generate_data_summary(request: DataSummaryRequest):
    """
    Generate a 5-line plain English data summary using LLM
    """
    try:
        logger.info(f"Generating data summary for {len(request.columns)} columns")
        
        # Build the prompt
        prompt_parts = [
            "You are a data analyst creating a concise data summary for model documentation.",
            "Generate a 5-line plain English summary that briefly describes what the data is about.",
            "",
            f"Number of columns: {len(request.columns)}",
            f"Column list: {', '.join(request.columns)}",
        ]
        
        if request.data_dictionary:
            prompt_parts.append(f"\nData Dictionary:\n{request.data_dictionary}")
        
        if request.model_objective:
            prompt_parts.append(f"\nModel Objective: {request.model_objective}")
        
        prompt_parts.extend([
            "",
            "Please provide a 5-line summary (approximately 5 sentences) that:",
            "1. Describes the nature of the dataset",
            "2. Mentions key variables or features",
            "3. Relates to the business context if available",
            "4. Is written in clear, professional language",
            "5. Is concise and informative",
            "",
            "Return ONLY the summary text, no additional formatting or preamble."
        ])
        
        prompt = "\n".join(prompt_parts)
        
        # Call LLM service
        summary = await llm_service.generate_text(
            prompt=prompt,
            max_tokens=300,
            temperature=0.3
        )
        
        if not summary or summary.strip() == "":
            return DataSummaryResponse(
                success=False,
                error="LLM returned empty response"
            )
        
        logger.info("Data summary generated successfully")
        return DataSummaryResponse(
            success=True,
            summary=summary.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating data summary: {str(e)}")
        return DataSummaryResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-data-quality-summary", response_model=DataQualitySummaryResponse)
async def generate_data_quality_summary(request: DataQualitySummaryRequest):
    """
    Generate a 5-6 line plain English data quality summary using LLM
    """
    try:
        logger.info(f"Generating data quality summary for dataset with {request.totalRows} rows and {request.totalColumns} columns")
        
        # Build the prompt
        prompt_parts = [
            "You are a data quality analyst creating a concise quality summary for model documentation.",
            "Generate a 5-6 line plain English summary that describes the data quality status.",
            "",
            f"Dataset: {request.totalRows} rows, {request.totalColumns} columns",
            "",
            "Quality Issues Detected:",
        ]
        
        if request.metrics.emptyColumns > 0:
            prompt_parts.append(f"- Empty Columns: {request.metrics.emptyColumns} columns have 100% missing values")
            prompt_parts.append(f"  Columns: {', '.join(request.metrics.emptyColumnNames)}")
        
        if request.metrics.constantColumns > 0:
            prompt_parts.append(f"- Constant Columns: {request.metrics.constantColumns} columns have only one unique value")
            prompt_parts.append(f"  Columns: {', '.join(request.metrics.constantColumnNames)}")
        
        if request.metrics.sparseColumns > 0:
            prompt_parts.append(f"- Sparse Columns: {request.metrics.sparseColumns} columns have >50% missing values")
            prompt_parts.append(f"  Columns: {', '.join(request.metrics.sparseColumnNames)}")
        
        if request.metrics.formattingIssues > 0:
            prompt_parts.append(f"- Formatting Issues: {request.metrics.formattingIssues} columns have formatting inconsistencies")
            prompt_parts.append(f"  Columns: {', '.join(request.metrics.formattingIssueColumnNames)}")
        
        if not any([request.metrics.emptyColumns, request.metrics.constantColumns, 
                   request.metrics.sparseColumns, request.metrics.formattingIssues]):
            prompt_parts.append("- No major quality issues detected")
        
        prompt_parts.append("")
        prompt_parts.append("Recommendations:")
        if request.recommendations:
            for rec in request.recommendations:
                prompt_parts.append(f"- {rec}")
        else:
            prompt_parts.append("- Data appears to be in good quality for modeling")
        
        prompt_parts.extend([
            "",
            "Please provide a bullet-point summary (5-6 bullet points) that:",
            "1. Gives an overall assessment of data quality",
            "2. Highlights key issues (if any)",
            "3. Mentions recommendations or actions needed",
            "4. Is written in clear, professional language",
            "5. Provides actionable insights",
            "",
            "FORMAT REQUIREMENTS:",
            "- Each point should start with a bullet marker (- or •)",
            "- Each bullet point should be on a separate line",
            "- Use concise, clear statements (one sentence per bullet point)",
            "- Do not use numbered lists, only bullet points",
            "",
            "Return ONLY the bullet-point text, no additional formatting or preamble."
        ])
        
        prompt = "\n".join(prompt_parts)
        
        # Call LLM service
        summary = await llm_service.generate_text(
            prompt=prompt,
            max_tokens=400,
            temperature=0.3
        )
        
        if not summary or summary.strip() == "":
            return DataQualitySummaryResponse(
                success=False,
                error="LLM returned empty response"
            )
        
        logger.info("Data quality summary generated successfully")
        return DataQualitySummaryResponse(
            success=True,
            summary=summary.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating data quality summary: {str(e)}")
        return DataQualitySummaryResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-target-definition", response_model=TargetDefinitionResponse)
async def generate_target_definition(request: TargetDefinitionRequest):
    """
    Generate target variable definition - first from data dictionary, fallback to LLM
    """
    try:
        logger.info(f"Generating definition for target variable: {request.target_variable}")
        
        definition = None
        
        # Try to find definition in data dictionary first
        if request.data_dictionary:
            logger.info("Searching for target variable in data dictionary")
            try:
                # Parse data dictionary (assuming CSV format)
                import csv
                from io import StringIO
                
                csv_reader = csv.reader(StringIO(request.data_dictionary))
                rows = list(csv_reader)
                
                if len(rows) > 1:  # Has header + data
                    # Search for target variable in first column
                    for row in rows[1:]:  # Skip header
                        if len(row) >= 2 and row[0].strip().lower() == request.target_variable.lower():
                            definition = row[1].strip()
                            logger.info(f"Found definition in data dictionary: {definition}")
                            break
                
            except Exception as e:
                logger.warning(f"Failed to parse data dictionary: {str(e)}")
        
        # If not found in data dictionary, use LLM
        if not definition:
            logger.info("Generating definition using LLM")
            
            prompt_parts = [
                "You are a data analyst providing a one-line definition for a target variable.",
                f"Target Variable: {request.target_variable}",
                "",
                "Context:",
            ]
            
            if request.problem_statement:
                prompt_parts.append(f"Problem Statement: {request.problem_statement}")
            
            prompt_parts.append(f"Available Columns: {', '.join(request.columns)}")
            
            prompt_parts.extend([
                "",
                "Please provide a ONE-LINE definition (one sentence) that:",
                "1. Clearly explains what this target variable represents",
                "2. Is concise and professional",
                "3. Uses business/domain language",
                "",
                "Return ONLY the definition, no additional text or preamble."
            ])
            
            prompt = "\n".join(prompt_parts)
            
            # Call LLM service
            definition = await llm_service.generate_text(
                prompt=prompt,
                max_tokens=150,
                temperature=0.3
            )
            
            if not definition or definition.strip() == "":
                return TargetDefinitionResponse(
                    success=False,
                    error="LLM returned empty response"
                )
            
            definition = definition.strip()
        
        logger.info("Target definition generated successfully")
        return TargetDefinitionResponse(
            success=True,
            definition=definition
        )
        
    except Exception as e:
        logger.error(f"Error generating target definition: {str(e)}")
        return TargetDefinitionResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-model-objective", response_model=ModelObjectiveResponse)
async def generate_model_objective(request: ModelObjectiveRequest):
    """
    Generate a 3-line objective text explaining why modeling is done
    """
    try:
        logger.info("Generating model objective text")
        
        # Build the prompt
        prompt_parts = [
            "You are a data analyst creating a model objective statement for model documentation.",
            "Generate a 3-line (approximately 3 sentences) text that explains why modeling is being done.",
            "The text should describe the objective of the modeling exercise based on the provided context.",
            "",
        ]
        
        # Add available inputs
        if request.project_description:
            prompt_parts.append(f"Project Description: {request.project_description}")
        
        if request.problem_statement:
            prompt_parts.append(f"Problem Statement: {request.problem_statement}")
        
        if request.data_summary:
            prompt_parts.append(f"Data Summary: {request.data_summary}")
        
        if request.target_variable_name:
            prompt_parts.append(f"Target Variable Name: {request.target_variable_name}")
        
        if request.target_definition:
            prompt_parts.append(f"Target Variable Definition: {request.target_definition}")
        
        prompt_parts.extend([
            "",
            "Please provide a 3-line objective statement (approximately 3 sentences) that:",
            "1. Explains why modeling is being done",
            "2. Describes the business objective or problem the model aims to solve",
            "3. Connects the modeling objective to the target variable and business context",
            "4. Is written in clear, professional language",
            "5. Is concise and informative",
            "",
            "Return ONLY the objective text (3 lines), no additional formatting or preamble."
        ])
        
        prompt = "\n".join(prompt_parts)
        
        # Call LLM service
        objective = await llm_service.generate_text(
            prompt=prompt,
            max_tokens=200,
            temperature=0.3
        )
        
        if not objective or objective.strip() == "":
            return ModelObjectiveResponse(
                success=False,
                error="LLM returned empty response"
            )
        
        logger.info("Model objective generated successfully")
        return ModelObjectiveResponse(
            success=True,
            objective=objective.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating model objective: {str(e)}")
        return ModelObjectiveResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-monotonicity-summary", response_model=MonotonicitySummaryResponse)
async def generate_monotonicity_summary(request: MonotonicitySummaryRequest):
    """
    Generate a 2-3 line writeup explaining monotonicity stats from the summary table
    """
    try:
        logger.info(f"Generating monotonicity summary for {len(request.models)} models")
        
        # Build the summary table data for the prompt
        table_data = []
        for model in request.models:
            model_name = model.get('modelName', 'Unknown')
            monotonicity_score = model.get('monotonicityScore', 0)
            ks_statistic = model.get('ksStatistic', 0)
            lift_top_decile = model.get('liftTopDecile')
            auc = model.get('auc', 0)
            gini = model.get('gini', 0)
            psi_value = None
            if model.get('psi') and isinstance(model.get('psi'), dict):
                psi_value = model.get('psi', {}).get('value')
            elif model.get('psi') is not None:
                psi_value = model.get('psi')
            
            lift_str = f"{lift_top_decile:.2f}x" if lift_top_decile is not None else "N/A"
            auc_gini_str = f"AUC {auc:.3f} / Gini {gini:.3f}"
            psi_str = f"{psi_value:.4f}" if psi_value is not None else "N/A"
            
            table_data.append({
                'model': model_name,
                'monotonicity_score': f"{monotonicity_score:.2f}%",
                'ks_statistic': f"{ks_statistic:.3f}",
                'lift': lift_str,
                'auc_gini': auc_gini_str,
                'psi': psi_str
            })
        
        # Build the prompt
        prompt_parts = [
            "You are a data analyst creating a summary writeup for model monotonicity analysis.",
            "Generate a 2-3 line (approximately 2-3 sentences) text that explains the overall monotonicity statistics across all models.",
            "The writeup should summarize the key findings from the monotonicity analysis table.",
            "",
            "Monotonicity Summary Table:",
            ""
        ]
        
        # Add table data
        prompt_parts.append("Model | Monotonicity Score | KS Statistic | Lift | AUC/Gini | PSI")
        prompt_parts.append("-" * 80)
        for row in table_data:
            prompt_parts.append(f"{row['model']} | {row['monotonicity_score']} | {row['ks_statistic']} | {row['lift']} | {row['auc_gini']} | {row['psi']}")
        
        prompt_parts.extend([
            "",
            "Please provide a 2-3 line summary (approximately 2-3 sentences) that:",
            "1. Explains the overall monotonicity performance across models",
            "2. Highlights key patterns or notable findings from the statistics",
            "3. Is written in clear, professional language",
            "4. Is concise and informative",
            "",
            "Return ONLY the summary text (2-3 lines), no additional formatting or preamble."
        ])
        
        prompt = "\n".join(prompt_parts)
        
        # Call LLM service
        writeup = await llm_service.generate_text(
            prompt=prompt,
            max_tokens=150,
            temperature=0.3
        )
        
        if not writeup or writeup.strip() == "":
            return MonotonicitySummaryResponse(
                success=False,
                error="LLM returned empty response"
            )
        
        logger.info("Monotonicity summary generated successfully")
        return MonotonicitySummaryResponse(
            success=True,
            writeup=writeup.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating monotonicity summary: {str(e)}")
        return MonotonicitySummaryResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/calculate-event-rate", response_model=EventRateResponse)
async def calculate_event_rate(request: EventRateRequest):
    """
    Calculate event rate for target variable (after removing duplicates)
    Uses the ENTIRE/MASTER dataset, not the dev/hold split
    """
    try:
        logger.info(f"Calculating event rate for dataset: {request.dataset_id}, target: {request.target_variable}")
        
        # Import dataframe state manager
        from app.services.dataframe_state_manager import dataframe_state_manager
        
        # Get the ORIGINAL FULL dataframe (before any split or transformations)
        # This is what was uploaded by the user
        df = dataframe_state_manager._full_dataframes.get(request.dataset_id)
        
        if df is not None:
            logger.info(f"Using original full dataframe (before split/transformations), shape: {df.shape}")
        
        if df is None:
            return EventRateResponse(
                success=False,
                error="Dataset not found"
            )
        
        if request.target_variable not in df.columns:
            return EventRateResponse(
                success=False,
                error=f"Target variable '{request.target_variable}' not found in dataset"
            )
        
        # Remove duplicate rows
        df_deduped = df.drop_duplicates()
        total_count = len(df_deduped)
        
        logger.info(f"Removed {len(df) - total_count} duplicate rows. Remaining: {total_count}")
        
        # Count events (where target = 1 or True)
        target_col = df_deduped[request.target_variable]
        
        # Handle different representations of "event"
        # Could be 1, True, "1", "True", "yes", "Yes", etc.
        event_count = 0
        
        # Try numeric first
        try:
            event_count = int((target_col == 1).sum())
        except:
            pass
        
        # If no numeric 1s found, try boolean True
        if event_count == 0:
            try:
                event_count = int((target_col == True).sum())
            except:
                pass
        
        # If still no events, try string representations
        if event_count == 0:
            try:
                event_count = int(target_col.astype(str).str.lower().isin(['1', 'true', 'yes']).sum())
            except:
                pass
        
        percentage = (event_count / total_count * 100) if total_count > 0 else 0
        
        logger.info(f"Event rate calculated on ENTIRE dataset: {event_count}/{total_count} ({percentage:.2f}%)")
        
        return EventRateResponse(
            success=True,
            event_count=int(event_count),
            total_count=int(total_count),
            percentage=float(percentage)
        )
        
    except Exception as e:
        logger.error(f"Error calculating event rate: {str(e)}")
        return EventRateResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/get-sampling-plan", response_model=SamplingPlanResponse)
async def get_sampling_plan(request: SamplingPlanRequest):
    """
    Get sampling plan data for Train/Hold splits with event rates
    """
    try:
        logger.info(f"Getting sampling plan for dataset: {request.dataset_id}")
        
        # Import dataframe state manager
        from app.services.dataframe_state_manager import dataframe_state_manager
        
        # Check if dataset has split indices
        has_split_indices = request.dataset_id in dataframe_state_manager._split_indices
        
        logger.info(f"Dataset has split indices: {has_split_indices}")
        
        # Get the ORIGINAL FULL dataframe (before any split or transformations)
        # Train and Hold should be splits of this original entire dataset
        master_df = dataframe_state_manager._full_dataframes.get(request.dataset_id)
        
        if master_df is not None:
            logger.info(f"Using original full dataframe (before split/transformations), shape: {master_df.shape}")
        
        if master_df is None:
            return SamplingPlanResponse(
                success=False,
                error="Dataset not found"
            )
        
        if request.target_variable not in master_df.columns:
            return SamplingPlanResponse(
                success=False,
                error=f"Target variable '{request.target_variable}' not found in dataset"
            )
        
        # Get train dataframe (filter from master using indices)
        train_df = master_df
        if has_split_indices:
            train_indices = dataframe_state_manager._split_indices[request.dataset_id].get('train', [])
            if len(train_indices) > 0:
                valid_train_indices = train_indices[train_indices < len(master_df)]
                train_df = master_df.iloc[valid_train_indices]
                logger.info(f"Filtered train dataframe, shape: {train_df.shape}")
        else:
            # If no split, use entire master as train
            logger.info(f"No split indices, using entire dataset as train, shape: {train_df.shape}")
        dev_df = train_df  # alias for backward compat
        
        # Calculate train/dev stats
        train_total = len(dev_df)
        train_target = dev_df[request.target_variable]
        
        # Count events in train
        train_event_count = 0
        try:
            train_event_count = int((train_target == 1).sum())
        except:
            pass
        if train_event_count == 0:
            try:
                train_event_count = int((train_target == True).sum())
            except:
                pass
        if train_event_count == 0:
            try:
                train_event_count = int(train_target.astype(str).str.lower().isin(['1', 'true', 'yes']).sum())
            except:
                pass
        
        train_event_rate = (train_event_count / train_total * 100) if train_total > 0 else 0
        
        train_data = {
            'total': int(train_total),
            'event_count': int(train_event_count),
            'event_rate': float(train_event_rate)
        }
        
        logger.info(f"Train data: {train_data}")
        
        # Calculate validation stats (if split exists)
        hold_data = {
            'total': 0,
            'event_count': 0,
            'event_rate': 0.0
        }
        
        if has_split_indices:
            validation_indices = dataframe_state_manager._split_indices[request.dataset_id].get('validation', [])
            if len(validation_indices) > 0:
                valid_validation_indices = validation_indices[validation_indices < len(master_df)]
                hold_df = master_df.iloc[valid_validation_indices]
                logger.info(f"Filtered validation dataframe, shape: {hold_df.shape}")
            else:
                hold_df = pd.DataFrame()
            
            if len(hold_df) > 0:
                hold_total = len(hold_df)
                hold_target = hold_df[request.target_variable]
                
                # Count events in hold
                hold_event_count = 0
                try:
                    hold_event_count = int((hold_target == 1).sum())
                except:
                    pass
                if hold_event_count == 0:
                    try:
                        hold_event_count = int((hold_target == True).sum())
                    except:
                        pass
                if hold_event_count == 0:
                    try:
                        hold_event_count = int(hold_target.astype(str).str.lower().isin(['1', 'true', 'yes']).sum())
                    except:
                        pass
                
                hold_event_rate = (hold_event_count / hold_total * 100) if hold_total > 0 else 0
                
                hold_data = {
                    'total': int(hold_total),
                    'event_count': int(hold_event_count),
                    'event_rate': float(hold_event_rate)
                }
                
                logger.info(f"Hold data: {hold_data}")
        
        return SamplingPlanResponse(
            success=True,
            has_split=has_split_indices,
            train=train_data,
            hold=hold_data
        )
        
    except Exception as e:
        logger.error(f"Error getting sampling plan: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return SamplingPlanResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-sampling-plan-writeup", response_model=SamplingPlanWriteupResponse)
async def generate_sampling_plan_writeup(request: SamplingPlanWriteupRequest):
    """
    Generate LLM writeup for sampling plan section
    """
    try:
        sampling_plan = request.sampling_plan
        has_split = sampling_plan.get('hasSplit', False)
        train_data = sampling_plan.get('train', {})
        hold_data = sampling_plan.get('hold', {}) if has_split else {}
        
        train_total = train_data.get('total', 0)
        train_event_count = train_data.get('eventCount', 0)
        train_event_rate = train_data.get('eventRate', 0)
        
        if has_split:
            hold_total = hold_data.get('total', 0)
            hold_event_count = hold_data.get('eventCount', 0)
            hold_event_rate = hold_data.get('eventRate', 0)
            
            # Calculate split ratio
            total_samples = train_total + hold_total
            train_pct = (train_total / total_samples * 100) if total_samples > 0 else 0
            hold_pct = (hold_total / total_samples * 100) if total_samples > 0 else 0
            
            prompt = f"""You are a data scientist writing model documentation. Write a concise 3-4 line paragraph explaining the sampling plan for model training and validation.

Sampling Plan Details:
- Train and Hold samples are split in a {train_pct:.0f}-{hold_pct:.0f}% ratio
- Train sample: {train_total:,} total records, {train_event_count:,} events, {train_event_rate:.2f}% event rate
- Hold sample: {hold_total:,} total records, {hold_event_count:,} events, {hold_event_rate:.2f}% event rate
- Split is done randomly ensuring similar event rate across both samples

Write in a professional, modeller's voice explaining why this split ensures robust validation testing. Mention the event rate similarity between train and hold samples. Keep it concise (3-4 lines maximum)."""
        else:
            prompt = f"""You are a data scientist writing model documentation. Write a concise 3-4 line paragraph explaining the sampling plan.

Sampling Plan Details:
- Train sample: {train_total:,} total records, {train_event_count:,} events, {train_event_rate:.2f}% event rate
- No hold sample was created

Write in a professional, modeller's voice. Keep it concise (3-4 lines maximum)."""
        
        writeup = await llm_service.generate_text(prompt=prompt)
        if writeup:
            # Clean up the writeup
            writeup = writeup.strip()
            writeup = writeup.replace('**', '').replace('*', '')
        
        return SamplingPlanWriteupResponse(
            success=True,
            writeup=writeup if writeup else None
        )
        
    except Exception as e:
        logger.error(f"Error generating sampling plan writeup: {str(e)}")
        return SamplingPlanWriteupResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-model-validation-writeup", response_model=ModelValidationWriteupResponse)
async def generate_model_validation_writeup(request: ModelValidationWriteupRequest):
    """
    Generate LLM writeup for model validation section
    """
    try:
        model_validation = request.model_validation
        best_model = model_validation.get('bestModel', {})
        model_name = best_model.get('modelName', 'Unknown')
        metrics = best_model.get('metrics', {})
        data_summary_content = request.data_summary or 'Not provided'
        
        # Extract key metrics
        accuracy = metrics.get('accuracy', 0) * 100
        precision = metrics.get('precision', 0)
        recall = metrics.get('recall', 0)
        f1_score = metrics.get('f1Score', 0)
        auc_roc = metrics.get('aucRoc', 0)
        auc_pr = metrics.get('aucPr', 0)
        log_loss = metrics.get('logLoss', 0)
        
        prompt = f"""You are a data scientist writing model documentation. Write a concise 3-4 line paragraph explaining why the best performing model is well-suited for this use case.

Model Information:
- Model Name: {model_name}
- Performance Metrics on Hold Dataset:
  * Accuracy: {accuracy:.2f}%
  * Precision: {precision:.4f}
  * Recall: {recall:.4f}
  * F1 Score: {f1_score:.4f}
  * AUC-ROC: {auc_roc:.4f}
  * AUC-PR: {auc_pr:.4f}
  * Log Loss: {log_loss:.4f}

Data Summary Context:
{data_summary_content}

Write in a professional, modeller's voice explaining why this model is well-suited for the problem. Reference the key performance metrics and how they align with the data characteristics and business objectives. Keep it concise (3-4 lines maximum)."""
        
        writeup = await llm_service.generate_text(prompt=prompt)
        if writeup:
            # Clean up the writeup
            writeup = writeup.strip()
            writeup = writeup.replace('**', '').replace('*', '')
        
        return ModelValidationWriteupResponse(
            success=True,
            writeup=writeup if writeup else None
        )
        
    except Exception as e:
        logger.error(f"Error generating model validation writeup: {str(e)}")
        return ModelValidationWriteupResponse(
            success=False,
            error=str(e)
        )


class ModelPerformanceRequest(BaseModel):
    model_id: str
    dataset_id: str
    data_dictionary: Optional[str] = None
    variable_categories: Optional[Dict[str, str]] = None  # feature_name -> category
    category_colors: Optional[Dict[str, str]] = None  # category -> color


class FeatureInfo(BaseModel):
    feature_name: str
    importance: float  # SHAP value with sign
    description: str  # From data dictionary or empty


class ModelPerformanceResponse(BaseModel):
    success: bool
    total_features: int
    used_features: List[str]
    top_features: List[FeatureInfo]
    category_distribution: Dict[str, int]  # category -> count
    category_colors: Dict[str, str]  # category -> color
    error: Optional[str] = None


@documentation_router.post("/documentation/get-model-performance", response_model=ModelPerformanceResponse)
async def get_model_performance(request: ModelPerformanceRequest):
    """
    Get model performance data including used features, SHAP importance, and descriptions
    """
    try:
        from app.models.model_evaluation_database import model_evaluation_db
        import pandas as pd
        
        logger.info(f"Getting model performance for model {request.model_id}, dataset {request.dataset_id}")
        logger.info(f"  - Has data dictionary: {request.data_dictionary is not None}")
        logger.info(f"  - Variable categories: {len(request.variable_categories)} features")
        
        # Get model evaluation data from database
        evaluation_data = model_evaluation_db.get_model_evaluation(
            request.model_id, 
            include_explainability=True,
            include_pdp=False
        )
        
        if not evaluation_data:
            return ModelPerformanceResponse(
                success=False,
                total_features=0,
                used_features=[],
                top_features=[],
                category_distribution={},
                category_colors={},
                error=f"Model {request.model_id} not found"
            )
        
        # Get used features from model evaluation data
        used_features = evaluation_data.get('used_features', [])
        
        if not used_features:
            # Fallback: try to get from feature_importance
            feature_importance = evaluation_data.get('feature_importance', [])
            if feature_importance:
                used_features = [f.get('feature_name') for f in feature_importance if f.get('feature_name')]
                logger.info(f"Got {len(used_features)} used features from feature_importance")
        else:
            logger.info(f"Got {len(used_features)} used features from evaluation_data.used_features")
        
        logger.info(f"Total used features: {len(used_features)}")
        if used_features:
            logger.info(f"Sample features: {used_features[:5]}")
        
        # Get SHAP importance values from explainability data
        explainability_data = evaluation_data.get('explainability_data', [])
        
        # Build feature importance map from SHAP summary data
        # Store signed SHAP values (mean of actual values, not absolute) for proper sign display
        feature_importance_map = {}
        for entry in explainability_data:
            if entry.get('data_type') == 'shap_summary' and entry.get('feature_name'):
                feature_name = entry.get('feature_name')
                values = entry.get('values', [])
                
                # Calculate mean SHAP value (signed, not absolute) for proper sign display
                if isinstance(values, list) and len(values) > 0:
                    import numpy as np
                    mean_shap = float(np.mean(values))  # Signed mean, not absolute
                    feature_importance_map[feature_name] = mean_shap
                else:
                    # Fallback: try metadata if values not available
                    metadata = entry.get('metadata', {})
                    if isinstance(metadata, dict):
                        # Use mean_abs as fallback, but note it's absolute
                        mean_abs_shap = metadata.get('mean_abs', 0)
                        feature_importance_map[feature_name] = mean_abs_shap
        
        logger.info(f"Found SHAP data for {len(feature_importance_map)} features")
        
        # Parse data dictionary if provided
        feature_descriptions = {}
        if request.data_dictionary:
            try:
                # Try to parse as CSV
                from io import StringIO
                df_dict = pd.read_csv(StringIO(request.data_dictionary))
                
                # Find the description column
                desc_col = None
                possible_desc_cols = ['description', 'datadesc', 'data_description', 'data_desc']
                for col in df_dict.columns:
                    if col.lower() in possible_desc_cols:
                        desc_col = col
                        break
                
                # If no description column found, use 2nd column
                if not desc_col and len(df_dict.columns) >= 2:
                    desc_col = df_dict.columns[1]
                
                if desc_col:
                    # Assume first column has feature names
                    feature_col = df_dict.columns[0]
                    for _, row in df_dict.iterrows():
                        feature_name = str(row[feature_col]).strip()
                        description = str(row[desc_col]).strip() if pd.notna(row[desc_col]) else ''
                        feature_descriptions[feature_name] = description
                    
                    logger.info(f"Parsed {len(feature_descriptions)} feature descriptions from data dictionary")
            except Exception as e:
                logger.warning(f"Failed to parse data dictionary: {e}")
        
        # Build top features list
        top_features = []
        for feature in used_features:
            importance = feature_importance_map.get(feature, 0)
            description = feature_descriptions.get(feature, '')
            top_features.append({
                'feature_name': feature,
                'importance': float(importance),
                'description': description
            })
        
        # Sort by absolute importance value (descending)
        top_features.sort(key=lambda x: abs(x['importance']), reverse=True)
        
        # Calculate category distribution
        category_distribution = {}
        if request.variable_categories:
            for feature in used_features:
                category = request.variable_categories.get(feature, 'Unknown')
                category_distribution[category] = category_distribution.get(category, 0) + 1
        
        logger.info(f"Category distribution: {category_distribution}")
        
        return ModelPerformanceResponse(
            success=True,
            total_features=len(used_features),
            used_features=used_features,
            top_features=top_features,
            category_distribution=category_distribution,
            category_colors=request.category_colors or {}
        )
            
    except Exception as e:
        logger.error(f"Error getting model performance: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return ModelPerformanceResponse(
            success=False,
            total_features=0,
            used_features=[],
            top_features=[],
            category_distribution={},
            category_colors={},
            error=str(e)
        )


@documentation_router.post("/documentation/generate-segmentation-understanding", response_model=SegmentationUnderstandingResponse)
async def generate_segmentation_understanding(request: SegmentationUnderstandingRequest):
    """
    Generate an 8-line explanation of why segmentation was done and why each segment was created.
    Uses LLM to analyze segmentation data, event rates, and chart characteristics.
    """
    try:
        logger.info("Generating segmentation understanding")
        
        # Prepare segment details for the prompt
        segment_details = []
        for i, segment in enumerate(request.segments):
            segment_num = i + 1
            rule = segment.get('rule', 'No rule specified')
            total = segment.get('total', 0)
            event_rate = segment.get('eventRate', 0)  # Already in percentage
            segment_dist = segment.get('segmentDistribution', 0)  # Already in percentage
            
            # Get IV data if available
            iv_data = ""
            if request.iv_report and request.iv_report.get('table') and i < len(request.iv_report['table']):
                iv_row = request.iv_report['table'][i]
                iv_data = f" WoE: {iv_row.get('woe', 0):.3f}, IV Contribution: {iv_row.get('iv_contribution', 0):.4f}, Bad Rate: {iv_row.get('bad_rate', 0)*100:.1f}%"
            
            segment_details.append(
                f"Segment {segment_num}: Rule='{rule}', Total Records={total:,}, Event Rate={event_rate:.2f}%, "
                f"Segment Distribution={segment_dist:.2f}%{iv_data}"
            )
        
        # Build comprehensive prompt
        prompt = f"""You are analyzing segmentation results for a predictive modeling project. Based on the following information, the focus has to be on explaining why segmentation was performed and why each segment was created.

DATA SUMMARY (from project objectives):
{request.data_summary}

SEGMENTATION DETAILS:
{chr(10).join(segment_details)}

SEGMENT STATISTICS:
- Segment Sizes: {', '.join([f'Segment {i+1}: {size:,}' for i, size in enumerate(request.segment_sizes)])}
- Segment Proportions: {', '.join([f'Segment {i+1}: {prop:.1f}%' for i, prop in enumerate(request.segment_proportions)])}
- Event Rates: {', '.join([f'Segment {i+1}: {rate:.2f}%' for i, rate in enumerate(request.event_rates)])}

ANALYSIS RULES:
1. Compare event rates across segments - significant differences indicate why segmentation was valuable
2. Segment sizes and proportions show the distribution of the population
3. Higher event rates in a segment suggest higher risk/event likelihood
4. Lower event rates suggest lower risk/event likelihood
5. Segment rules (conditions) explain the characteristics that define each segment
6. If IV (Information Value) data is available, WoE (Weight of Evidence) and IV contributions indicate predictive power
7. Segments with high IV contributions are more predictive and valuable

REQUIREMENTS:
- Start with 1-2 sentences explaining why segmentation was done overall (based on event rate differences, business need, etc.)
- Then provide bullet points for each segment (one bullet per segment) explaining why it was created based on:
  * Its event rate compared to other segments
  * Its size/proportion in the population
  * The characteristics defined by its rules
  * Its predictive power (if IV data available)
- Format: Use "- " prefix for each bullet point (e.g., "- Segment 1: ...")
- Each bullet point should be maximum 25 words
- Write in clear, professional language suitable for model documentation
- Focus on business and statistical rationale for segmentation

OUTPUT FORMAT EXAMPLE:
Segmentation was performed to identify distinct risk groups within the population based on observable characteristics and event rate differences.

- Segment 1: [maximum 25 words]
- Segment 2: [maximum 25 words]
- Segment 3: [maximum 25 words]

Generate the explanation following this exact format:"""
        
        # Call LLM service
        understanding = await llm_service.generate_text(prompt)
        
        if not understanding:
            return SegmentationUnderstandingResponse(
                success=False,
                error="Failed to generate segmentation understanding"
            )
        
        logger.info("Segmentation understanding generated successfully")
        return SegmentationUnderstandingResponse(
            success=True,
            understanding=understanding.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating segmentation understanding: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return SegmentationUnderstandingResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/get-quality-check-plan", response_model=QualityCheckPlanResponse)
async def get_quality_check_plan(request: QualityCheckPlanRequest):
    """
    Get the analysis plan (quality check plan) from MessageState for a dataset.
    Returns the plan data with missing_values, outliers, and duplicates.
    """
    try:
        logger.info(f"Getting quality check plan for dataset: {request.dataset_id}")
        
        # Get MessageState for the dataset
        state = message_state_manager.create_or_load_state(request.dataset_id, "")
        plan_data = state.get("plan", "")
        
        if not plan_data:
            return QualityCheckPlanResponse(
                success=False,
                error="No analysis plan found. Please run data quality checks first."
            )
        
        # Parse the plan JSON string
        try:
            if isinstance(plan_data, str):
                plan_dict = json.loads(plan_data)
            else:
                plan_dict = plan_data
            
            # Convert plan to table format with Issue, Variable, Observation, Treatment columns
            table_rows = []
            
            # Process missing_values
            if plan_dict.get("missing_values"):
                for item in plan_dict["missing_values"]:
                    table_rows.append({
                        "Issue": "Missing Values",
                        "Variable": item.get("variable", item.get("name", "Unknown")),
                        "Observation": item.get("detection", "N/A"),
                        "Treatment": item.get("treatment", item.get("final_treatment", "N/A"))
                    })
            
            # Process outliers
            if plan_dict.get("outliers"):
                for item in plan_dict["outliers"]:
                    table_rows.append({
                        "Issue": "Outliers",
                        "Variable": item.get("variable", item.get("name", "Unknown")),
                        "Observation": item.get("detection", "N/A"),
                        "Treatment": item.get("treatment", item.get("final_treatment", "N/A"))
                    })
            
            # Process duplicates
            if plan_dict.get("duplicates"):
                for item in plan_dict["duplicates"]:
                    table_rows.append({
                        "Issue": "Duplicates",
                        "Variable": item.get("variable", item.get("name", "Dataset")),
                        "Observation": item.get("detection", "N/A"),
                        "Treatment": item.get("treatment", item.get("final_treatment", "N/A"))
                    })
            
            logger.info(f"Quality check plan retrieved: {len(table_rows)} rows")
            return QualityCheckPlanResponse(
                success=True,
                plan={"table": table_rows}
            )
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse plan data: {e}")
            return QualityCheckPlanResponse(
                success=False,
                error=f"Failed to parse plan data: {str(e)}"
            )
        
    except Exception as e:
        logger.error(f"Error getting quality check plan: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return QualityCheckPlanResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/get-column-stats", response_model=ColumnStatsResponse)
async def get_column_stats(request: ColumnStatsRequest):
    """
    Get column statistics for the processed dataset.
    Returns the same column stats table that appears after code execution.
    """
    try:
        logger.info(f"Getting column stats for dataset: {request.dataset_id}")
        
        # Get processed DataFrame from state manager
        processed_df = dataframe_state_manager.get_dataframe(request.dataset_id)
        
        if processed_df is None:
            # Fallback to MessageState
            try:
                state = message_state_manager.create_or_load_state(request.dataset_id, "")
                processed_df = state.get("datasetFile")
            except:
                pass
        
        if processed_df is None:
            return ColumnStatsResponse(
                success=False,
                error="No processed dataset found. Please run code execution first."
            )
        
        # Calculate column statistics - import locally to avoid circular dependency
        from app.api.routes import calculate_column_info
        columns_info = calculate_column_info(processed_df)
        logger.info(f"Column stats calculated: {len(columns_info)} columns")
        
        # Convert to table format matching the UI
        stats_data = []
        for col_info in columns_info:
            stats_row = {
                'Column': col_info.column_name,
                'Type': col_info.data_type,
                'Missing': col_info.missing_count,
                'Unique': col_info.unique_count,
                'Mean': col_info.mean,
                'Median': col_info.median,
                'Mode': col_info.mode,
                'Std': col_info.standard_deviation,
                'Var': col_info.variance,
                'Min': col_info.min_value,
                'p5%': col_info.percentile_5,
                'p25%': col_info.percentile_25,
                'p50%': col_info.percentile_50,
                'p75%': col_info.percentile_75,
                'p95%': col_info.percentile_95,
                'p99%': col_info.percentile_99,
                'Max': col_info.max_value
            }
            stats_data.append(stats_row)
        
        return ColumnStatsResponse(
            success=True,
            stats=stats_data
        )
        
    except Exception as e:
        logger.error(f"Error getting column stats: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return ColumnStatsResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-quality-changes-writeup", response_model=QualityChangesWriteupResponse)
async def generate_quality_changes_writeup(request: QualityChangesWriteupRequest):
    """
    Generate a 3-4 line write-up explaining what quality level changes have happened.
    Uses both the quality check plan and column stats tables as inputs.
    """
    try:
        logger.info("Generating quality changes write-up")
        
        # Format the quality check plan table for the prompt
        plan_table_text = ""
        if request.quality_check_plan.get("table"):
            plan_table_text = "\n".join([
                f"Issue: {row.get('Issue', 'N/A')}, Variable: {row.get('Variable', 'N/A')}, "
                f"Observation: {row.get('Observation', 'N/A')}, Treatment: {row.get('Treatment', 'N/A')}"
                for row in request.quality_check_plan["table"]
            ])
        
        # Format the column stats summary for the prompt
        stats_summary = f"Total columns analyzed: {len(request.column_stats)}\n"
        if request.column_stats:
            # Count columns with missing values
            cols_with_missing = sum(1 for row in request.column_stats if row.get('Missing', 0) > 0)
            stats_summary += f"Columns with missing values: {cols_with_missing}\n"
            
            # Sample a few key stats
            sample_stats = request.column_stats[:3]
            for stat in sample_stats:
                stats_summary += f"{stat.get('Column', 'N/A')}: Type={stat.get('Type', 'N/A')}, Missing={stat.get('Missing', 0)}, Unique={stat.get('Unique', 'N/A')}\n"
        
        # Build prompt
        prompt = f"""You are analyzing data quality treatment results for a predictive modeling project. Based on the following information, write a concise 3-4 line summary explaining what quality level changes have been implemented.

QUALITY CHECK PLAN (Issues Identified and Treatments Planned):
{plan_table_text}

COLUMN STATISTICS (After Treatment Implementation):
{stats_summary}

REQUIREMENTS:
- Write 3-4 lines maximum
- Explain in simple terms what quality changes have happened
- Focus on the main treatments applied (missing value imputation, outlier capping, duplicate removal, etc.)
- Mention the scope of changes (how many variables were affected)
- Write in clear, professional language suitable for model documentation

Generate the summary:"""
        
        # Call LLM service
        writeup = await llm_service.generate_text(prompt)
        
        if not writeup:
            return QualityChangesWriteupResponse(
                success=False,
                error="Failed to generate quality changes write-up"
            )
        
        logger.info("Quality changes write-up generated successfully")
        return QualityChangesWriteupResponse(
            success=True,
            writeup=writeup.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating quality changes write-up: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return QualityChangesWriteupResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-feature-engineering-writeup", response_model=FeatureEngineeringWriteupResponse)
async def generate_feature_engineering_writeup(request: FeatureEngineeringWriteupRequest):
    """
    Generate a 3-4 line explanation of the transformations done on variables using LLM.
    """
    try:
        logger.info(f"Generating feature engineering write-up for {len(request.transformed_variables)} transformed variables")
        
        if not request.transformed_variables or len(request.transformed_variables) == 0:
            return FeatureEngineeringWriteupResponse(
                success=False,
                error="No transformed variables provided"
            )
        
        # Build the prompt
        prompt_parts = [
            "You are a data scientist creating a concise explanation of feature engineering transformations for model documentation.",
            "Generate a 3-4 line plain English explanation that describes the transformations applied to variables.",
            "",
            f"Number of transformed variables: {len(request.transformed_variables)}",
            "",
            "Transformed Variables:",
        ]
        
        # Add sample of transformed variables (limit to first 10 to avoid token limits)
        sample_vars = request.transformed_variables[:10]
        for var in sample_vars:
            var_name = var.get('new_variable_name', '')
            var_type = var.get('var_type', '')
            method = var.get('transformation_methods', '')
            prompt_parts.append(f"- {var_name} ({var_type}): {method}")
        
        if len(request.transformed_variables) > 10:
            prompt_parts.append(f"... and {len(request.transformed_variables) - 10} more transformations")
        
        prompt_parts.extend([
            "",
            "Please provide a 3-4 line explanation (approximately 3-4 sentences) that:",
            "1. Describes the types of transformations applied (e.g., WOE, Log, One Hot Encoding)",
            "2. Mentions the purpose or benefit of these transformations",
            "3. Is written in clear, professional language",
            "4. Is concise and informative",
            "",
            "Return ONLY the explanation text, no additional formatting or preamble."
        ])
        
        prompt = "\n".join(prompt_parts)
        
        # Call LLM service
        writeup = await llm_service.generate_text(
            prompt=prompt,
            max_tokens=200,
            temperature=0.3
        )
        
        if not writeup or writeup.strip() == "":
            return FeatureEngineeringWriteupResponse(
                success=False,
                error="LLM returned empty response"
            )
        
        logger.info("Feature engineering write-up generated successfully")
        return FeatureEngineeringWriteupResponse(
            success=True,
            writeup=writeup.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating feature engineering write-up: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return FeatureEngineeringWriteupResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/generate-decile-progression-writeup", response_model=DecileProgressionWriteupResponse)
async def generate_decile_progression_writeup(request: DecileProgressionWriteupRequest):
    """
    Generate a 2-4 line explanation of the decile progression for monotonicity analysis using LLM.
    """
    try:
        logger.info(f"Generating decile progression write-up for model {request.model_name}")
        
        if not request.deciles or len(request.deciles) == 0:
            return DecileProgressionWriteupResponse(
                success=False,
                error="No decile data provided"
            )
        
        # Build the prompt
        prompt_parts = [
            "You are a data scientist creating a concise explanation of decile progression for model documentation.",
            "Generate a 2-4 line plain English explanation that describes the decile progression and monotonicity behavior.",
            "",
            f"Model: {request.model_name}",
            f"Monotonicity Score: {request.monotonicity_score:.2f}%",
            "",
            "Decile Table (showing progression from lowest to highest risk):",
            "Decile | Count | Bads | Goods | Bad Rate | Avg Score | Lift | Cum Bad Rate",
            "-" * 80,
        ]
        
        # Add full decile table data
        for decile in request.deciles:
            decile_num = decile.get('Decile', 0)
            count = decile.get('Count', 0)
            bads = decile.get('Bads', 0)
            goods = decile.get('Goods', 0)
            bad_rate = decile.get('Bad_Rate', 0)
            avg_score = decile.get('Avg_Score', 0)
            lift = decile.get('Lift', 0)
            cum_bad_rate = decile.get('Cum_Bad_Rate', 0)
            
            bad_rate_pct = f"{(bad_rate * 100):.2f}%" if bad_rate is not None else "N/A"
            avg_score_str = f"{avg_score:.3f}" if avg_score is not None else "N/A"
            lift_str = f"{lift:.2f}x" if lift is not None else "N/A"
            cum_bad_rate_pct = f"{(cum_bad_rate * 100):.2f}%" if cum_bad_rate is not None else "N/A"
            
            prompt_parts.append(
                f"{decile_num} | {count} | {bads} | {goods} | {bad_rate_pct} | {avg_score_str} | {lift_str} | {cum_bad_rate_pct}"
            )
        
        prompt_parts.extend([
            "",
            "Please provide a 2-4 line explanation (approximately 2-4 sentences) that:",
            "1. Describes the overall decile progression pattern based on the table data (e.g., increasing bad rates from low to high deciles)",
            "2. Mentions the monotonicity score and key observations from the table",
            "3. Explains what the deciles represent in simple terms",
            "4. Is written in clear, professional language",
            "5. Is concise and informative",
            "",
            "Return ONLY the explanation text, no additional formatting or preamble."
        ])
        
        prompt = "\n".join(prompt_parts)
        
        # Call LLM service
        writeup = await llm_service.generate_text(
            prompt=prompt,
            max_tokens=200,
            temperature=0.3
        )
        
        if not writeup or writeup.strip() == "":
            return DecileProgressionWriteupResponse(
                success=False,
                error="LLM returned empty response"
            )
        
        logger.info("Decile progression write-up generated successfully")
        return DecileProgressionWriteupResponse(
            success=True,
            writeup=writeup.strip()
        )
        
    except Exception as e:
        logger.error(f"Error generating decile progression write-up: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return DecileProgressionWriteupResponse(
            success=False,
            error=str(e)
        )


class AIExplainabilityWriteupRequest(BaseModel):
    beeswarm_data: Optional[List[Dict[str, Any]]] = None
    waterfall_data: Optional[List[Dict[str, Any]]] = None
    pdp_data: Optional[List[Dict[str, Any]]] = None


class AIExplainabilityWriteupResponse(BaseModel):
    success: bool
    writeup: Optional[str] = None
    error: Optional[str] = None


@documentation_router.post("/documentation/generate-ai-explainability-writeup", response_model=AIExplainabilityWriteupResponse)
async def generate_ai_explainability_writeup(request: AIExplainabilityWriteupRequest):
    """
    Generate a 3-4 line explanation of what AI explainability charts (Beeswarm Plot, Waterfall, PDP/ICE Lines) tell and explain.
    """
    try:
        logger.info("Generating AI explainability write-up")
        
        # Build summary of the charts
        chart_summaries = []
        
        # Beeswarm Plot summary
        if request.beeswarm_data and len(request.beeswarm_data) > 0:
            beeswarm_features = len(request.beeswarm_data)
            sample_feature = request.beeswarm_data[0]
            feature_name = sample_feature.get('featureName', 'feature')
            values_count = len(sample_feature.get('values', []))
            chart_summaries.append(f"Beeswarm Plot: Shows SHAP value distributions for {beeswarm_features} features, with {values_count} predictions per feature. Each dot represents a single prediction, colored by feature value.")
        
        # Waterfall summary
        if request.waterfall_data and len(request.waterfall_data) > 0:
            waterfall_features = len(request.waterfall_data)
            chart_summaries.append(f"Waterfall Plot: Explains a single prediction by showing how {waterfall_features} features contribute, starting from a base value.")
        
        # PDP/ICE Lines summary
        if request.pdp_data and len(request.pdp_data) > 0:
            pdp_features = len(request.pdp_data)
            sample_pdp = request.pdp_data[0]
            ice_lines_count = len(sample_pdp.get('ice_lines', []))
            chart_summaries.append(f"PDP/ICE Lines: Shows partial dependence for {pdp_features} features with {ice_lines_count} individual conditional expectation (ICE) lines, revealing how predictions change as each feature varies.")
        
        if not chart_summaries:
            return AIExplainabilityWriteupResponse(
                success=False,
                error="No explainability chart data provided"
            )
        
        # Build the prompt
        prompt_parts = [
            "You are a data scientist creating a concise explanation of AI explainability charts for model documentation.",
            "Generate a 3-4 line plain English explanation that describes what these AI explainability charts tell us about the model's behavior and feature importance.",
            "",
            "The following explainability charts are available:",
            ""
        ]
        
        prompt_parts.extend(chart_summaries)
        
        prompt_parts.extend([
            "",
            "Write a clear, concise 3-4 line explanation that helps readers understand:",
            "1. What these charts collectively reveal about the model",
            "2. How they help interpret feature importance and model predictions",
            "3. What insights they provide for understanding model behavior",
            "",
            "Keep the explanation simple and accessible, avoiding technical jargon where possible."
        ])
        
        prompt = "\n".join(prompt_parts)
        
        # Call LLM service
        response = await llm_service.generate_text(
            prompt=prompt,
            max_tokens=200,
            temperature=0.7
        )
        
        if not response or not response.strip():
            return AIExplainabilityWriteupResponse(
                success=False,
                error="Failed to generate AI explainability write-up"
            )
        
        writeup = response.strip()
        
        logger.info("AI explainability write-up generated successfully")
        return AIExplainabilityWriteupResponse(
            success=True,
            writeup=writeup
        )
        
    except Exception as e:
        logger.error(f"Error generating AI explainability write-up: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return AIExplainabilityWriteupResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/get-transformed-variables", response_model=TransformedVariablesResponse)
async def get_transformed_variables(request: TransformedVariablesRequest):
    """
    Get transformed variables data from chat history.
    Extracts transformation response data from feature engineering messages.
    """
    try:
        logger.info(f"Getting transformed variables for dataset: {request.dataset_id}")
        
        # Get chat history for this dataset from MessageState
        state = message_state_manager.create_or_load_state(request.dataset_id, "")
        chat_history = state.get("chat_history", [])
        if not chat_history:
            logger.warning(f"No chat history found for dataset: {request.dataset_id}")
            return TransformedVariablesResponse(
                success=True,
                transformed_variables=[]
            )
        
        transformed_variables = []
        
        # Search through chat history for feature transformation responses
        for message in reversed(chat_history):  # Start from most recent
            try:
                # Check if this is a feature transformation response
                if message.get("role") == "assistant":
                    content = message.get("content", "")
                    if not content:
                        continue
                    
                    # Try to parse as JSON
                    try:
                        parsed = json.loads(content) if isinstance(content, str) else content
                    except:
                        # If not JSON, check if it's a string containing transformation data
                        if "feature-transformation" in str(content).lower() or "transformation" in str(content).lower():
                            # Try to extract from string
                            continue
                        continue
                    
                    # Check for response_data with transformation structure
                    if isinstance(parsed, dict):
                        # Check for direct response_data
                        response_data = parsed.get("response_data")
                        if response_data and isinstance(response_data, list):
                            # Check if it looks like transformed variables data
                            if len(response_data) > 0 and isinstance(response_data[0], dict):
                                if "new_variable_name" in response_data[0] or "transformation_methods" in response_data[0]:
                                    transformed_variables = response_data
                                    logger.info(f"Found {len(transformed_variables)} transformed variables in chat history")
                                    break
                        
                        # Check nested structure (response -> response_data)
                        response = parsed.get("response")
                        if response:
                            if isinstance(response, str):
                                try:
                                    inner_response = json.loads(response)
                                    if isinstance(inner_response, dict):
                                        response_data = inner_response.get("response_data")
                                        if response_data and isinstance(response_data, list):
                                            if len(response_data) > 0 and isinstance(response_data[0], dict):
                                                if "new_variable_name" in response_data[0] or "transformation_methods" in response_data[0]:
                                                    transformed_variables = response_data
                                                    logger.info(f"Found {len(transformed_variables)} transformed variables in nested chat history")
                                                    break
                                except:
                                    pass
                            elif isinstance(response, dict):
                                response_data = response.get("response_data")
                                if response_data and isinstance(response_data, list):
                                    if len(response_data) > 0 and isinstance(response_data[0], dict):
                                        if "new_variable_name" in response_data[0] or "transformation_methods" in response_data[0]:
                                            transformed_variables = response_data
                                            logger.info(f"Found {len(transformed_variables)} transformed variables in nested chat history")
                                            break
            except Exception as e:
                logger.debug(f"Error parsing message for transformed variables: {str(e)}")
                continue
        
        # Filter to only include required columns
        filtered_variables = []
        for var in transformed_variables:
            filtered_variables.append({
                "new_variable_name": var.get("new_variable_name", ""),
                "var_type": var.get("var_type", ""),
                "variable_definition": var.get("variable_definition", ""),
                "transformation_methods": var.get("transformation_methods", "")
            })
        
        logger.info(f"Returning {len(filtered_variables)} transformed variables")
        return TransformedVariablesResponse(
            success=True,
            transformed_variables=filtered_variables
        )
        
    except Exception as e:
        logger.error(f"Error getting transformed variables: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return TransformedVariablesResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/get-variable-analysis", response_model=VariableAnalysisResponse)
async def get_variable_analysis(request: VariableAnalysisRequest):
    """
    Get variable analysis data (variable_statistics with Correlation, VIF, IV) from MessageState.
    """
    try:
        logger.info(f"Getting variable analysis for dataset: {request.dataset_id}")
        
        # Get MessageState for the dataset
        state = message_state_manager.create_or_load_state(request.dataset_id, "")
        
        # Try to get variable_analysis from multiple locations
        variable_analysis = (
            state.get("variable_analysis") or
            state.get("variableAnalysis") or
            state.get("training_context", {}).get("variable_analysis") if isinstance(state.get("training_context"), dict) else None
        )
        
        # Get variable_statistics (flattened version if available)
        variable_statistics = (
            state.get("variable_statistics") or
            (variable_analysis.get("variable_statistics") if isinstance(variable_analysis, dict) else None) or
            []
        )
        
        if variable_statistics and len(variable_statistics) > 0:
            logger.info(f"Found {len(variable_statistics)} variable statistics in MessageState")
            return VariableAnalysisResponse(
                success=True,
                variable_analysis=variable_analysis if isinstance(variable_analysis, dict) else None,
                variable_statistics=variable_statistics if isinstance(variable_statistics, list) else []
            )
        else:
            logger.warning(f"No variable statistics found in MessageState for dataset: {request.dataset_id}")
            return VariableAnalysisResponse(
                success=True,
                variable_analysis=None,
                variable_statistics=[]
            )
        
    except Exception as e:
        logger.error(f"Error getting variable analysis: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return VariableAnalysisResponse(
            success=False,
            error=str(e)
        )


@documentation_router.post("/documentation/get-data-insights", response_model=DataInsightsResponse)
async def get_data_insights(request: DataInsightsRequest):
    """
    Get data insights by calling the bivariate analysis API directly.
    Extracts bivariate analysis data and generates EDA report and insights.
    Also tries to get aggregated LLM insights from chat history.
    """
    try:
        logger.info(f"Getting data insights for dataset: {request.dataset_id}")
        
        # Get target variable from dataset info
        dataset_info = dataset_manager.get_dataset_info(request.dataset_id)
        if not dataset_info:
            logger.warning(f"Dataset info not found for dataset: {request.dataset_id}")
            return DataInsightsResponse(
                success=False,
                error=f"Dataset not found: {request.dataset_id}"
            )
        
        target_variable = dataset_info.get('target_variable')
        if not target_variable:
            logger.warning(f"Target variable not found for dataset: {request.dataset_id}")
            return DataInsightsResponse(
                success=False,
                error=f"Target variable not configured for dataset: {request.dataset_id}"
            )
        
        logger.info(f"Target variable: {target_variable}")
        
        # Try to get used_features from latest model evaluation for this dataset
        used_features = []
        try:
            from app.models.model_evaluation_database import model_evaluation_db
            # Get all models for this dataset
            all_models = model_evaluation_db.list_all_models()
            # Find the latest model for this dataset
            for model in reversed(all_models):  # Reverse to get latest first
                try:
                    # Extract model_id from model dict
                    model_id = model.get('id') if isinstance(model, dict) else model
                    if not model_id:
                        continue
                    eval_data = model_evaluation_db.get_model_evaluation(model_id, include_explainability=False)
                    if eval_data:
                        model_info = eval_data.get('model', {})
                        if model_info.get('dataset_id') == request.dataset_id:
                            used_features = eval_data.get('used_features', [])
                            if not used_features:
                                # Fallback to feature_importance
                                feature_importance = eval_data.get('feature_importance', [])
                                if feature_importance:
                                    used_features = [f.get('feature_name') for f in feature_importance if f.get('feature_name')]
                            if used_features:
                                logger.info(f"Found {len(used_features)} used_features from model {model_id}")
                                break
                except Exception as model_error:
                    # Skip this model if there's an error and continue to next
                    logger.debug(f"Skipping model {model.get('id', 'unknown') if isinstance(model, dict) else 'unknown'} due to error: {str(model_error)}")
                    continue
        except Exception as e:
            logger.warning(f"Could not get used_features: {str(e)}")
        
        # Check if user has actually requested Data Insights analysis (specifically bivariate or IV)
        user_requested_bivariate = False
        user_requested_iv = False
        user_requested_correlation = False
        user_requested_correlation_analysis = False
        user_requested_vif = False
        chat_history = []
        try:
            state = message_state_manager.create_or_load_state(request.dataset_id, "")
            chat_history = state.get("chat_history", [])
            userquery = state.get("userquery", "")
            messages = state.get("messages", [])
            intent = state.get("intent", "")
            logger.info(f"Loaded chat history with {len(chat_history)} messages")
            logger.info(f"Loaded {len(messages)} messages from state")
            logger.info(f"Current userquery in state: {userquery[:200] if userquery else 'None'}...")
            logger.info(f"Current intent in state: {intent}")
            
            # Define keywords for detection
            # IMPORTANT: Order matters - check specific patterns first, then general ones
            bivariate_keywords = [
                "bivariate_analysis",  # Most specific first
                "bivariate analysis",
                "bivariate"
            ]
            iv_keywords = [
                "iv_analysis",  # Most specific first
                "iv analysis",
                "information value"
            ]
            correlation_keywords = [
                "correlation_matrix",  # Most specific first
                "correlation matrix",
                "generate correlation",  # More specific - only match when generating
                "correlation analysis"  # More specific
            ]
            correlation_analysis_keywords = [  # NEW
                "correlation_analysis",  # Most specific first
                "correlation analysis"
            ]
            vif_keywords = [  # NEW
                "variance_inflation_factor",  # Most specific first
                "variance inflation factor",
                "vif"
            ]
            # Note: "iv" as standalone word only (not part of "insights")
            # Note: "generate insights" only if not followed by specific analysis type
            
            # ZERO: Check intent field - if intent is 'data_insight', user likely requested data insights
            # But we still need to check chat history to determine which specific type (bivariate/IV)
            if intent == "data_insight":
                logger.info("Intent is 'data_insight' - user has requested data insights, will check chat history for specific type")
            
            # Helper function to check if a query string contains a keyword as a whole word or exact phrase
            def check_keywords_in_query(query_str: str, keywords: list, require_word_boundary: bool = False) -> bool:
                if not query_str:
                    return False
                query_lower = query_str.lower()
                for keyword in keywords:
                    keyword_lower = keyword.lower()
                    # For short keywords like "iv", require word boundaries to avoid matching "insights"
                    if require_word_boundary and len(keyword_lower) <= 3:
                        # Use regex word boundaries: \biv\b matches "iv" but not "insights"
                        import re
                        pattern = r'\b' + re.escape(keyword_lower) + r'\b'
                        if re.search(pattern, query_lower):
                            return True
                    else:
                        # For longer keywords, simple substring match is fine
                        if keyword_lower in query_lower:
                            return True
                return False
            
            # FIRST: Check the current userquery field in state (most recent query)
            # This is the most reliable indicator since it's the latest user query
            # Parse the query intelligently - if it says "Generate insights for: X", only match X
            if userquery:
                userquery_lower = userquery.lower()
                logger.info(f"Checking current userquery for keywords: {userquery_lower[:150]}...")
                
                # Check if query has format "Generate insights for: X" or "Generate insights for X"
                # If so, only check what comes after "for:" or "for"
                # Updated regex to capture everything after "for:" including commas
                match = re.search(r'generate\s+insights\s+for\s*:?\s*(.+)', userquery_lower)
                if match:
                    # Extract what comes after "for:" - this may contain comma-separated values
                    requested_types_str = match.group(1).strip()
                    logger.info(f"Detected specific request format, extracted types: '{requested_types_str}'")
                    
                    # Split by comma to handle multiple analysis types
                    requested_types = [t.strip() for t in requested_types_str.split(',')]
                    logger.info(f"Parsed {len(requested_types)} analysis type(s): {requested_types}")
                    
                    # Check each type independently (not using elif, so multiple can be True)
                    for requested_type in requested_types:
                        requested_type_lower = requested_type.lower()
                        
                        if "bivariate" in requested_type_lower:
                            user_requested_bivariate = True
                            logger.info(f"✓ User requested Bivariate Analysis (from specific request): {requested_type}")
                        
                        if "iv" in requested_type_lower or "information value" in requested_type_lower:
                            user_requested_iv = True
                            logger.info(f"✓ User requested IV Analysis (from specific request): {requested_type}")
                        
                        if "correlation" in requested_type_lower:
                            # Check if it's correlation_analysis or correlation_matrix
                            if "correlation_analysis" in requested_type_lower or "correlation analysis" in requested_type_lower:
                                user_requested_correlation_analysis = True
                                logger.info(f"✓ User requested Correlation Analysis (from specific request): {requested_type}")
                            elif "correlation_matrix" in requested_type_lower or "correlation matrix" in requested_type_lower:
                                user_requested_correlation = True
                                logger.info(f"✓ User requested Correlation Matrix (from specific request): {requested_type}")
                            else:
                                # Default to correlation_matrix if just "correlation" is mentioned
                                user_requested_correlation = True
                                logger.info(f"✓ User requested Correlation Matrix (from specific request, default): {requested_type}")
                        
                        if "vif" in requested_type_lower or "variance_inflation_factor" in requested_type_lower or "variance inflation factor" in requested_type_lower:
                            user_requested_vif = True
                            logger.info(f"✓ User requested VIF Analysis (from specific request): {requested_type}")
                    # Don't check general keywords if we found a specific request format
                else:
                    # No specific format - check all keywords
                    if check_keywords_in_query(userquery, bivariate_keywords):
                        user_requested_bivariate = True
                        logger.info(f"✓ User requested Bivariate Analysis (matched in userquery): {userquery[:100]}...")
                    # For IV, require word boundaries to avoid matching "insights"
                    if check_keywords_in_query(userquery, iv_keywords, require_word_boundary=True):
                        user_requested_iv = True
                        logger.info(f"✓ User requested IV Analysis (matched in userquery): {userquery[:100]}...")
                    if check_keywords_in_query(userquery, correlation_keywords):
                        user_requested_correlation = True
                        logger.info(f"✓ User requested Correlation Matrix (matched in userquery): {userquery[:100]}...")
                    if check_keywords_in_query(userquery, correlation_analysis_keywords):  # NEW
                        user_requested_correlation_analysis = True
                        logger.info(f"✓ User requested Correlation Analysis (matched in userquery): {userquery[:100]}...")
                    if check_keywords_in_query(userquery, vif_keywords):  # NEW
                        user_requested_vif = True
                        logger.info(f"✓ User requested VIF Analysis (matched in userquery): {userquery[:100]}...")
                    # Only check general keywords if no specific analysis was requested
                    if not user_requested_bivariate and not user_requested_iv and not user_requested_correlation and not user_requested_correlation_analysis:
                        general_keywords = ["auto data insights", "generate auto insights"]
                        if check_keywords_in_query(userquery, general_keywords):
                            user_requested_bivariate = True
                            user_requested_iv = True
                            logger.info(f"✓ User requested general Data Insights (matched in userquery): {userquery[:100]}...")
            
            # SECOND: Check messages array for HumanMessage objects (LangChain format)
            if messages and not user_requested_bivariate and not user_requested_iv and not user_requested_correlation:
                logger.info(f"Checking {len(messages)} messages from state for user queries...")
                for msg in reversed(messages):  # Check latest first
                    # Check if it's a HumanMessage (user message)
                    if hasattr(msg, 'content') and hasattr(msg, '__class__') and 'Human' in str(msg.__class__):
                        msg_content = str(msg.content) if msg.content else ""
                        if msg_content:
                            logger.info(f"Found HumanMessage: {msg_content[:150]}...")
                            # Use same intelligent parsing as userquery
                            match = re.search(r'generate\s+insights\s+for\s*:?\s*([^,]+)', msg_content.lower())
                            if match:
                                requested_type = match.group(1).strip()
                                if "bivariate" in requested_type:
                                    user_requested_bivariate = True
                                    logger.info(f"✓ User requested Bivariate Analysis (from HumanMessage): {requested_type}")
                                elif "iv" in requested_type or "information value" in requested_type:
                                    user_requested_iv = True
                                    logger.info(f"✓ User requested IV Analysis (from HumanMessage): {requested_type}")
                                elif "correlation_matrix" in requested_type or "correlation matrix" in requested_type:
                                    user_requested_correlation = True
                                    logger.info(f"✓ User requested Correlation Matrix (from HumanMessage): {requested_type}")
                            else:
                                if check_keywords_in_query(msg_content, bivariate_keywords):
                                    user_requested_bivariate = True
                                    logger.info(f"✓ User requested Bivariate Analysis (matched in HumanMessage): {msg_content[:100]}...")
                                if check_keywords_in_query(msg_content, iv_keywords, require_word_boundary=True):
                                    user_requested_iv = True
                                    logger.info(f"✓ User requested IV Analysis (matched in HumanMessage): {msg_content[:100]}...")
                                if check_keywords_in_query(msg_content, correlation_keywords):
                                    user_requested_correlation = True
                                    logger.info(f"✓ User requested Correlation Matrix (matched in HumanMessage): {msg_content[:100]}...")
                            if user_requested_bivariate or user_requested_iv or user_requested_correlation:
                                break
            
            # Check if user has requested specific analyses in their messages
            # Also check assistant messages for parsed selections (from agentic system logs)
            
            # Check user messages for explicit requests
            # Note: chat_history format is: {"role": "user"/"assistant", "content": [{"type": "text", "text": "..."}]}
            # OR: {"role": "user"/"assistant", "content": "..."} (string format)
            # OR: {"query": "...", "response": "...", "intent": "..."} (old format)
            for idx, message in enumerate(chat_history):
                # Handle old format (query/response/intent)
                if "query" in message and "role" not in message:
                    user_query = message.get("query", "")
                    if user_query:
                        user_query_lower = user_query.lower()
                        logger.info(f"Checking user message {idx} (old format) for keywords: {user_query[:150]}...")
                        # Use same intelligent parsing
                        match = re.search(r'generate\s+insights\s+for\s*:?\s*([^,]+)', user_query_lower)
                        if match:
                            requested_type = match.group(1).strip()
                            if "bivariate" in requested_type:
                                user_requested_bivariate = True
                                logger.info(f"✓ User requested Bivariate Analysis (from old format): {requested_type}")
                            elif "iv" in requested_type or "information value" in requested_type:
                                user_requested_iv = True
                                logger.info(f"✓ User requested IV Analysis (from old format): {requested_type}")
                            elif "correlation_matrix" in requested_type or "correlation matrix" in requested_type:
                                user_requested_correlation = True
                                logger.info(f"✓ User requested Correlation Matrix (from old format): {requested_type}")
                        else:
                            for keyword in bivariate_keywords:
                                if keyword.lower() in user_query_lower:
                                    user_requested_bivariate = True
                                    logger.info(f"✓ User requested Bivariate Analysis (matched keyword '{keyword}'): {user_query[:100]}...")
                                    break
                            # For IV, use word boundaries
                            if check_keywords_in_query(user_query, iv_keywords, require_word_boundary=True):
                                user_requested_iv = True
                                logger.info(f"✓ User requested IV Analysis (matched in old format): {user_query[:100]}...")
                    continue
                
                # Handle new format (role/content)
                if message.get("role") == "user":
                    user_query = message.get("content", "")
                    # Handle different content formats
                    if isinstance(user_query, list) and len(user_query) > 0:
                        # Format: [{"type": "text", "text": "..."}]
                        user_query = user_query[0].get("text", "") if isinstance(user_query[0], dict) else str(user_query[0])
                    elif not isinstance(user_query, str):
                        user_query = str(user_query)
                    
                    if user_query:
                        user_query_lower = user_query.lower()
                        logger.info(f"Checking user message {idx} for keywords: {user_query[:150]}...")
                        # Use same intelligent parsing
                        match = re.search(r'generate\s+insights\s+for\s*:?\s*([^,]+)', user_query_lower)
                        if match:
                            requested_type = match.group(1).strip()
                            if "bivariate" in requested_type:
                                user_requested_bivariate = True
                                logger.info(f"✓ User requested Bivariate Analysis (from chat_history): {requested_type}")
                            elif "iv" in requested_type or "information value" in requested_type:
                                user_requested_iv = True
                                logger.info(f"✓ User requested IV Analysis (from chat_history): {requested_type}")
                            elif "correlation_matrix" in requested_type or "correlation matrix" in requested_type:
                                user_requested_correlation = True
                                logger.info(f"✓ User requested Correlation Matrix (from chat_history): {requested_type}")
                        else:
                            # Check for specific analyses
                            for keyword in bivariate_keywords:
                                if keyword.lower() in user_query_lower:
                                    user_requested_bivariate = True
                                    logger.info(f"✓ User requested Bivariate Analysis (matched keyword '{keyword}'): {user_query[:100]}...")
                                    break
                            # For IV, use word boundaries
                            if check_keywords_in_query(user_query, iv_keywords, require_word_boundary=True):
                                user_requested_iv = True
                                logger.info(f"✓ User requested IV Analysis (matched in chat_history): {user_query[:100]}...")
                            # For correlation, only match specific request patterns to avoid false positives from system prompts
                            # Check for explicit request patterns first
                            correlation_request_patterns = [
                                r'generate\s+(?:insights?\s+for\s*:?\s*)?correlation',
                                r'correlation\s+matrix',
                                r'correlation_matrix',
                                r'correlation\s+analysis'
                            ]
                            matched_correlation = False
                            for pattern in correlation_request_patterns:
                                if re.search(pattern, user_query_lower):
                                    user_requested_correlation = True
                                    matched_correlation = True
                                    logger.info(f"✓ User requested Correlation Matrix (matched pattern '{pattern}' in chat_history): {user_query[:100]}...")
                                    break
                            # Only check general keywords if no specific analysis was requested
                            if not user_requested_bivariate and not user_requested_iv and not user_requested_correlation:
                                general_keywords = ["auto data insights", "generate auto insights"]
                                if check_keywords_in_query(user_query, general_keywords):
                                    user_requested_bivariate = True
                                    user_requested_iv = True
                                    user_requested_correlation = True
                                    logger.info(f"✓ User requested general Data Insights (matched in chat_history): {user_query[:100]}...")
            
            # Also check assistant messages for bivariate_analysis data (indicates user requested it)
            # This is a fallback in case keyword matching fails - if we find the data, user must have requested it
            # Check assistant messages regardless of keyword match status, as a more reliable indicator
            # We check even if one is already True, to ensure we catch both if they exist
            assistant_count = len([m for m in chat_history if m.get('role') == 'assistant' or ('response' in m and 'role' not in m)])
            logger.info(f"Checking {assistant_count} assistant messages for bivariate_analysis data")
            for idx, message in enumerate(reversed(chat_history)):  # Check latest first
                text_content = None
                
                # Handle old format (query/response/intent)
                if "response" in message and "role" not in message:
                    text_content = message.get("response", "")
                    if not isinstance(text_content, str):
                        text_content = str(text_content)
                # Handle new format (role/content)
                elif message.get("role") == "assistant":
                    content = message.get("content", "")
                    if isinstance(content, list) and len(content) > 0:
                        # Format: [{"type": "text", "text": "..."}]
                        text_content = content[0].get("text", "") if isinstance(content[0], dict) else str(content[0])
                    elif isinstance(content, str):
                        text_content = content
                    else:
                        text_content = str(content)
                else:
                    continue  # Skip non-assistant messages
                
                # Log first 200 chars of assistant message for debugging
                if idx == 0 and text_content:  # Latest message
                    logger.info(f"Latest assistant message preview: {text_content[:200]}...")
                
                # If we find bivariate_analysis data, user must have requested it
                if text_content and "bivariate_analysis" in text_content and isinstance(text_content, str) and not user_requested_bivariate:
                    logger.info(f"Found 'bivariate_analysis' string in assistant message {idx}, attempting to parse...")
                    try:
                        parsed = json.loads(text_content)
                        response_str = parsed.get("response", "")
                        if isinstance(response_str, str):
                            inner_response = json.loads(response_str)
                            inner_response_data = inner_response.get("response", {})
                            if isinstance(inner_response_data, dict):
                                bivariate_array = inner_response_data.get("bivariate_analysis", [])
                                if isinstance(bivariate_array, list) and len(bivariate_array) > 0:
                                    user_requested_bivariate = True
                                    logger.info(f"✓ Found bivariate_analysis data in chat - user must have requested it ({len(bivariate_array)} entries)")
                                    if user_requested_iv:  # If we found both, we can break
                                        break
                    except (json.JSONDecodeError, KeyError, TypeError) as e:
                        logger.warning(f"Error parsing bivariate_analysis from assistant message {idx}: {str(e)}")
                        logger.debug(f"Message content type: {type(text_content)}, length: {len(text_content) if isinstance(text_content, str) else 'N/A'}")
                        # Try alternative parsing - maybe it's already a dict
                        try:
                            if isinstance(text_content, dict):
                                response_data = text_content.get("response", {})
                                if isinstance(response_data, str):
                                    inner_response = json.loads(response_data)
                                    inner_response_data = inner_response.get("response", {})
                                    if isinstance(inner_response_data, dict):
                                        bivariate_array = inner_response_data.get("bivariate_analysis", [])
                                        if isinstance(bivariate_array, list) and len(bivariate_array) > 0:
                                            user_requested_bivariate = True
                                            logger.info(f"✓ Found bivariate_analysis data (alternative parse) - user must have requested it ({len(bivariate_array)} entries)")
                        except Exception as e2:
                            logger.debug(f"Alternative parse also failed: {str(e2)}")
                        pass
                
                # If we find iv_analysis_summary or iv_insight, user must have requested IV
                if text_content and ("iv_analysis_summary" in text_content or "iv_insight" in text_content) and isinstance(text_content, str) and not user_requested_iv:
                    try:
                        parsed = json.loads(text_content)
                        response_str = parsed.get("response", "")
                        if isinstance(response_str, str):
                            inner_response = json.loads(response_str)
                            inner_response_data = inner_response.get("response", {})
                            inner_data = inner_response.get("data", {})
                            
                            # Check for IV analysis summary
                            if isinstance(inner_response_data, dict):
                                iv_summary = inner_response_data.get("iv_analysis_summary", [])
                                if (isinstance(iv_summary, list) and len(iv_summary) > 0) or (isinstance(iv_summary, dict)):
                                    user_requested_iv = True
                                    logger.info(f"Found iv_analysis_summary in chat - user must have requested it")
                            
                            # Check for IV insights
                            if isinstance(inner_data, dict):
                                iv_insight_list = inner_data.get("iv_insight", [])
                                if isinstance(iv_insight_list, list) and len(iv_insight_list) > 0:
                                    user_requested_iv = True
                                    logger.info(f"Found iv_insight in chat - user must have requested it ({len(iv_insight_list)} insights)")
                            
                            if user_requested_iv and user_requested_bivariate and user_requested_correlation:  # If we found all, we can break
                                break
                    except (json.JSONDecodeError, KeyError, TypeError) as e:
                        logger.debug(f"Error parsing IV data from assistant message: {str(e)}")
                        pass
                
                # If we find correlation_matrix_heatmap or correlation_matrix_insight, user must have requested correlation
                if text_content and ("correlation_matrix_heatmap" in text_content or "correlation_matrix_insight" in text_content) and isinstance(text_content, str) and not user_requested_correlation:
                    logger.info(f"Found 'correlation_matrix' string in assistant message {idx}, attempting to parse...")
                    try:
                        parsed = json.loads(text_content)
                        response_str = parsed.get("response", "")
                        if isinstance(response_str, str):
                            inner_response = json.loads(response_str)
                            inner_response_data = inner_response.get("response", {})
                            inner_data = inner_response.get("data", {})
                            
                            # Check for correlation matrix heatmap
                            if isinstance(inner_response_data, dict):
                                corr_heatmap = inner_response_data.get("correlation_matrix_heatmap", [])
                                if isinstance(corr_heatmap, list) and len(corr_heatmap) > 0:
                                    user_requested_correlation = True
                                    # Extract the first table from the list
                                    corr_table = corr_heatmap[0]
                                    if isinstance(corr_table, dict):
                                        correlation_matrix_heatmap = corr_table
                                        logger.info(f"✓ Found correlation_matrix_heatmap in chat - user must have requested it ({len(corr_table.get('rows', []))} rows)")
                                    else:
                                        logger.warning(f"correlation_matrix_heatmap[0] is not a dict: {type(corr_table)}")
                                elif isinstance(corr_heatmap, dict):
                                    # Sometimes it might be a dict directly, not a list
                                    user_requested_correlation = True
                                    correlation_matrix_heatmap = corr_heatmap
                                    logger.info(f"✓ Found correlation_matrix_heatmap (direct dict) in chat - user must have requested it ({len(corr_heatmap.get('rows', []))} rows)")

                            
                            # Check for correlation matrix insights
                            if isinstance(inner_data, dict):
                                corr_insight_list = inner_data.get("correlation_matrix_insight", [])
                                if isinstance(corr_insight_list, list) and len(corr_insight_list) > 0:
                                    user_requested_correlation = True
                                    logger.info(f"✓ Found correlation_matrix_insight in chat - user must have requested it ({len(corr_insight_list)} insights)")
                            
                            if user_requested_bivariate and user_requested_iv and user_requested_correlation:  # If we found all, we can break
                                break
                    except (json.JSONDecodeError, KeyError, TypeError) as e:
                        logger.warning(f"Error parsing correlation_matrix from assistant message {idx}: {str(e)}")
                        pass
            
            logger.info(f"Final detection results: user_requested_bivariate={user_requested_bivariate}, user_requested_iv={user_requested_iv}, user_requested_correlation={user_requested_correlation}, user_requested_correlation_analysis={user_requested_correlation_analysis}, user_requested_vif={user_requested_vif}")
            if not user_requested_bivariate and not user_requested_iv and not user_requested_correlation and not user_requested_correlation_analysis and not user_requested_vif:
                logger.info("User has not requested any Data Insights - will not show sections")
        except Exception as e:
            logger.warning(f"Could not check if user requested data insights: {str(e)}")
        
        # If user hasn't requested any data insights, return empty response
        if not user_requested_bivariate and not user_requested_iv and not user_requested_correlation and not user_requested_correlation_analysis and not user_requested_vif:
            return DataInsightsResponse(
                success=True,
                insights={}
            )
        
        # Try to get aggregated LLM insights from chat history (for bivariate, IV, and correlation)
        aggregated_insights = []
        iv_aggregated_insights = []
        correlation_aggregated_insights = []
        correlation_analysis_insights = []  # NEW
        vif_aggregated_insights = []  # NEW
        correlation_numeric_table = None 
        iv_analysis_summary = None
        correlation_matrix_heatmap = None
        vif_analysis_table = None  # NEW
        
        try:
            logger.info(f"Searching through {len(chat_history)} messages in chat history for insights")
            
            # Search for insights in chat history (reverse order for latest first)
            # Search ALL assistant messages to find IV data even if it's in a different message
            for idx, message in enumerate(reversed(chat_history)):
                if message.get("role") == "assistant":
                    content = message.get("content", "")
                    if isinstance(content, list) and len(content) > 0:
                        text_content = content[0].get("text", "") if isinstance(content[0], dict) else str(content[0])
                    elif isinstance(content, str):
                        text_content = content
                    else:
                        text_content = str(content)
                    
                    # Check if this message contains data insights (bivariate or IV)
                    # Be more lenient - check for any data insight indicators
                    has_insights_indicator = (
                        "bivariate_insight" in text_content or 
                        "bivariate_analysis" in text_content or
                        "iv_insight" in text_content or
                        "iv_analysis" in text_content or
                        "iv_analysis_summary" in text_content or
                        ("data" in text_content.lower() and "insight" in text_content.lower()) or
                        ("standard_insights" in text_content and ("bivariate" in text_content.lower() or "iv" in text_content.lower()))
                    )
                    
                    # Process ALL assistant messages that might contain JSON data
                    # (IV or correlation data might be in a separate message from bivariate)
                    should_process = has_insights_indicator or (
                        "iv_insight" in text_content or
                        "iv_analysis" in text_content or
                        "iv_analysis_summary" in text_content or
                        "correlation_matrix_heatmap" in text_content or
                        "correlation_matrix_insight" in text_content or
                        "correlation_numeric" in text_content or  # NEW
                        "correlation_insight" in text_content or  # NEW
                        "correlation_analysis" in text_content or  # NEW
                        "vif_analysis" in text_content or  # NEW
                        "vif_insight" in text_content or  # NEW
                        ("response" in text_content and "{" in text_content)  # Likely JSON structure
                    )
                    
                    if should_process:
                        logger.debug(f"Processing message at index {idx} for insights extraction")
                        try:
                            # Try to parse as JSON
                            parsed = json.loads(text_content)
                            
                            # Structure from agentic_system.py:
                            # payload = {"response": json.dumps(inner_response)}
                            # inner_response = {"response": response_data, "data": {"bivariate_insight": [...], "iv_insight": [...], ...}}
                            
                            # First, check if parsed["response"] is a JSON string
                            response_str = parsed.get("response", "")
                            if isinstance(response_str, str):
                                try:
                                    inner_response = json.loads(response_str)
                                    # inner_response has "response" (tables) and "data" (insights)
                                    inner_data = inner_response.get("data", {})
                                    inner_response_data = inner_response.get("response", {})
                                    
                                    if isinstance(inner_data, dict):
                                        # Extract bivariate insights
                                        if not aggregated_insights:
                                            bivariate_insight_list = inner_data.get("bivariate_insight", [])
                                            if isinstance(bivariate_insight_list, list) and len(bivariate_insight_list) > 0:
                                                aggregated_insights = bivariate_insight_list
                                                logger.info(f"Found {len(aggregated_insights)} bivariate LLM insights from chat history")
                                        
                                        # Extract IV insights
                                        if not iv_aggregated_insights:
                                            iv_insight_list = inner_data.get("iv_insight", [])
                                            if isinstance(iv_insight_list, list) and len(iv_insight_list) > 0:
                                                iv_aggregated_insights = iv_insight_list
                                                logger.info(f"Found {len(iv_aggregated_insights)} IV LLM insights from chat history")
                                        
                                        # Extract correlation matrix insights
                                        if not correlation_aggregated_insights:
                                            corr_insight_list = inner_data.get("correlation_matrix_insight", [])
                                            if isinstance(corr_insight_list, list) and len(corr_insight_list) > 0:
                                                correlation_aggregated_insights = corr_insight_list
                                                logger.info(f"Found {len(correlation_aggregated_insights)} Correlation Matrix LLM insights from chat history")

                                        # Extract correlation analysis insights (NEW)
                                        if not correlation_analysis_insights:
                                            corr_analysis_insight_list = inner_data.get("correlation_insight", [])
                                            if isinstance(corr_analysis_insight_list, list) and len(corr_analysis_insight_list) > 0:
                                                correlation_analysis_insights = corr_analysis_insight_list
                                                logger.info(f"Found {len(correlation_analysis_insights)} Correlation Analysis LLM insights from chat history")
                                        
                                        # Extract VIF insights (NEW)
                                        if not vif_aggregated_insights:
                                            vif_insight_list = inner_data.get("vif_insight", [])
                                            if isinstance(vif_insight_list, list) and len(vif_insight_list) > 0:
                                                vif_aggregated_insights = vif_insight_list
                                                logger.info(f"Found {len(vif_aggregated_insights)} VIF LLM insights from chat history")

                                    # Extract IV analysis summary (from response_data)
                                    # Structure: response_data["iv_analysis_summary"] = [{"columns": [...], "rows": [...], "title": ...}]
                                    if isinstance(inner_response_data, dict) and not iv_analysis_summary:
                                        iv_summary_tables = inner_response_data.get("iv_analysis_summary", [])
                                        logger.debug(f"Found iv_analysis_summary key: {iv_summary_tables is not None}, type: {type(iv_summary_tables)}")
                                        if isinstance(iv_summary_tables, list) and len(iv_summary_tables) > 0:
                                            # Get the first table (should be the summary)
                                            iv_table = iv_summary_tables[0]
                                            if isinstance(iv_table, dict):
                                                iv_analysis_summary = iv_table
                                                rows = iv_table.get('rows', [])
                                                logger.info(f"Found IV analysis summary with {len(rows)} variables")
                                                if len(rows) > 0:
                                                    logger.debug(f"Sample IV row: {rows[0]}")
                                        elif isinstance(iv_summary_tables, dict):
                                            # Sometimes it might be a dict directly, not a list
                                            iv_analysis_summary = iv_summary_tables
                                            rows = iv_summary_tables.get('rows', [])
                                            logger.info(f"Found IV analysis summary (direct dict) with {len(rows)} variables")
                                        elif iv_summary_tables is not None:
                                            logger.debug(f"iv_analysis_summary is not a list or dict: {type(iv_summary_tables)}, value: {str(iv_summary_tables)[:200]}")
                                    # Extract correlation numeric table (from response_data) - NEW
                                    # Structure: response_data["correlation_numeric"] = [{"columns": [...], "rows": [...], "title": ...}]
                                    if isinstance(inner_response_data, dict) and not correlation_numeric_table:
                                        corr_numeric_tables = inner_response_data.get("correlation_numeric", [])
                                        logger.debug(f"Found correlation_numeric key: {corr_numeric_tables is not None}, type: {type(corr_numeric_tables)}")
                                        if isinstance(corr_numeric_tables, list) and len(corr_numeric_tables) > 0:
                                            # Get the first table (should be the correlation numeric table)
                                            corr_numeric_table_obj = corr_numeric_tables[0]
                                            if isinstance(corr_numeric_table_obj, dict):
                                                correlation_numeric_table = corr_numeric_table_obj
                                                rows = corr_numeric_table_obj.get('rows', [])
                                                columns = corr_numeric_table_obj.get('columns', [])
                                                logger.info(f"Found Correlation Numeric table with {len(rows)} rows, {len(columns)} columns")
                                                if len(rows) > 0:
                                                    logger.debug(f"Sample correlation numeric row: {rows[0]}")
                                        elif isinstance(corr_numeric_tables, dict):
                                            # Sometimes it might be a dict directly, not a list
                                            correlation_numeric_table = corr_numeric_tables
                                            rows = corr_numeric_tables.get('rows', [])
                                            columns = corr_numeric_tables.get('columns', [])
                                            logger.info(f"Found Correlation Numeric table (direct dict) with {len(rows)} rows, {len(columns)} columns")
                                        elif corr_numeric_tables is not None:
                                            logger.debug(f"correlation_numeric is not a list or dict: {type(corr_numeric_tables)}, value: {str(corr_numeric_tables)[:200]}")
                                    # Extract correlation matrix heatmap (from response_data)
                                    # Structure: response_data["correlation_matrix_heatmap"] = [{"columns": [...], "rows": [...], "title": "Correlation Matrix"}]
                                    if isinstance(inner_response_data, dict) and not correlation_matrix_heatmap:
                                        corr_heatmap_tables = inner_response_data.get("correlation_matrix_heatmap", [])
                                        logger.debug(f"Found correlation_matrix_heatmap key: {corr_heatmap_tables is not None}, type: {type(corr_heatmap_tables)}")
                                        if isinstance(corr_heatmap_tables, list) and len(corr_heatmap_tables) > 0:
                                            # Get the first table (should be the correlation matrix)
                                            corr_table = corr_heatmap_tables[0]
                                            if isinstance(corr_table, dict):
                                                correlation_matrix_heatmap = corr_table
                                                rows = corr_table.get('rows', [])
                                                columns = corr_table.get('columns', [])
                                                logger.info(f"Found Correlation Matrix heatmap with {len(rows)} rows, {len(columns)} columns")
                                                if len(rows) > 0:
                                                    logger.debug(f"Sample correlation row: {rows[0]}")
                                        elif isinstance(corr_heatmap_tables, dict):
                                            # Sometimes it might be a dict directly, not a list
                                            correlation_matrix_heatmap = corr_heatmap_tables
                                            rows = corr_heatmap_tables.get('rows', [])
                                            columns = corr_heatmap_tables.get('columns', [])
                                            logger.info(f"Found Correlation Matrix heatmap (direct dict) with {len(rows)} rows, {len(columns)} columns")
                                        elif corr_heatmap_tables is not None:
                                            logger.debug(f"correlation_matrix_heatmap is not a list or dict: {type(corr_heatmap_tables)}, value: {str(corr_heatmap_tables)[:200]}")
                                    
                                    # Extract VIF analysis table (from response_data) - NEW
                                    # Structure: response_data["vif_analysis"] = [{"columns": [...], "rows": [...], "title": ...}]
                                    if isinstance(inner_response_data, dict) and not vif_analysis_table:
                                        vif_tables = inner_response_data.get("vif_analysis", [])
                                        logger.debug(f"Found vif_analysis key: {vif_tables is not None}, type: {type(vif_tables)}")
                                        if isinstance(vif_tables, list) and len(vif_tables) > 0:
                                            # Get the first table (should be the VIF analysis)
                                            vif_table_obj = vif_tables[0]
                                            if isinstance(vif_table_obj, dict):
                                                vif_analysis_table = vif_table_obj
                                                rows = vif_table_obj.get('rows', [])
                                                columns = vif_table_obj.get('columns', [])
                                                logger.info(f"Found VIF analysis table with {len(rows)} rows, {len(columns)} columns")
                                                if len(rows) > 0:
                                                    logger.debug(f"Sample VIF row: {rows[0]}")
                                        elif isinstance(vif_tables, dict):
                                            # Sometimes it might be a dict directly, not a list
                                            vif_analysis_table = vif_tables
                                            rows = vif_tables.get('rows', [])
                                            columns = vif_tables.get('columns', [])
                                            logger.info(f"Found VIF analysis table (direct dict) with {len(rows)} rows, {len(columns)} columns")
                                        elif vif_tables is not None:
                                            logger.debug(f"vif_analysis is not a list or dict: {type(vif_tables)}, value: {str(vif_tables)[:200]}")

                                except json.JSONDecodeError as je:
                                    logger.debug(f"Failed to parse response_str as JSON: {str(je)}")
                                    pass
                            
                            # Fallback: try direct access from parsed["data"]
                            if not aggregated_insights or not iv_aggregated_insights or not correlation_aggregated_insights:
                                data_meta = parsed.get("data", {})
                                if isinstance(data_meta, dict):
                                    if not aggregated_insights:
                                        bivariate_insight_list = data_meta.get("bivariate_insight", [])
                                        if isinstance(bivariate_insight_list, list) and len(bivariate_insight_list) > 0:
                                            aggregated_insights = bivariate_insight_list
                                            logger.info(f"Found {len(aggregated_insights)} bivariate insights (fallback)")
                                    
                                    if not iv_aggregated_insights:
                                        iv_insight_list = data_meta.get("iv_insight", [])
                                        if isinstance(iv_insight_list, list) and len(iv_insight_list) > 0:
                                            iv_aggregated_insights = iv_insight_list
                                            logger.info(f"Found {len(iv_aggregated_insights)} IV insights (fallback)")
                                    
                                    if not correlation_aggregated_insights:
                                        corr_insight_list = data_meta.get("correlation_matrix_insight", [])
                                        if isinstance(corr_insight_list, list) and len(corr_insight_list) > 0:
                                            correlation_aggregated_insights = corr_insight_list
                                            logger.info(f"Found {len(correlation_aggregated_insights)} Correlation Matrix insights (fallback)")
                                    # Extract correlation analysis insights (fallback) - NEW
                                    if not correlation_analysis_insights:
                                        corr_analysis_insight_list = data_meta.get("correlation_insight", [])
                                        if isinstance(corr_analysis_insight_list, list) and len(corr_analysis_insight_list) > 0:
                                            correlation_analysis_insights = corr_analysis_insight_list
                                            logger.info(f"Found {len(correlation_analysis_insights)} Correlation Analysis insights (fallback)")
                            
                            # Additional fallback: check if parsed itself is the inner_response structure
                            if not aggregated_insights or not iv_aggregated_insights or not iv_analysis_summary or not correlation_aggregated_insights or not correlation_matrix_heatmap:
                                inner_data = parsed.get("data", {})
                                inner_response_data_fallback = parsed.get("response", {})
                                
                                if isinstance(inner_data, dict):
                                    if not aggregated_insights:
                                        bivariate_insight_list = inner_data.get("bivariate_insight", [])
                                        if isinstance(bivariate_insight_list, list) and len(bivariate_insight_list) > 0:
                                            aggregated_insights = bivariate_insight_list
                                    
                                    if not iv_aggregated_insights:
                                        iv_insight_list = inner_data.get("iv_insight", [])
                                        if isinstance(iv_insight_list, list) and len(iv_insight_list) > 0:
                                            iv_aggregated_insights = iv_insight_list
                                    
                                    if not correlation_aggregated_insights:
                                        corr_insight_list = inner_data.get("correlation_matrix_insight", [])
                                        if isinstance(corr_insight_list, list) and len(corr_insight_list) > 0:
                                            correlation_aggregated_insights = corr_insight_list

                                    # Extract correlation analysis insights (fallback) - NEW
                                    if not correlation_analysis_insights:
                                        corr_analysis_insight_list = inner_data.get("correlation_insight", [])
                                        if isinstance(corr_analysis_insight_list, list) and len(corr_analysis_insight_list) > 0:
                                            correlation_analysis_insights = corr_analysis_insight_list
                                            logger.info(f"Found {len(correlation_analysis_insights)} Correlation Analysis insights (fallback)")

                                # Also check for IV analysis summary in parsed["response"]
                                if not iv_analysis_summary and isinstance(inner_response_data_fallback, dict):
                                    iv_summary_tables = inner_response_data_fallback.get("iv_analysis_summary", [])
                                    if isinstance(iv_summary_tables, list) and len(iv_summary_tables) > 0:
                                        iv_table = iv_summary_tables[0]
                                        if isinstance(iv_table, dict):
                                            iv_analysis_summary = iv_table
                                            logger.info(f"Found IV analysis summary (fallback) with {len(iv_table.get('rows', []))} variables")
                                
                                # Also check for correlation matrix heatmap in parsed["response"]
                                if not correlation_matrix_heatmap and isinstance(inner_response_data_fallback, dict):
                                    corr_heatmap_tables = inner_response_data_fallback.get("correlation_matrix_heatmap", [])
                                    if isinstance(corr_heatmap_tables, list) and len(corr_heatmap_tables) > 0:
                                        corr_table = corr_heatmap_tables[0]
                                        if isinstance(corr_table, dict):
                                            correlation_matrix_heatmap = corr_table
                                            logger.info(f"Found Correlation Matrix heatmap (fallback) with {len(corr_table.get('rows', []))} variables")
                                # Also check for correlation numeric table in parsed["response"] - NEW
                                if not correlation_numeric_table and isinstance(inner_response_data_fallback, dict):
                                    corr_numeric_tables = inner_response_data_fallback.get("correlation_numeric", [])
                                    if isinstance(corr_numeric_tables, list) and len(corr_numeric_tables) > 0:
                                        corr_numeric_table_obj = corr_numeric_tables[0]
                                        if isinstance(corr_numeric_table_obj, dict):
                                            correlation_numeric_table = corr_numeric_table_obj
                                            logger.info(f"Found Correlation Numeric table (fallback) with {len(corr_numeric_table_obj.get('rows', []))} rows")
                                
                                # Also check for VIF analysis table in parsed["response"] - NEW
                                if not vif_analysis_table and isinstance(inner_response_data_fallback, dict):
                                    vif_tables = inner_response_data_fallback.get("vif_analysis", [])
                                    if isinstance(vif_tables, list) and len(vif_tables) > 0:
                                        vif_table_obj = vif_tables[0]
                                        if isinstance(vif_table_obj, dict):
                                            vif_analysis_table = vif_table_obj
                                            logger.info(f"Found VIF analysis table (fallback) with {len(vif_table_obj.get('rows', []))} rows")

                        except (json.JSONDecodeError, KeyError, TypeError) as e:
                            logger.debug(f"Error parsing chat message for insights: {str(e)}")
                            continue
                    
                    # Continue searching - don't break early, we need to find all insights
                    # (We'll break at the end of the loop if we found everything)
                    
                    # Log progress
                    if aggregated_insights or iv_aggregated_insights or iv_analysis_summary or correlation_aggregated_insights or correlation_matrix_heatmap or vif_aggregated_insights or vif_analysis_table:
                        logger.info(f"Progress: bivariate={len(aggregated_insights) if aggregated_insights else 0}, iv_insights={len(iv_aggregated_insights) if iv_aggregated_insights else 0}, iv_summary={'found' if iv_analysis_summary else 'not found'}, corr_insights={len(correlation_aggregated_insights) if correlation_aggregated_insights else 0}, corr_heatmap={'found' if correlation_matrix_heatmap else 'not found'}, vif_insights={len(vif_aggregated_insights) if vif_aggregated_insights else 0}, vif_table={'found' if vif_analysis_table else 'not found'}")
                        
        except Exception as e:
            logger.warning(f"Could not get aggregated insights from chat history: {str(e)}")
        
        # Also check messages array (LangChain format) if chat_history was empty - NEW
        if (not correlation_numeric_table or not correlation_analysis_insights or not vif_analysis_table or not vif_aggregated_insights) and messages:
            logger.info(f"Chat history was empty, checking {len(messages)} messages from state for correlation_numeric/vif data")
            try:
                for msg in reversed(messages):  # Check latest first
                    # Check if it's an AIMessage (assistant message)
                    if hasattr(msg, 'content') and hasattr(msg, '__class__') and 'AI' in str(msg.__class__):
                        msg_content = str(msg.content) if msg.content else ""
                        if msg_content and ("correlation_numeric" in msg_content or "correlation_insight" in msg_content or "vif_analysis" in msg_content or "vif_insight" in msg_content):
                            logger.info(f"Found correlation_numeric/correlation_insight/vif in AIMessage, attempting to parse...")
                            try:
                                # Try to parse as JSON
                                parsed = json.loads(msg_content)
                                
                                # Structure from agentic_system.py:
                                # payload = {"response": json.dumps(inner_response)}
                                # inner_response = {"response": response_data, "data": {"correlation_insight": [...], ...}}
                                
                                # First, check if parsed["response"] is a JSON string
                                response_str = parsed.get("response", "")
                                if isinstance(response_str, str):
                                    try:
                                        inner_response = json.loads(response_str)
                                        inner_data = inner_response.get("data", {})
                                        inner_response_data = inner_response.get("response", {})
                                        
                                        # Extract correlation analysis insights
                                        if not correlation_analysis_insights and isinstance(inner_data, dict):
                                            corr_analysis_insight_list = inner_data.get("correlation_insight", [])
                                            if isinstance(corr_analysis_insight_list, list) and len(corr_analysis_insight_list) > 0:
                                                correlation_analysis_insights = corr_analysis_insight_list
                                                logger.info(f"Found {len(correlation_analysis_insights)} Correlation Analysis LLM insights from messages array")
                                        
                                        # Extract correlation numeric table
                                        if not correlation_numeric_table and isinstance(inner_response_data, dict):
                                            corr_numeric_tables = inner_response_data.get("correlation_numeric", [])
                                            logger.debug(f"Found correlation_numeric key in messages: {corr_numeric_tables is not None}, type: {type(corr_numeric_tables)}")
                                            if isinstance(corr_numeric_tables, list) and len(corr_numeric_tables) > 0:
                                                corr_numeric_table_obj = corr_numeric_tables[0]
                                                if isinstance(corr_numeric_table_obj, dict):
                                                    correlation_numeric_table = corr_numeric_table_obj
                                                    rows = corr_numeric_table_obj.get('rows', [])
                                                    columns = corr_numeric_table_obj.get('columns', [])
                                                    logger.info(f"Found Correlation Numeric table from messages array with {len(rows)} rows, {len(columns)} columns")
                                                    if len(rows) > 0:
                                                        logger.debug(f"Sample correlation numeric row: {rows[0]}")
                                            elif isinstance(corr_numeric_tables, dict):
                                                correlation_numeric_table = corr_numeric_tables
                                                rows = corr_numeric_tables.get('rows', [])
                                                columns = corr_numeric_tables.get('columns', [])
                                                logger.info(f"Found Correlation Numeric table (direct dict) from messages array with {len(rows)} rows, {len(columns)} columns")
                                        
                                        # Extract VIF insights
                                        if not vif_aggregated_insights and isinstance(inner_data, dict):
                                            vif_insight_list = inner_data.get("vif_insight", [])
                                            if isinstance(vif_insight_list, list) and len(vif_insight_list) > 0:
                                                vif_aggregated_insights = vif_insight_list
                                                logger.info(f"Found {len(vif_aggregated_insights)} VIF LLM insights from messages array")
                                        
                                        # Extract VIF analysis table
                                        if not vif_analysis_table and isinstance(inner_response_data, dict):
                                            vif_tables = inner_response_data.get("vif_analysis", [])
                                            logger.debug(f"Found vif_analysis key in messages: {vif_tables is not None}, type: {type(vif_tables)}")
                                            if isinstance(vif_tables, list) and len(vif_tables) > 0:
                                                vif_table_obj = vif_tables[0]
                                                if isinstance(vif_table_obj, dict):
                                                    vif_analysis_table = vif_table_obj
                                                    rows = vif_table_obj.get('rows', [])
                                                    columns = vif_table_obj.get('columns', [])
                                                    logger.info(f"Found VIF analysis table from messages array with {len(rows)} rows, {len(columns)} columns")
                                                    if len(rows) > 0:
                                                        logger.debug(f"Sample VIF row: {rows[0]}")
                                            elif isinstance(vif_tables, dict):
                                                vif_analysis_table = vif_tables
                                                rows = vif_tables.get('rows', [])
                                                columns = vif_tables.get('columns', [])
                                                logger.info(f"Found VIF analysis table (direct dict) from messages array with {len(rows)} rows, {len(columns)} columns")
                                        
                                        # If we found everything we need, we can break
                                        if correlation_numeric_table and correlation_analysis_insights and vif_analysis_table and vif_aggregated_insights:
                                            break
                                    except json.JSONDecodeError as je:
                                        logger.debug(f"Failed to parse response_str as JSON from messages: {str(je)}")
                                        pass
                            except (json.JSONDecodeError, KeyError, TypeError) as e:
                                logger.debug(f"Error parsing correlation_numeric from AIMessage: {str(e)}")
                                pass
            except Exception as e:
                logger.warning(f"Could not get correlation_numeric from messages array: {str(e)}")
        
        # Log final results
        logger.info(f"Final extraction results: bivariate_insights={len(aggregated_insights) if aggregated_insights else 0}, iv_insights={len(iv_aggregated_insights) if iv_aggregated_insights else 0}, iv_summary={'found' if iv_analysis_summary else 'not found'}, corr_insights={len(correlation_aggregated_insights) if correlation_aggregated_insights else 0}, corr_heatmap={'found' if correlation_matrix_heatmap else 'not found'}, corr_analysis_insights={len(correlation_analysis_insights) if correlation_analysis_insights else 0}, corr_numeric_table={'found' if correlation_numeric_table else 'not found'}")
        
        # Fallback: Generate correlation matrix from dataset if user requested it but not found in chat history
        if user_requested_correlation and not correlation_matrix_heatmap:
            logger.info("Correlation matrix requested but not found in chat history - attempting to generate from dataset")
            try:
                from app.services.dataframe_state_manager import dataframe_state_manager
                from app.utils.helpers import generate_correlation_matrix_table
                import pandas as pd
                import numpy as np
                
                # Get dataframe
                df = dataframe_state_manager.get_dataframe(request.dataset_id)
                if df is None:
                    df = dataset_manager.load_dataset(request.dataset_id)
                
                if df is not None:
                    # Get target variable
                    dataset_info = dataset_manager.get_dataset_info(request.dataset_id)
                    target_variable = dataset_info.get('target_variable') if dataset_info else None
                    
                    # Select numeric columns only
                    numeric_df = df.select_dtypes(include='number')
                    numeric_df = numeric_df.replace([np.inf, -np.inf], np.nan).dropna(axis=1, how='all')
                    
                    # Remove ID-like columns
                    try:
                        n_rows = numeric_df.shape[0]
                        nunique_all = numeric_df.nunique(dropna=True)
                        id_like_cols = nunique_all[nunique_all == n_rows].index.tolist()
                        if id_like_cols:
                            numeric_df = numeric_df.drop(columns=id_like_cols)
                    except Exception as e:
                        logger.warning(f"Failed to drop ID-like columns: {e}")
                    
                    # Drop columns with no variance
                    nunique_series = numeric_df.nunique(dropna=True)
                    cols_with_variance = nunique_series[nunique_series > 1].index.tolist()
                    numeric_df = numeric_df[cols_with_variance]
                    
                    # Drop target variable if present
                    if target_variable and target_variable in numeric_df.columns:
                        numeric_df = numeric_df.drop(columns=[target_variable])
                    
                    if numeric_df.shape[1] >= 2:
                        # Calculate correlation matrix
                        corr_matrix = numeric_df.corr()
                        
                        # Generate correlation matrix table
                        variable_names = list(numeric_df.columns)
                        corr_table_result = generate_correlation_matrix_table(corr_matrix, variable_names)
                        
                        if corr_table_result and not corr_table_result.get("error"):
                            correlation_matrix_heatmap = corr_table_result.get("correlation_matrix")
                            if correlation_matrix_heatmap:
                                logger.info(f"Generated Correlation Matrix heatmap from dataset (fallback) with {len(correlation_matrix_heatmap.get('rows', []))} rows")
                            else:
                                logger.warning("Generated correlation table but correlation_matrix key not found")
                        else:
                            logger.warning(f"Failed to generate correlation matrix table: {corr_table_result.get('error', 'Unknown error')}")
                    else:
                        logger.warning("Not enough numeric variables for correlation matrix (fallback)")
                else:
                    logger.warning(f"Could not load dataset {request.dataset_id} for correlation matrix fallback")
            except Exception as e:
                logger.warning(f"Failed to generate correlation matrix from dataset (fallback): {str(e)}")
                import traceback
                logger.debug(traceback.format_exc())
        
        # Generate bivariate tables only if user requested bivariate
        bivariate_tables = []
        bivariate_analysis_data_from_chat = []  # Store bivariate_analysis array from chat history
        eda_report = []
        bivariate_insights = []
        
        if user_requested_bivariate:
            # First, try to extract bivariate_analysis array from chat history
            # This contains the exact data structure from the Excel file/chat window
            for idx, message in enumerate(reversed(chat_history)):
                # Handle old format (query/response/intent)
                if "response" in message and "role" not in message:
                    text_content = message.get("response", "")
                    if not isinstance(text_content, str):
                        text_content = str(text_content)
                # Handle new format (role/content)
                elif message.get("role") == "assistant":
                    content = message.get("content", "")
                    if isinstance(content, list) and len(content) > 0:
                        # Format: [{"type": "text", "text": "..."}]
                        text_content = content[0].get("text", "") if isinstance(content[0], dict) else str(content[0])
                    elif isinstance(content, str):
                        text_content = content
                    else:
                        text_content = str(content)
                else:
                    continue  # Skip non-assistant messages
                
                if "bivariate_analysis" in text_content and not bivariate_analysis_data_from_chat:
                    logger.info(f"Found 'bivariate_analysis' string in assistant message {idx}, attempting to parse...")
                    try:
                        parsed = json.loads(text_content)
                        response_str = parsed.get("response", "")
                        if isinstance(response_str, str):
                            inner_response = json.loads(response_str)
                            inner_response_data = inner_response.get("response", {})
                            if isinstance(inner_response_data, dict):
                                bivariate_array = inner_response_data.get("bivariate_analysis", [])
                                if isinstance(bivariate_array, list) and len(bivariate_array) > 0:
                                    bivariate_analysis_data_from_chat = bivariate_array
                                    logger.info(f"✓ Found {len(bivariate_analysis_data_from_chat)} bivariate analysis entries from chat history")
                                    break
                    except (json.JSONDecodeError, KeyError, TypeError) as e:
                        logger.debug(f"Error extracting bivariate_analysis from chat message {idx}: {str(e)}")
                        continue
            
            # If we found bivariate_analysis in chat history, use it; otherwise generate it
            if bivariate_analysis_data_from_chat and len(bivariate_analysis_data_from_chat) > 0:
                logger.info(f"Using bivariate_analysis data from chat history ({len(bivariate_analysis_data_from_chat)} entries)")
                bivariate_tables = bivariate_analysis_data_from_chat
            else:
                # Import the same function that generates the Excel file
                from app.utils.helpers import generate_bivariate_tables_for_standard_insights
                
                # Generate bivariate tables using the SAME function as the Excel file
                # This ensures we get the exact same data structure and insights
                bivariate_tables = generate_bivariate_tables_for_standard_insights(
                    dataset_id=request.dataset_id,
                    target_variable=target_variable,
                    top_categories=10,
                    bins=10,
                    binning_method="quantile"
                )
                logger.info(f"Generated {len(bivariate_tables)} bivariate tables (same as Excel file)")
            
            # Note: We don't filter by used_features here - let the frontend handle filtering
            # This allows users to see all data when selecting "All" or numeric filters
            logger.info(f"Building EDA report from {len(bivariate_tables)} bivariate tables (all variables, frontend will filter)")
            if used_features:
                logger.info(f"Used features available: {len(used_features)} features (frontend will filter based on user selection)")
            
            # Build EDA Report table from bivariate_tables (same structure as Excel file)
            logger.info(f"Building EDA report from {len(bivariate_tables)} bivariate tables")
            for table in bivariate_tables:
                variable_name = table.get('variable_name', '')
                if not variable_name:
                    logger.debug(f"Skipping table with no variable_name")
                    continue
                
                # Extract insights array (contains "Event rate range: X.XX% to Y.YY%" and "Monotonic decreasing pattern")
                insights = table.get('insights', [])
                if not isinstance(insights, list):
                    insights = []
                
                logger.debug(f"Processing variable {variable_name}: {len(insights)} insights found")
                if insights:
                    logger.debug(f"Sample insights for {variable_name}: {insights[:2]}")
                
                # Extract event rate range from insights (exact format from Excel: "Event rate range: 7.72% to 13.09%")
                event_rate_range = ""
                for insight in insights:
                    if isinstance(insight, str) and ("Event rate range" in insight or "event rate range" in insight.lower()):
                        # Extract the range (e.g., "7.72% to 13.09%")
                        # The insight format is: "Event rate range: 7.72% to 13.09%"
                        # Try regex first (more precise)
                        range_match = re.search(r'(\d+\.?\d*%?\s*to\s*\d+\.?\d*%)', insight, re.IGNORECASE)
                        if range_match:
                            event_rate_range = range_match.group(1).strip()
                            logger.debug(f"Extracted event rate range for {variable_name}: {event_rate_range}")
                        else:
                            # Fallback: extract from insight text directly
                            event_rate_range = insight.replace("Event rate range: ", "").replace("event rate range: ", "").strip()
                            logger.debug(f"Extracted event rate range (fallback) for {variable_name}: {event_rate_range}")
                        break
                
                # Extract meaningful insight text (e.g., "Monotonic decreasing pattern")
                insight_text = ""
                for insight in insights:
                    if not isinstance(insight, str):
                        continue
                    # Skip "Event rate range" and "Highest event rate" insights
                    if "Event rate range" in insight or "Highest event rate" in insight or "Lowest event rate" in insight:
                        continue
                    # Use pattern insights (e.g., "Monotonic decreasing pattern")
                    if "Monotonic" in insight or "monotonic" in insight.lower():
                        insight_text = insight
                        break
                    elif "pattern" in insight.lower() or "trend" in insight.lower():
                        insight_text = insight
                        break
                
                # If no pattern insight found, use the first non-event-rate insight
                if not insight_text:
                    for insight in insights:
                        if isinstance(insight, str) and "Event rate range" not in insight and "Highest event rate" not in insight and "Lowest event rate" not in insight:
                            insight_text = insight
                            break
                
                # Default insight if still empty
                if not insight_text:
                    insight_text = "Analysis completed"
                
                # Add to EDA report
                eda_report.append({
                    "Variable": variable_name,
                    "Event rate range": event_rate_range,
                    "Insight": insight_text
                })
            
            logger.info(f"Built EDA report with {len(eda_report)} entries")
            
            # Use aggregated insights if available, otherwise generate new ones
            bivariate_insights = aggregated_insights
            
            # If no aggregated insights, generate LLM insights
            if not bivariate_insights and eda_report:
                logger.info("Generating LLM insights for bivariate analysis")
                try:
                    # Create a summary of the analysis for LLM
                    analysis_summary = f"Bivariate analysis was performed on {len(eda_report)} variables against target variable {target_variable}. "
                    
                    # Format variables for prompt
                    vars_summary = "\n".join([
                        f"- {r['Variable']}: Event rate range {r.get('Event rate range', 'N/A')}"
                        for r in eda_report[:15]
                    ])
                    
                    prompt = f"""Based on the following bivariate analysis results, generate 5-7 concise bullet-point insights that summarize the overall patterns and findings across all variables. Focus on:
1. Overall patterns (monotonicity, consistency)
2. Event rate range variations
3. Notable variables with extreme ranges
4. Any anomalies or special cases
5. Overall assessment for modeling

Variables analyzed:
{vars_summary}

Generate insights in bullet format, one per line, starting with "- ". Each insight should be 1-2 sentences maximum."""
                    
                    llm_insights_text = await llm_service.generate_text(prompt=prompt)
                    # Parse bullet points
                    bivariate_insights = [
                        line.strip("- ").strip() 
                        for line in llm_insights_text.split("\n") 
                        if line.strip().startswith("-") and len(line.strip()) > 3
                    ]
                    logger.info(f"Generated {len(bivariate_insights)} LLM insights")
                except Exception as e:
                    logger.warning(f"Failed to generate LLM insights: {str(e)}")
        
        # Always generate IV analysis directly when user requested IV (similar to bivariate approach)
        # This ensures IV data is always available, matching the bivariate approach
        if user_requested_iv:
            if not iv_analysis_summary:
                logger.info("IV analysis summary not found in chat history - generating directly")
            else:
                logger.info("IV analysis summary found in chat history, but regenerating to ensure consistency")
            
            try:
                from app.utils.helpers import generate_iv_analysis_tables_pipeline_style
                
                # Generate IV analysis tables directly (same function used for Excel export)
                logger.info(f"Generating IV analysis for dataset {request.dataset_id}, target {target_variable}")
                iv_sections = generate_iv_analysis_tables_pipeline_style(
                    dataset_id=request.dataset_id,
                    target_variable=target_variable,
                    bins=10
                )
                
                logger.info(f"Generated {len(iv_sections)} IV sections")
                
                # Extract the summary table
                for section in iv_sections:
                    if section.get("analysis_kind") == "iv_analysis_summary":
                        iv_analysis_summary = {
                            "title": section.get("title", "Information Value (IV) Summary"),
                            "columns": section.get("columns", ["Feature Name", "IV"]),
                            "rows": section.get("rows", [])
                        }
                        logger.info(f"Generated IV analysis summary with {len(iv_analysis_summary.get('rows', []))} variables")
                        break
                
                # If no summary found, try to extract from any section with rows
                if not iv_analysis_summary and iv_sections:
                    for section in iv_sections:
                        if "rows" in section and "columns" in section:
                            # Check if it looks like a summary (has Feature Name and IV columns)
                            cols = section.get("columns", [])
                            if "Feature Name" in cols or "IV" in cols:
                                iv_analysis_summary = {
                                    "title": section.get("title", "Information Value (IV) Summary"),
                                    "columns": section.get("columns", []),
                                    "rows": section.get("rows", [])
                                }
                                logger.info(f"Extracted IV analysis summary from section: {len(iv_analysis_summary.get('rows', []))} variables")
                                break
                
                if not iv_analysis_summary:
                    logger.warning("IV sections generated but no iv_analysis_summary found - checking all sections")
                    if iv_sections:
                        logger.info(f"Available section analysis_kinds: {[s.get('analysis_kind') for s in iv_sections]}")
            except Exception as e:
                logger.error(f"Failed to generate IV analysis directly: {str(e)}", exc_info=True)
        
        # Build IV EDA Report table from IV analysis summary
        iv_eda_report = []
        if iv_analysis_summary:
            iv_rows = iv_analysis_summary.get("rows", [])
            logger.info(f"Building IV EDA report from {len(iv_rows)} rows")
            
            # Note: We don't filter by used_features here - let the frontend handle filtering
            # This allows users to see all data when selecting "All" or numeric filters
            logger.info(f"Building IV EDA report from {len(iv_rows)} rows (all variables, frontend will filter)")
            if used_features:
                logger.info(f"Used features available: {len(used_features)} features (frontend will filter based on user selection)")
            
            for row in iv_rows:
                feature_name = row.get("Feature Name", "")
                iv_value = row.get("IV", 0)
                if feature_name:
                    iv_eda_report.append({
                        "Variable": feature_name,
                        "IV": iv_value
                    })
            logger.info(f"Built IV EDA report with {len(iv_eda_report)} entries")
        else:
            logger.warning("IV analysis summary not found - cannot build IV EDA report")
        
        # Build Correlation Matrix EDA Report table from correlation_matrix_heatmap
        correlation_eda_report = []
        if correlation_matrix_heatmap:
            corr_rows = correlation_matrix_heatmap.get("rows", [])
            corr_columns = correlation_matrix_heatmap.get("columns", [])
            logger.info(f"Building Correlation Matrix EDA report from {len(corr_rows)} rows, {len(corr_columns)} columns")
            
            # Note: We don't filter by used_features here - let the frontend handle filtering
            # This allows users to see all data when selecting "All" or numeric filters
            # However, we still need to filter columns to keep the matrix manageable
            # But we'll keep all rows so frontend can filter them
            logger.info(f"Building Correlation Matrix EDA report from {len(corr_rows)} rows, {len(corr_columns)} columns (all variables, frontend will filter rows)")
            if used_features:
                logger.info(f"Used features available: {len(used_features)} features (frontend will filter rows based on user selection)")
            
            # Build EDA report: each row represents a variable, columns are all other variables
            # Structure: Variable | var1 | var2 | ... for all other variables
            for row_data in corr_rows:
                variable_name = row_data.get("Variable", "")
                if not variable_name:
                    continue
                
                # Create a row entry with Variable and all correlation values
                eda_row = {"Variable": variable_name}
                
                # Add correlation values for each column variable (skip "Variable" column)
                for col_name in corr_columns[1:]:  # Skip first column which is "Variable"
                    corr_value = row_data.get(col_name, None)
                    if corr_value is not None:
                        # Round to 4 decimal places for display
                        if isinstance(corr_value, (int, float)):
                            eda_row[col_name] = round(corr_value, 4)
                        else:
                            eda_row[col_name] = corr_value
                
                correlation_eda_report.append(eda_row)
            
            logger.info(f"Built Correlation Matrix EDA report with {len(correlation_eda_report)} entries")
        else:
            logger.warning("Correlation Matrix heatmap not found - cannot build Correlation Matrix EDA report")
        
        # Build Correlation Analysis EDA Report table from correlation_numeric table - NEW
        correlation_analysis_eda_report = []
        if correlation_numeric_table:
            corr_numeric_rows = correlation_numeric_table.get("rows", [])
            logger.info(f"Building Correlation Analysis EDA report from {len(corr_numeric_rows)} rows")
            
            # Build EDA report: Variable Name | Type of Variable | Pearson Coefficient | Spearman Coefficient
            for row_data in corr_numeric_rows:
                eda_row = {
                    "Variable Name": row_data.get("Variable Name", ""),
                    "Type of Variable": row_data.get("Type of Variable", ""),
                    "Pearson Coefficient": row_data.get("Pearson Coefficient", None),
                    "Spearman Coefficient": row_data.get("Spearman Coefficient", None)
                }
                
                # Round numeric values to 4 decimal places
                if eda_row["Pearson Coefficient"] is not None and isinstance(eda_row["Pearson Coefficient"], (int, float)):
                    eda_row["Pearson Coefficient"] = round(eda_row["Pearson Coefficient"], 4)
                if eda_row["Spearman Coefficient"] is not None and isinstance(eda_row["Spearman Coefficient"], (int, float)):
                    eda_row["Spearman Coefficient"] = round(eda_row["Spearman Coefficient"], 4)
                
                correlation_analysis_eda_report.append(eda_row)
            
            logger.info(f"Built Correlation Analysis EDA report with {len(correlation_analysis_eda_report)} entries")
        else:
            logger.warning("Correlation Numeric table not found - cannot build Correlation Analysis EDA report")
        
        # Build VIF EDA Report table from vif_analysis table - NEW
        vif_eda_report = []
        if vif_analysis_table:
            vif_rows = vif_analysis_table.get("rows", [])
            logger.info(f"Building VIF EDA report from {len(vif_rows)} rows")
            
            # Build EDA report: Variable | VIF | Interpretation
            for row_data in vif_rows:
                eda_row = {
                    "Variable": row_data.get("Variable", ""),
                    "VIF": row_data.get("VIF", None),
                    "Interpretation": row_data.get("Interpretation", "")
                }
                
                # Round VIF to 2 decimal places if numeric
                if eda_row["VIF"] is not None and isinstance(eda_row["VIF"], (int, float)):
                    eda_row["VIF"] = round(eda_row["VIF"], 2)
                
                vif_eda_report.append(eda_row)
            
            logger.info(f"Built VIF EDA report with {len(vif_eda_report)} entries")
        else:
            logger.warning("VIF analysis table not found - cannot build VIF EDA report")

        # Only include sections that user requested
        insights_data = {}
        
        # Add bivariate analysis section only if user requested it
        if user_requested_bivariate:
            insights_data["bivariate_analysis"] = {
                "insights": bivariate_insights if isinstance(bivariate_insights, list) else [],
                "eda_report": eda_report,
                "rows_to_show": 20  # Default value for row filter
            }
            logger.info(f"Added bivariate analysis section: {len(bivariate_insights)} insights, {len(eda_report)} EDA rows")
        
        # Add IV analysis section only if user requested it
        iv_insights = []
        if user_requested_iv:
            # Generate IV insights if not found in chat history
            iv_insights = iv_aggregated_insights if isinstance(iv_aggregated_insights, list) else []
            if not iv_insights and iv_eda_report:
                logger.info("Generating LLM insights for IV analysis")
                try:
                    # Create a summary of the IV analysis for LLM
                    analysis_summary = f"Information Value (IV) analysis was performed on {len(iv_eda_report)} variables against target variable {target_variable}. "
                    
                    # Format variables for prompt
                    vars_summary = "\n".join([
                        f"- {r['Variable']}: IV = {r.get('IV', 'N/A')}"
                        for r in iv_eda_report[:15]
                    ])
                    
                    prompt = f"""Based on the following Information Value (IV) analysis results, generate 5-7 concise bullet-point insights that summarize the overall patterns and findings. Focus on:
1. Overall IV strength distribution
2. Variables with strong predictive power (IV > 0.3)
3. Variables with weak predictive power (IV < 0.02)
4. Notable patterns or anomalies
5. Overall assessment for feature selection

Variables analyzed:
{vars_summary}

Generate insights in bullet format, one per line, starting with "- ". Each insight should be 1-2 sentences maximum."""
                    
                    llm_insights_text = await llm_service.generate_text(prompt=prompt)
                    # Parse bullet points
                    iv_insights = [
                        line.strip("- ").strip() 
                        for line in llm_insights_text.split("\n") 
                        if line.strip().startswith("-") and len(line.strip()) > 3
                    ]
                    logger.info(f"Generated {len(iv_insights)} LLM insights for IV")
                except Exception as e:
                    logger.warning(f"Failed to generate LLM insights for IV: {str(e)}")
            
            insights_data["iv_analysis"] = {
                "insights": iv_insights,
                "eda_report": iv_eda_report,
                "rows_to_show": 20  # Default value for row filter
            }
            logger.info(f"Added IV analysis section: {len(iv_insights)} insights, {len(iv_eda_report)} EDA rows")
        
        # Add Correlation Matrix section only if user requested it
        correlation_insights = []
        if user_requested_correlation:
            # Generate correlation insights if not found in chat history
            correlation_insights = correlation_aggregated_insights if isinstance(correlation_aggregated_insights, list) else []
            if not correlation_insights and correlation_eda_report:
                logger.info("Generating LLM insights for Correlation Matrix analysis")
                try:
                    # Create a summary of the correlation matrix for LLM
                    analysis_summary = f"Correlation Matrix analysis was performed on {len(correlation_eda_report)} variables. "
                    
                    # Format variables for prompt (show sample correlations)
                    vars_summary = "\n".join([
                        f"- {r['Variable']}: correlations with other variables"
                        for r in correlation_eda_report[:10]
                    ])
                    
                    prompt = f"""Based on the following Correlation Matrix analysis results, generate 5-7 concise bullet-point insights that summarize the overall patterns and findings. Focus on:
1. Overall correlation strength distribution
2. Variables with high correlations (multicollinearity concerns)
3. Variables with low correlations (independent features)
4. Notable correlation patterns or groups
5. Overall assessment for feature selection and multicollinearity

Variables analyzed:
{vars_summary}

Generate insights in bullet format, one per line, starting with "- ". Each insight should be 1-2 sentences maximum."""
                    
                    llm_insights_text = await llm_service.generate_text(prompt=prompt)
                    # Parse bullet points
                    correlation_insights = [
                        line.strip("- ").strip() 
                        for line in llm_insights_text.split("\n") 
                        if line.strip().startswith("-") and len(line.strip()) > 3
                    ]
                    logger.info(f"Generated {len(correlation_insights)} LLM insights for Correlation Matrix")
                except Exception as e:
                    logger.warning(f"Failed to generate LLM insights for Correlation Matrix: {str(e)}")
            
            insights_data["correlation_analysis"] = {
                "insights": correlation_insights if isinstance(correlation_insights, list) else [],
                "eda_report": correlation_eda_report,
                "rows_to_show": 20  # Default value for row filter
            }
            logger.info(f"Added Correlation Matrix analysis section: {len(correlation_insights)} insights, {len(correlation_eda_report)} EDA rows")
        # Add Correlation Analysis section only if user requested it - NEW
        correlation_analysis_insights_list = []
        if user_requested_correlation_analysis:
            # Generate correlation analysis insights if not found in chat history
            correlation_analysis_insights_list = correlation_analysis_insights if isinstance(correlation_analysis_insights, list) else []
            if not correlation_analysis_insights_list and correlation_analysis_eda_report:
                logger.info("Generating LLM insights for Correlation Analysis")
                try:
                    # Create a summary of the correlation analysis for LLM
                    analysis_summary = f"Correlation Analysis was performed on {len(correlation_analysis_eda_report)} variables. "
                    
                    # Format variables for prompt (show sample correlations)
                    vars_summary = "\n".join([
                        f"- {r.get('Variable Name', '')}: Pearson={r.get('Pearson Coefficient', 'N/A')}, Spearman={r.get('Spearman Coefficient', 'N/A')}"
                        for r in correlation_analysis_eda_report[:10]
                    ])
                    
                    prompt = f"""Based on the following Correlation Analysis results, generate 5-7 concise bullet-point insights that summarize the overall patterns and findings. Focus on:
1. Overall correlation strength distribution
2. Variables with strong correlations (positive or negative)
3. Variables with weak correlations
4. Notable patterns or relationships
5. Overall assessment for feature selection

Variables analyzed:
{vars_summary}

Generate insights in bullet format, one per line, starting with "- ". Each insight should be 1-2 sentences maximum."""
                    
                    llm_insights_text = await llm_service.generate_text(prompt=prompt)
                    # Parse bullet points
                    correlation_analysis_insights_list = [
                        line.strip("- ").strip() 
                        for line in llm_insights_text.split("\n") 
                        if line.strip().startswith("-") and len(line.strip()) > 3
                    ]
                    logger.info(f"Generated {len(correlation_analysis_insights_list)} LLM insights for Correlation Analysis")
                except Exception as e:
                    logger.warning(f"Failed to generate LLM insights for Correlation Analysis: {str(e)}")
            
            insights_data["correlation_analysis_numeric"] = {  # Changed from correlationAnalysisNumeric
                "insights": correlation_analysis_insights_list if isinstance(correlation_analysis_insights_list, list) else [],
                "eda_report": correlation_analysis_eda_report,
                "rows_to_show": 20  # Default value for row filter
            }
            logger.info(f"Added Correlation Analysis section: {len(correlation_analysis_insights_list)} insights, {len(correlation_analysis_eda_report)} EDA rows")
        
        # Add VIF Analysis section only if user requested it - NEW
        vif_insights_list = []
        if user_requested_vif:
            # Generate VIF insights if not found in chat history
            vif_insights_list = vif_aggregated_insights if isinstance(vif_aggregated_insights, list) else []
            if not vif_insights_list and vif_eda_report:
                logger.info("Generating LLM insights for VIF Analysis")
                try:
                    # Create a summary of the VIF analysis for LLM
                    analysis_summary = f"VIF (Variance Inflation Factor) analysis was performed on {len(vif_eda_report)} variables. "
                    
                    # Format variables for prompt
                    vars_summary = "\n".join([
                        f"- {r.get('Variable', '')}: VIF={r.get('VIF', 'N/A')}, {r.get('Interpretation', '')}"
                        for r in vif_eda_report[:15]
                    ])
                    
                    prompt = f"""Based on the following VIF (Variance Inflation Factor) analysis results, generate 5-7 concise bullet-point insights that summarize the overall patterns and findings. Focus on:
1. Overall VIF distribution and multicollinearity severity
2. Variables with severe multicollinearity (VIF > 10)
3. Variables with acceptable VIF (VIF < 5)
4. Notable patterns or groups of correlated variables
5. Overall assessment for feature selection and multicollinearity concerns

Variables analyzed:
{vars_summary}

Generate insights in bullet format, one per line, starting with "- ". Each insight should be 1-2 sentences maximum."""
                    
                    llm_insights_text = await llm_service.generate_text(prompt=prompt)
                    # Parse bullet points
                    vif_insights_list = [
                        line.strip("- ").strip() 
                        for line in llm_insights_text.split("\n") 
                        if line.strip().startswith("-") and len(line.strip()) > 3
                    ]
                    logger.info(f"Generated {len(vif_insights_list)} LLM insights for VIF Analysis")
                except Exception as e:
                    logger.warning(f"Failed to generate LLM insights for VIF Analysis: {str(e)}")
            
            insights_data["vif_analysis"] = {
                "insights": vif_insights_list if isinstance(vif_insights_list, list) else [],
                "eda_report": vif_eda_report,
                "rows_to_show": 20  # Default value for row filter
            }
            logger.info(f"Added VIF Analysis section: {len(vif_insights_list)} insights, {len(vif_eda_report)} EDA rows")
        
        logger.info(f"Data insights retrieved: bivariate={user_requested_bivariate} ({len(eda_report)} EDA rows, {len(bivariate_insights)} insights), IV={user_requested_iv} ({len(iv_eda_report)} EDA rows, {len(iv_insights) if user_requested_iv else 0} insights), Correlation={user_requested_correlation} ({len(correlation_eda_report)} EDA rows, {len(correlation_insights) if user_requested_correlation else 0} insights), CorrelationAnalysis={user_requested_correlation_analysis} ({len(correlation_analysis_eda_report)} EDA rows, {len(correlation_analysis_insights_list) if user_requested_correlation_analysis else 0} insights), VIF={user_requested_vif} ({len(vif_eda_report)} EDA rows, {len(vif_insights_list) if user_requested_vif else 0} insights)")

        
        return DataInsightsResponse(
            success=True,
            insights=insights_data
        )
        
    except Exception as e:
        logger.error(f"Error getting data insights: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return DataInsightsResponse(
            success=False,
            error=str(e)
        )


def set_table_auto_width_and_wrap(table, num_cols: int):
    """
    Helper function to set auto column width and enable text wrapping for a table.
    This ensures columns adjust automatically and text wraps within cells.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches
    
    # Set table to auto-fit columns
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Set table layout to auto (auto-fit)
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'autofit')
        tblPr.append(tblLayout)
    else:
        tblLayout.set(qn('w:type'), 'autofit')
    
    # Set column widths proportionally and enable text wrapping for all cells
    # Calculate proportional widths based on content type
    # First column (Column name) gets more width, numeric columns get less
    column_widths = []
    if num_cols == 17:  # EDA Report or Implemented Quality Changes table
        # Column name: 1.5 inches, Data Types/Type: 1.0 inch, others: 0.6 inches each
        column_widths = [Inches(1.5), Inches(1.0)] + [Inches(0.6)] * 15
    else:
        # Default: equal distribution
        total_width = Inches(10)  # Total table width
        width_per_col = total_width / num_cols
        column_widths = [width_per_col] * num_cols
    
    # Set column widths
    for i, col in enumerate(table.columns):
        if i < len(column_widths):
            col.width = column_widths[i]
    
    # Enable text wrapping for all cells
    for row in table.rows:
        for cell in row.cells:
            # Get or create cell properties
            tcPr = cell._element.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell._element.insert(0, tcPr)
            
            # Set text wrapping (noWrap = False means wrap is enabled)
            noWrap = tcPr.find(qn('w:noWrap'))
            if noWrap is not None:
                tcPr.remove(noWrap)
            
            # Set vertical alignment to top for better text wrapping display
            vAlign = tcPr.find(qn('w:vAlign'))
            if vAlign is None:
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'top')
                tcPr.append(vAlign)
            else:
                vAlign.set(qn('w:val'), 'top')
            
            # Set cell margins for better text wrapping
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is None:
                tcMar = OxmlElement('w:tcMar')
                tcPr.append(tcMar)
            
            # Set small margins (0.1 inches) for all sides
            for margin_name in ['top', 'left', 'bottom', 'right']:
                margin = tcMar.find(qn(f'w:{margin_name}'))
                if margin is None:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '144')  # 0.1 inch in twips (1440 twips = 1 inch)
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)


@documentation_router.post("/documentation/download")
async def download_documentation(documentation_data: Dict[str, Any]):
    """
    Download documentation as .docx file
    """
    def safe_get_rows_to_show(rows_to_show_value, default=20, max_rows=None):
        """
        Safely convert rowsToShow to an integer for slicing.
        Handles string values like "all", "used_features", or numeric strings.
        """
        if rows_to_show_value is None:
            return default
        
        if isinstance(rows_to_show_value, str):
            # If it's a string like "all" or "used_features", return None to show all
            if rows_to_show_value.lower() in ['all', 'used_features']:
                return None
            # Try to parse as integer
            try:
                parsed = int(rows_to_show_value)
                return parsed if parsed > 0 else default
            except (ValueError, TypeError):
                return default
        elif isinstance(rows_to_show_value, (int, float)):
            # Convert to int and ensure it's positive
            parsed = int(rows_to_show_value)
            return parsed if parsed > 0 else default
        else:
            return default
    
    def safe_slice(data_list, rows_to_show_value, default=20):
        """
        Safely slice a list based on rowsToShow value.
        Returns the full list if rowsToShow is "all" or "used_features".
        """
        rows_to_show = safe_get_rows_to_show(rows_to_show_value, default)
        if rows_to_show is None:
            return data_list
        return data_list[:rows_to_show] if rows_to_show > 0 else data_list
    
    try:
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=500, detail="python-docx library not installed")
        
        logger.info("Generating documentation .docx file")
        
        # Create a temporary directory for Excel files
        temp_dir = tempfile.mkdtemp()
        excel_files = []  # Track all Excel files created
        
        # Create a new Document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Model Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Generated on: {documentation_data.get('meta', {}).get('lastUpdated', 'N/A')}").italic = True
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Add spacing
        
        # Extract objectives data
        objectives = documentation_data.get('objectives', {})
        model_objective = objectives.get('modelObjective', {})
        data_summary = objectives.get('dataSummary', {})
        
        # Check if OBJECTIVE section has data
        generated_objective = model_objective.get('generatedObjective', '')
        data_summary_content = data_summary.get('content', '')
        has_objective_data = bool(generated_objective or (data_summary_content and data_summary_content != 'Not provided'))
        
        # Add OBJECTIVE section
        # Section heading with light blue background (simulated with shading)
        heading = doc.add_heading('1. OBJECTIVE', 1)
        # Note: python-docx doesn't support background colors easily, so we'll use bold and larger font
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue color
        
        if not has_objective_data:
            # Show default message
            doc.add_paragraph('You skipped to the documentation, upload the Data first')
        else:
            # Model Objective sub-section
            doc.add_heading('Model Objective', 2)
            
            # Add generated objective content
            if generated_objective:
                doc.add_paragraph(generated_objective)
            else:
                doc.add_paragraph('Model objective will be generated after data summary and target definition are available.')
            
            # Data Summary sub-section
            doc.add_heading('Data Summary', 2)
            doc.add_paragraph(data_summary_content if data_summary_content else 'Not provided')
        
        doc.add_paragraph()  # Add spacing
        
        # MODEL DESIGN section
        model_design = documentation_data.get('modelDesign', {})
        data_overview = model_design.get('dataOverview', {})
        
        # Check if MODEL DESIGN section has data
        dataset_stats = data_overview.get('datasetStats', {})
        eda_report_table = data_overview.get('edaReport', {}).get('table', [])
        data_quality = data_overview.get('dataQuality', {})
        var_categorization = data_overview.get('variableCategorization', {})
        categories = var_categorization.get('categories', {})
        target_definition = model_design.get('targetDefinition', {})
        sampling_plan = model_design.get('samplingPlan', {})
        model_validation = model_design.get('modelValidation', {})
        has_model_design_data = bool(
            dataset_stats or 
            (eda_report_table and len(eda_report_table) > 0) or
            data_quality or
            (categories and len(categories) > 0) or
            target_definition or
            sampling_plan or
            model_validation
        )
        
        heading = doc.add_heading('2. MODEL DESIGN', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        if not has_model_design_data:
            # Show default message
            doc.add_paragraph('You skipped to the documentation, upload the Data first')
        else:
            # 2.1 Data Overview
            doc.add_heading('2.1 Data Overview', 2)
        
        # Dataset Statistics
        dataset_stats = data_overview.get('datasetStats', {})
        if dataset_stats:
            doc.add_paragraph('Dataset Information:', style='Heading 3')
            stats_para = doc.add_paragraph()
            stats_para.add_run(f"Total Rows: {dataset_stats.get('totalRows', 0):,}").bold = True
            stats_para.add_run(f" | Total Columns: {dataset_stats.get('totalColumns', 0)}")
            stats_para.add_run(f" | Numerical: {dataset_stats.get('numericalColumns', 0)}")
            stats_para.add_run(f" | Categorical: {dataset_stats.get('categoricalColumns', 0)}")
            stats_para.add_run(f" | Date: {dataset_stats.get('dateColumns', 0)}")
        
        # EDA Report - Create Excel file instead of table
        eda_report = data_overview.get('edaReport', {})
        eda_report_table = eda_report.get('table', [])
        if eda_report_table and len(eda_report_table) > 0:
            doc.add_paragraph('EDA Report:', style='Heading 3')
            
            # Handle rowsToShow - it can be a string ("all", "used_features") or a number
            table_data = safe_slice(eda_report_table, eda_report.get('rowsToShow', 20))
            
            # Create Excel file
            headers = ['Column', 'Data Types', 'Unique', 'Missing', 'Mean', 'Median', 'Mode', 'Std', 'Var', 'Min', 'p5%', 'p25%', 'p50%', 'p75%', 'p95%', 'p99%', 'Max']
            excel_filename = 'RawData_EDA_Report.xlsx'
            excel_path = os.path.join(temp_dir, excel_filename)
            
            create_excel_from_table_data(table_data, headers, excel_path)
            excel_files.append(excel_path)
            
            # Embed Excel file as linked object
            embed_excel_as_linked_object(doc, excel_path, 'RawData_EDA_Report')
            
            doc.add_paragraph()  # Add spacing
        
        # Data Quality Assessment
        data_quality = data_overview.get('dataQuality', {})
        if data_quality:
            doc.add_paragraph('Data Quality Assessment:', style='Heading 3')
            quality_summary = data_quality.get('summary', 'No quality assessment available')
            if quality_summary:
                # Parse bullet points (lines starting with "- " or "• ")
                lines = quality_summary.split('\n')
                for line in lines:
                    line = line.strip()
                    if line:
                        # Check if it's a bullet point
                        if line.startswith('- ') or line.startswith('• '):
                            # Remove bullet marker and add as bullet point
                            bullet_text = line[2:].strip() if line.startswith('- ') else line[2:].strip()
                            para = doc.add_paragraph(bullet_text, style='List Bullet')
                        else:
                            # Regular paragraph (for any intro text before bullets)
                            doc.add_paragraph(line)
            else:
                doc.add_paragraph('No quality assessment available')
        
        # Variable Categorization
        var_categorization = data_overview.get('variableCategorization', {})
        categories = var_categorization.get('categories', {})
        if categories and len(categories) > 0:
            doc.add_paragraph('Variable Categorization Distribution:', style='Heading 3')
            
            # Generate pie chart using matplotlib (same approach as Segment Proportions)
            try:
                import matplotlib
                matplotlib.use('Agg')
                import matplotlib.pyplot as plt
                
                # Prepare data for pie chart
                labels = list(categories.keys())
                data = list(categories.values())
                category_colors = var_categorization.get('colors', {})
                
                # Get colors for each category, with fallback defaults
                default_colors = ['#6366f1', '#3b82f6', '#10b981', '#f59e0b', '#ec4899', '#ef4444', '#8b5cf6', '#06b6d4']
                colors = [category_colors.get(label, default_colors[i % len(default_colors)]) for i, label in enumerate(labels)]
                
                if labels and data:
                    # Calculate percentages
                    total = sum(data)
                    percentages = [(d / total * 100) for d in data]
                    
                    fig, ax = plt.subplots(figsize=(6, 4))
                    # Generate pie chart without labels on segments (to avoid overlap)
                    pie_result = ax.pie(
                        data, 
                        labels=None,  # No labels on pie
                        autopct=None,  # No percentage text on pie segments
                        colors=colors[:len(data)], 
                        startangle=90
                    )
                    wedges = pie_result[0]
                    
                    # Add legend with labels and counts/percentages
                    legend_labels = [f'{label}: {count} variables ({percentages[i]:.1f}%)' for i, (label, count) in enumerate(zip(labels, data))]
                    ax.legend(wedges, legend_labels, loc='center left', bbox_to_anchor=(1, 0, 0.5, 1), fontsize=9)
                    # No title - heading is already in DOCX
                    plt.tight_layout()
                    
                    # Save to BytesIO
                    img_buffer = BytesIO()
                    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
                    img_buffer.seek(0)
                    
                    # Validate image data
                    if img_buffer.getvalue():
                        # Add image to document (6 inches width - 2x the original 3 inches, maintaining aspect ratio)
                        doc.add_picture(img_buffer, width=Inches(6))
                        
                        # Center the image
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        doc.add_paragraph()  # Add spacing after image
                        
                        logger.info(f"Variable Categorization pie chart generated and added successfully ({len(labels)} categories)")
                    else:
                        raise ValueError("Empty image buffer generated")
                    
                    plt.close()
                else:
                    logger.warning("Variable Categorization data is empty, skipping pie chart generation")
            except ImportError:
                logger.warning("matplotlib not available, skipping Variable Categorization pie chart generation")
            except Exception as e:
                logger.error(f"Failed to generate Variable Categorization pie chart: {str(e)}")
                import traceback
                logger.error(traceback.format_exc())
            
            doc.add_paragraph()  # Add spacing
        else:
            doc.add_paragraph('Variable Categorization Distribution:', style='Heading 3')
            doc.add_paragraph('Variable categorization not available. Generate the Knowledge Graph to see variable categories.')
        
        doc.add_paragraph()  # Add spacing
        
        # 2.2 Target Definition
        target_definition = model_design.get('targetDefinition', {})
        if target_definition:
            doc.add_heading('2.2 Target Definition', 2)
            
            # Target Variable Name
            target_name = target_definition.get('targetVariableName', 'Not configured')
            doc.add_paragraph('Target Variable Name:', style='Heading 3')
            target_para = doc.add_paragraph(target_name)
            target_para.runs[0].font.bold = True
            
            # Definition
            doc.add_paragraph('Definition:', style='Heading 3')
            definition = target_definition.get('definition', 'Not available')
            doc.add_paragraph(definition if definition else 'Not available')
            
            # Event Rate
            doc.add_paragraph('Event Rate:', style='Heading 3')
            event_rate = target_definition.get('eventRate', {})
            event_count = event_rate.get('eventCount', 0)
            total_count = event_rate.get('totalCount', 0)
            percentage = event_rate.get('percentage', 0)
            
            event_rate_para = doc.add_paragraph()
            event_rate_para.add_run(f"{event_count:,} / {total_count:,}").bold = True
            event_rate_para.add_run(f" ({percentage:.2f}%)")
            
        doc.add_paragraph()  # Add spacing
        
        # 2.3 Sampling Plan
        sampling_plan = model_design.get('samplingPlan', {})
        if sampling_plan:
            doc.add_heading('2.3 Sampling Plan', 2)
            
            # Generate or retrieve LLM writeup about sampling plan
            sampling_writeup = sampling_plan.get('writeup', '')
            if not sampling_writeup:
                try:
                    has_split = sampling_plan.get('hasSplit', False)
                    train_data = sampling_plan.get('train', {})
                    hold_data = sampling_plan.get('hold', {}) if has_split else {}
                    
                    train_total = train_data.get('total', 0)
                    train_event_count = train_data.get('eventCount', 0)
                    train_event_rate = train_data.get('eventRate', 0)
                    
                    if has_split:
                        hold_total = hold_data.get('total', 0)
                        hold_event_count = hold_data.get('eventCount', 0)
                        hold_event_rate = hold_data.get('eventRate', 0)
                        
                        # Calculate split ratio
                        total_samples = train_total + hold_total
                        train_pct = (train_total / total_samples * 100) if total_samples > 0 else 0
                        hold_pct = (hold_total / total_samples * 100) if total_samples > 0 else 0
                        
                        prompt = f"""You are a data scientist writing model documentation. Write a concise 3-4 line paragraph explaining the sampling plan for model training and validation.

Sampling Plan Details:
- Train and Hold samples are split in a {train_pct:.0f}-{hold_pct:.0f}% ratio
- Train sample: {train_total:,} total records, {train_event_count:,} events, {train_event_rate:.2f}% event rate
- Hold sample: {hold_total:,} total records, {hold_event_count:,} events, {hold_event_rate:.2f}% event rate
- Split is done randomly ensuring similar event rate across both samples

Write in a professional, modeller's voice explaining why this split ensures robust validation testing. Mention the event rate similarity between train and hold samples. Keep it concise (3-4 lines maximum)."""
                    else:
                        prompt = f"""You are a data scientist writing model documentation. Write a concise 3-4 line paragraph explaining the sampling plan.

Sampling Plan Details:
- Train sample: {train_total:,} total records, {train_event_count:,} events, {train_event_rate:.2f}% event rate
- No hold sample was created

Write in a professional, modeller's voice. Keep it concise (3-4 lines maximum)."""
                    
                    sampling_writeup = await llm_service.generate_text(prompt=prompt)
                    if sampling_writeup:
                        # Clean up the writeup (remove any markdown formatting, extra newlines)
                        sampling_writeup = sampling_writeup.strip()
                        # Remove markdown formatting if present
                        sampling_writeup = sampling_writeup.replace('**', '').replace('*', '')
                        # Store in documentation_data for frontend
                        if 'samplingPlan' not in model_design:
                            model_design['samplingPlan'] = {}
                        model_design['samplingPlan']['writeup'] = sampling_writeup
                except Exception as e:
                    logger.warning(f"Failed to generate sampling plan writeup: {str(e)}")
                    # Continue without writeup if LLM fails
            
            # Display writeup if available
            if sampling_writeup:
                # Split into paragraphs if multiple lines
                lines = [line.strip() for line in sampling_writeup.split('\n') if line.strip()]
                for line in lines:
                    doc.add_paragraph(line, style='Body Text')
                doc.add_paragraph()  # Add spacing before table
            
            # Create table
            has_split = sampling_plan.get('hasSplit', False)
            num_rows = 2 if has_split else 1  # +1 for header
            table = doc.add_table(rows=num_rows + 1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Sample'
            header_cells[1].text = 'Total'
            header_cells[2].text = 'Event'
            header_cells[3].text = 'Event Rate'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Train row
            train_data = sampling_plan.get('train', {})
            train_cells = table.rows[1].cells
            train_cells[0].text = 'Train'
            train_cells[1].text = f"{train_data.get('total', 0):,}"
            train_cells[2].text = f"{train_data.get('eventCount', 0):,}"
            train_cells[3].text = f"{train_data.get('eventRate', 0):.2f}%"
            
            # Validation row (if split exists)
            if has_split:
                validation_data = sampling_plan.get('validation', {})
                validation_cells = table.rows[2].cells
                validation_cells[0].text = 'Validation'
                validation_cells[1].text = f"{validation_data.get('total', 0):,}"
                validation_cells[2].text = f"{validation_data.get('eventCount', 0):,}"
                validation_cells[3].text = f"{validation_data.get('eventRate', 0):.2f}%"
            
            doc.add_paragraph()  # Add spacing after table
            
            # Sampling Identifier
            sampling_identifier = sampling_plan.get('samplingIdentifier', '')
            doc.add_paragraph('Sampling Identifier:', style='Heading 3')
            identifier_text = sampling_identifier if sampling_identifier else 'No Variable selected by user'
            doc.add_paragraph(identifier_text)
        
        doc.add_paragraph()  # Add spacing
        
        # 2.4 Model Validation (moved up, was 2.5)
        doc.add_heading('2.4 Model Validation', level=2)
        
        model_validation = model_design.get('modelValidation', {})
        has_hold_dataset = model_validation.get('hasHoldDataset', False)
        
        if not has_hold_dataset:
            doc.add_paragraph('No Hold dataset was available during Model Training', style='Body Text')
        else:
            best_model = model_validation.get('bestModel', {})
            model_name = best_model.get('modelName', 'Unknown')
            metrics = best_model.get('metrics', {})
            
            # Best Performing Model
            best_model_para = doc.add_paragraph()
            best_model_para.add_run('Best Performing Model: ')
            model_name_run = best_model_para.add_run(model_name)
            model_name_run.bold = True
            model_name_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue color
            
            doc.add_paragraph()  # Add spacing
            
            # Generate or retrieve LLM writeup explaining why this model is well-suited
            model_validation_writeup = model_validation.get('writeup', '')
            if not model_validation_writeup:
                try:
                    # Get data summary from objectives
                    data_summary_content = data_summary.get('content', '')
                    
                    # Extract key metrics
                    accuracy = metrics.get('accuracy', 0) * 100
                    precision = metrics.get('precision', 0)
                    recall = metrics.get('recall', 0)
                    f1_score = metrics.get('f1Score', 0)
                    auc_roc = metrics.get('aucRoc', 0)
                    auc_pr = metrics.get('aucPr', 0)
                    log_loss = metrics.get('logLoss', 0)
                    
                    prompt = f"""You are a data scientist writing model documentation. Write a concise 3-4 line paragraph explaining why the best performing model is well-suited for this use case.

Model Information:
- Model Name: {model_name}
- Performance Metrics on Hold Dataset:
  * Accuracy: {accuracy:.2f}%
  * Precision: {precision:.4f}
  * Recall: {recall:.4f}
  * F1 Score: {f1_score:.4f}
  * AUC-ROC: {auc_roc:.4f}
  * AUC-PR: {auc_pr:.4f}
  * Log Loss: {log_loss:.4f}

Data Summary Context:
{data_summary_content if data_summary_content else 'Not provided'}

Write in a professional, modeller's voice explaining why this model is well-suited for the problem. Reference the key performance metrics and how they align with the data characteristics and business objectives. Keep it concise (3-4 lines maximum)."""
                    
                    model_validation_writeup = await llm_service.generate_text(prompt=prompt)
                    if model_validation_writeup:
                        # Clean up the writeup (remove any markdown formatting, extra newlines)
                        model_validation_writeup = model_validation_writeup.strip()
                        # Remove markdown formatting if present
                        model_validation_writeup = model_validation_writeup.replace('**', '').replace('*', '')
                        # Store in documentation_data for frontend
                        if 'modelValidation' not in model_design:
                            model_design['modelValidation'] = {}
                        model_design['modelValidation']['writeup'] = model_validation_writeup
                except Exception as e:
                    logger.warning(f"Failed to generate model validation writeup: {str(e)}")
                    # Continue without writeup if LLM fails
            
            # Display writeup if available
            if model_validation_writeup:
                # Split into paragraphs if multiple lines
                lines = [line.strip() for line in model_validation_writeup.split('\n') if line.strip()]
                for line in lines:
                    doc.add_paragraph(line, style='Body Text')
                doc.add_paragraph()  # Add spacing before table
            
            # Metrics Table
            doc.add_paragraph('On Hold Dataset', style='Heading 3')
            
            metrics_table = doc.add_table(rows=2, cols=7)
            metrics_table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = metrics_table.rows[0].cells
            header_cells[0].text = 'Accuracy'
            header_cells[1].text = 'Precision'
            header_cells[2].text = 'Recall'
            header_cells[3].text = 'F1 Score'
            header_cells[4].text = 'AUC-ROC'
            header_cells[5].text = 'AUC-PR'
            header_cells[6].text = 'Log Loss'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data row
            data_cells = metrics_table.rows[1].cells
            data_cells[0].text = f"{(metrics.get('accuracy', 0) * 100):.2f}%"
            data_cells[1].text = f"{metrics.get('precision', 0):.4f}"
            data_cells[2].text = f"{metrics.get('recall', 0):.4f}"
            data_cells[3].text = f"{metrics.get('f1Score', 0):.4f}"
            data_cells[4].text = f"{metrics.get('aucRoc', 0):.4f}"
            data_cells[5].text = f"{metrics.get('aucPr', 0):.4f}"
            data_cells[6].text = f"{metrics.get('logLoss', 0):.4f}"
        
        # DATA TREATMENT section
        doc.add_paragraph()  # Add spacing before new section
        heading = doc.add_heading('3. DATA TREATMENT', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        model_design = documentation_data.get('modelDesign', {})
        data_treatment = model_design.get('dataTreatment', {})
        
        # Check if DATA TREATMENT section has data
        implemented_changes = data_treatment.get('implementedQualityChanges', {})
        writeup = implemented_changes.get('writeup')
        quality_check_plan = data_treatment.get('qualityCheckPlan', {})
        plan_table = quality_check_plan.get('table', [])
        column_stats = implemented_changes.get('columnStats', [])
        has_data_treatment_data = bool(
            (writeup and writeup.get('content')) or
            (plan_table and len(plan_table) > 0) or
            (column_stats and len(column_stats) > 0)
        )
        
        if not has_data_treatment_data:
            # Show default message
            doc.add_paragraph('You didn\'t clean your data! Go to Data Treatment page and pre-process the data.')
        else:
            # Writeup - moved before 3.1 Quality Check Plan
            implemented_changes = data_treatment.get('implementedQualityChanges', {})
            writeup = implemented_changes.get('writeup')
            if writeup and writeup.get('content'):
                writeup_para = doc.add_paragraph(writeup['content'])
                writeup_para.style = 'Body Text'
                doc.add_paragraph()  # Add spacing
            
            # 3.1 Quality Check Plan
            doc.add_heading('3.1 Quality Check Plan', 2)
            
            quality_check_plan = data_treatment.get('qualityCheckPlan', {})
            plan_table = quality_check_plan.get('table', [])
            
            if plan_table:
                # Create Excel file instead of table
                table_data = safe_slice(plan_table, quality_check_plan.get('rowsToShow', 20))
                headers = ['Issue', 'Variable', 'Observation', 'Treatment']
                excel_filename = 'Quality_Check_Plan.xlsx'
                excel_path = os.path.join(temp_dir, excel_filename)
                
                create_excel_from_table_data(table_data, headers, excel_path)
                excel_files.append(excel_path)
                
                # Embed Excel file as linked object
                embed_excel_as_linked_object(doc, excel_path, 'Quality_Check_Plan')
                
                doc.add_paragraph()  # Add spacing after embedded file
            else:
                doc.add_paragraph('No quality check plan data available.', style='Body Text')
            
            # 3.2 Implemented Quality Changes
            doc.add_heading('3.2 Implemented Quality Changes', 2)
            
            implemented_changes = data_treatment.get('implementedQualityChanges', {})
            column_stats = implemented_changes.get('columnStats', [])
            
            if column_stats:
                # Create Excel file instead of table
                table_data = safe_slice(column_stats, implemented_changes.get('rowsToShow', 20))
                headers = ['Column', 'Type', 'Missing', 'Unique', 'Mean', 'Median', 'Mode', 'Std', 'Var', 
                          'Min', 'p5%', 'p25%', 'p50%', 'p75%', 'p95%', 'p99%', 'Max']
                excel_filename = 'PreProcessedData_EDA_Report.xlsx'
                excel_path = os.path.join(temp_dir, excel_filename)
                
                create_excel_from_table_data(table_data, headers, excel_path)
                excel_files.append(excel_path)
                
                # Embed Excel file as linked object
                embed_excel_as_linked_object(doc, excel_path, 'PreProcessedData_EDA_Report')
                
                doc.add_paragraph()  # Add spacing after embedded file
            else:
                doc.add_paragraph('No column stats data available.', style='Body Text')
        
        # 4. DATA INSIGHTS section
        data_insights = documentation_data.get('dataInsights', {})
        
        # Check if DATA INSIGHTS section has data
        has_data_insights_data = False
        if data_insights:
            bivariate_analysis = data_insights.get('bivariateAnalysis')
            iv_analysis = data_insights.get('ivAnalysis')
            correlation_analysis = data_insights.get('correlationAnalysis')
            correlation_analysis_insights = data_insights.get('correlationAnalysisInsights')
            vif_analysis = data_insights.get('vifAnalysis')
            has_data_insights_data = bool(
                (bivariate_analysis and (bivariate_analysis.get('insights') or bivariate_analysis.get('edaReport'))) or
                (iv_analysis and (iv_analysis.get('insights') or iv_analysis.get('edaReport'))) or
                (correlation_analysis and correlation_analysis.get('edaReport')) or
                (correlation_analysis_insights and correlation_analysis_insights.get('insights')) or
                (vif_analysis and (vif_analysis.get('insights') or vif_analysis.get('edaReport')))
            )
        
        doc.add_paragraph()  # Add spacing before new section
        heading = doc.add_heading('4. DATA INSIGHTS', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        if not has_data_insights_data:
            # Show default message
            doc.add_paragraph('You didn\'t derive any insights on this data! Go to Data Insights page and generate your insights.')
        elif data_insights and has_data_insights_data:
            
            # 4.1 Bivariate Analysis
            bivariate_analysis = data_insights.get('bivariateAnalysis')
            if bivariate_analysis:
                doc.add_heading('4.1 Bivariate Analysis', 2)
                
                # Bullet point insights
                insights = bivariate_analysis.get('insights', [])
                if insights:
                    for insight in insights:
                        para = doc.add_paragraph(insight, style='List Bullet')
                
                doc.add_paragraph()  # Add spacing
                
                # 4.1.1 Insights - Create Excel file instead of table
                doc.add_heading('4.1.1 Insights', 3)
                eda_report = bivariate_analysis.get('edaReport', [])
                if eda_report:
                    # Create Excel file
                    table_data = safe_slice(eda_report, bivariate_analysis.get('rowsToShow', 20))
                    headers = ['Variable', 'Event rate range', 'Insight']
                    excel_filename = 'Bivariate_Analysis_Insights.xlsx'
                    excel_path = os.path.join(temp_dir, excel_filename)
                    
                    create_excel_from_table_data(table_data, headers, excel_path)
                    excel_files.append(excel_path)
                    
                    # Embed Excel file as linked object
                    embed_excel_as_linked_object(doc, excel_path, 'Bivariate_Analysis_Insights')
                else:
                    doc.add_paragraph('No EDA report data available.', style='Body Text')
            else:
                doc.add_paragraph('Bivariate Analysis data not available.', style='Body Text')
            
            # Information Value (IV) - Dynamic numbering based on what sections exist
            iv_analysis = data_insights.get('ivAnalysis')
            if iv_analysis:
                # Determine section number dynamically
                # Count existing sections: bivariate = 1, so IV = 2 if bivariate exists, else 1
                section_counter = 1
                if bivariate_analysis:
                    section_counter = 2
                # Future sections would increment this counter
                section_num = f'4.{section_counter}'
                subsection_num = f'{section_num}.1'
                
                doc.add_heading(f'{section_num} Information Value (IV)', 2)
                
                # Bullet point insights
                iv_insights = iv_analysis.get('insights', [])
                if iv_insights:
                    for insight in iv_insights:
                        para = doc.add_paragraph(insight, style='List Bullet')
                
                doc.add_paragraph()  # Add spacing
                
                # 4.x.1 Insights - Create Excel file instead of table
                doc.add_heading(f'{subsection_num} Insights', 3)
                iv_eda_report = iv_analysis.get('edaReport', [])
                if iv_eda_report:
                    # Create Excel file
                    table_data = safe_slice(iv_eda_report, iv_analysis.get('rowsToShow', 20))
                    headers = ['Variable', 'IV']
                    excel_filename = 'IV_Analysis_Insights.xlsx'
                    excel_path = os.path.join(temp_dir, excel_filename)
                    
                    create_excel_from_table_data(table_data, headers, excel_path)
                    excel_files.append(excel_path)
                    
                    # Embed Excel file as linked object
                    embed_excel_as_linked_object(doc, excel_path, 'IV_Analysis_Insights')
                else:
                    doc.add_paragraph('No EDA report data available.', style='Body Text')
            
            # Correlation Matrix - Should come before Correlation Analysis (Numeric)
            correlation_analysis = data_insights.get('correlationAnalysis')
            if correlation_analysis:
                # Determine section number dynamically
                # Count existing sections: bivariate = 1, IV = 2, correlation_matrix = 3
                section_counter = 1
                if bivariate_analysis:
                    section_counter = 2
                if iv_analysis:
                    section_counter = 3
                # Correlation Matrix is always 4.3 if bivariate and IV exist
                section_num = f'4.{section_counter}'
                subsection_num = f'{section_num}.1'
                
                doc.add_heading(f'{section_num} Correlation Matrix', 2)
                
                # Bullet point insights
                corr_insights = correlation_analysis.get('insights', [])
                if corr_insights:
                    for insight in corr_insights:
                        para = doc.add_paragraph(insight, style='List Bullet')
                
                doc.add_paragraph()  # Add spacing
                
                # 4.x.1 Insights - Create Excel file instead of table
                doc.add_heading(f'{subsection_num} Insights', 3)
                corr_eda_report = correlation_analysis.get('edaReport', [])
                if corr_eda_report:
                    # Get column names from first row (Variable + all other variables)
                    first_row = corr_eda_report[0] if corr_eda_report else {}
                    column_names = list(first_row.keys())
                    
                    if len(column_names) > 0:
                        # Create Excel file
                        table_data = safe_slice(corr_eda_report, correlation_analysis.get('rowsToShow', 20))
                        excel_filename = 'Correlation_Matrix_Insights.xlsx'
                        excel_path = os.path.join(temp_dir, excel_filename)
                        
                        create_excel_from_table_data(table_data, column_names, excel_path)
                        excel_files.append(excel_path)
                        
                        # Embed Excel file as linked object
                        embed_excel_as_linked_object(doc, excel_path, 'Correlation_Matrix_Insights')
                    else:
                        doc.add_paragraph('No correlation matrix data available.', style='Body Text')
                else:
                    doc.add_paragraph('No EDA report data available.', style='Body Text')
            
            # Correlation Analysis (Numeric) - Dynamic numbering based on what sections exist - NEW
            correlation_analysis_numeric = data_insights.get('correlationAnalysisNumeric')
            if correlation_analysis_numeric:
                # Determine section number dynamically
                # Count existing sections: bivariate = 1, IV = 2, correlation_matrix = 3, correlation_analysis = 4
                section_counter = 1
                if bivariate_analysis:
                    section_counter = 2
                if iv_analysis:
                    section_counter = 3
                if correlation_analysis:  # correlation_matrix
                    section_counter = 4
                # If correlation_analysis_numeric exists, it's the next one
                section_num = f'4.{section_counter}'
                subsection_num = f'{section_num}.1'
                
                doc.add_heading(f'{section_num} Correlation Analysis', 2)
                
                # Bullet point insights
                corr_analysis_insights = correlation_analysis_numeric.get('insights', [])
                if corr_analysis_insights:
                    for insight in corr_analysis_insights:
                        para = doc.add_paragraph(insight, style='List Bullet')
                
                doc.add_paragraph()  # Add spacing
                
                # 4.x.1 Insights - Create Excel file instead of table
                doc.add_heading(f'{subsection_num} Insights', 3)
                corr_analysis_eda_report = correlation_analysis_numeric.get('edaReport', [])
                if corr_analysis_eda_report:
                    # Create Excel file
                    table_data = safe_slice(corr_analysis_eda_report, correlation_analysis_numeric.get('rowsToShow', 20))
                    headers = ['Variable Name', 'Type of Variable', 'Pearson Coefficient', 'Spearman Coefficient']
                    excel_filename = 'Correlation_Analysis_Insights.xlsx'
                    excel_path = os.path.join(temp_dir, excel_filename)
                    
                    create_excel_from_table_data(table_data, headers, excel_path)
                    excel_files.append(excel_path)
                    
                    # Embed Excel file as linked object
                    embed_excel_as_linked_object(doc, excel_path, 'Correlation_Analysis_Insights')
                else:
                    doc.add_paragraph('No EDA report data available.', style='Body Text')
            
            # VIF Analysis - Dynamic numbering based on what sections exist - NEW
            vif_analysis = data_insights.get('vifAnalysis')
            if vif_analysis:
                # Determine section number dynamically
                # Count existing sections: bivariate = 1, IV = 2, correlation_matrix = 3, correlation_analysis = 4, vif = 5
                section_counter = 1
                if bivariate_analysis:
                    section_counter = 2
                if iv_analysis:
                    section_counter = 3
                if correlation_analysis:  # correlation_matrix
                    section_counter = 4
                if correlation_analysis_numeric:  # correlation_analysis
                    section_counter = 5
                # If vif_analysis exists, it's the next one
                section_num = f'4.{section_counter}'
                subsection_num = f'{section_num}.1'
                
                doc.add_heading(f'{section_num} VIF Analysis', 2)
                
                # Bullet point insights
                vif_insights = vif_analysis.get('insights', [])
                if vif_insights:
                    for insight in vif_insights:
                        para = doc.add_paragraph(insight, style='List Bullet')
                
                doc.add_paragraph()  # Add spacing
                
                # 4.x.1 Insights
                doc.add_heading(f'{subsection_num} Insights', 3)
                vif_eda_report = vif_analysis.get('edaReport', [])
                if vif_eda_report:
                    # Create Excel file instead of table
                    table_data = safe_slice(vif_eda_report, vif_analysis.get('rowsToShow', 20))
                    headers = ['Variable', 'VIF', 'Interpretation']
                    excel_filename = 'VIF_Analysis_Insights.xlsx'
                    excel_path = os.path.join(temp_dir, excel_filename)
                    
                    create_excel_from_table_data(table_data, headers, excel_path)
                    excel_files.append(excel_path)
                    
                    # Embed Excel file as linked object
                    embed_excel_as_linked_object(doc, excel_path, 'VIF_Analysis_Insights')
                else:
                    doc.add_paragraph('No EDA report data available.', style='Body Text')
        
        # 5. SEGMENTATION section (moved from 2.4)
        segmentation = model_design.get('segmentation', {})
        
        # Check if SEGMENTATION section has data
        has_segmentation = segmentation.get('hasSegmentation', False) if segmentation else False
        has_segmentation_data = bool(
            segmentation and has_segmentation and (
                segmentation.get('understanding') or
                segmentation.get('variablesUsed') or
                segmentation.get('method') or
                segmentation.get('segments') or
                segmentation.get('segmentSizesChart') or
                segmentation.get('segmentProportionsChart') or
                segmentation.get('ivVisualizationCharts')
            )
        )
        
        doc.add_paragraph()  # Add spacing before new section
        heading = doc.add_heading('5. SEGMENTATION', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        if not has_segmentation_data:
            # Show default message
            doc.add_paragraph('You didn\'t do Segmentation over your data! Go to Segmentation page and make segments of your data.')
        elif segmentation:
            if not has_segmentation:
                # No segmentation done
                no_seg_para = doc.add_paragraph('No Segmentation done by the user')
                no_seg_para.italic = True
            else:
                # Segmentation was performed
                # Understand the Segments Section (moved to first)
                understanding = segmentation.get('understanding')
                if understanding and understanding.get('content'):
                    doc.add_paragraph('Understand the Segments', style='Heading 3')
                    
                    # Parse and format the content with bullet points
                    content = understanding['content']
                    lines = content.split('\n')
                    
                    # First, add the introductory paragraph (lines that don't start with "- ")
                    intro_lines = []
                    bullet_lines = []
                    for line in lines:
                        line = line.strip()
                        if line:
                            if line.startswith('- '):
                                bullet_lines.append(line[2:])  # Remove "- " prefix
                            else:
                                intro_lines.append(line)
                    
                    # Add intro paragraph(s)
                    if intro_lines:
                        intro_text = ' '.join(intro_lines)
                        intro_para = doc.add_paragraph(intro_text)
                        intro_para.style = 'Body Text'
                    
                    # Add bullet points using List Bullet style
                    if bullet_lines:
                        for bullet_text in bullet_lines:
                            bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
                    
                    # Fallback: if no bullets detected, add as regular paragraph
                    if not intro_lines and not bullet_lines:
                        understanding_para = doc.add_paragraph(content)
                        understanding_para.style = 'Body Text'
                    
                    doc.add_paragraph()  # Add spacing
                
                # Variables used
                variables_used = segmentation.get('variablesUsed', [])
                if variables_used:
                    vars_para = doc.add_paragraph()
                    vars_para.add_run('Variables used: ').bold = True
                    vars_para.add_run(', '.join(variables_used) if isinstance(variables_used, list) else str(variables_used))
                    doc.add_paragraph()  # Add spacing
                
                # Method used
                method = segmentation.get('method', '')
                if method:
                    method_para = doc.add_paragraph()
                    method_para.add_run('Method used: ').bold = True
                    method_para.add_run(str(method).upper())
                    doc.add_paragraph()  # Add spacing
                
                # Segments
                segments = segmentation.get('segments', [])
                
                for segment in segments:
                    segment_number = segment.get('segmentNumber', 0)
                    rule = segment.get('rule', 'No rule specified')
                    total = segment.get('total', 0)
                    event_rate = segment.get('eventRate', 0)
                    segment_distribution = segment.get('segmentDistribution', 0)
                    
                    # Segment heading with rule
                    doc.add_paragraph(f'Segment {segment_number}: {rule}', style='Heading 3')
                    
                    # Create table for segment stats
                    seg_table = doc.add_table(rows=2, cols=3)
                    seg_table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = seg_table.rows[0].cells
                    header_cells[0].text = 'Total'
                    header_cells[1].text = 'Event Rate'
                    header_cells[2].text = 'Segment Distribution'
                    
                    # Make header bold
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data row
                    data_cells = seg_table.rows[1].cells
                    data_cells[0].text = f"{total:,}"
                    data_cells[1].text = f"{event_rate:.2f}%"
                    data_cells[2].text = f"{segment_distribution:.2f}%"
                    
                    doc.add_paragraph()  # Add spacing after table
                
                # Add Segment Sizes and Segment Proportions Charts (side by side)
                segment_sizes_chart = segmentation.get('segmentSizesChart')
                segment_proportions_chart = segmentation.get('segmentProportionsChart')
                
                if segment_sizes_chart and segment_proportions_chart:
                    doc.add_paragraph()  # Add spacing
                    
                    # Create a table to place charts side by side
                    charts_table = doc.add_table(rows=1, cols=2)
                    charts_table.style = 'Normal Table'
                    
                    # Remove borders properly using python-docx API
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    for row in charts_table.rows:
                        for cell in row.cells:
                            tcPr = cell._element.get_or_add_tcPr()
                            # Remove existing borders if any
                            existing_borders = tcPr.find(qn('w:tcBorders'))
                            if existing_borders is not None:
                                tcPr.remove(existing_borders)
                            # Add nil borders
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'nil')
                                border.set(qn('w:sz'), '0')
                                border.set(qn('w:space'), '0')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                    
                    # Chart 1: Segment Sizes (Bar Chart with Line for Event Rate)
                    try:
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        import numpy as np
                        
                        # Generate Segment Sizes bar chart with event rate line
                        labels = segment_sizes_chart.get('labels', [])
                        data = segment_sizes_chart.get('data', [])
                        event_rates = segment_sizes_chart.get('eventRates', [])
                        
                        if labels and data:
                            fig, ax1 = plt.subplots(figsize=(6, 4))
                            
                            # Bar chart for sizes
                            colors = ['#6366f1', '#3b82f6', '#10b981', '#f59e0b', '#ec4899', '#ef4444']
                            bars = ax1.bar(labels, data, color=colors[:len(data)], edgecolor=[c.replace('0.8', '1') for c in colors[:len(data)]], linewidth=2, alpha=0.8)
                            ax1.set_xlabel('Segment', fontsize=10)
                            ax1.set_ylabel('Total Records', fontsize=10, color='black')
                            ax1.tick_params(axis='y', labelcolor='black')
                            ax1.tick_params(axis='x', labelsize=10, rotation=45)
                            
                            # Line chart for event rates (if available)
                            if event_rates and len(event_rates) == len(data):
                                ax2 = ax1.twinx()
                                line = ax2.plot(labels, event_rates, color='#2563eb', linewidth=3, marker='o', markersize=6, label='Event Rate')
                                ax2.set_ylabel('Event Rate', fontsize=10, color='#2563eb')
                                ax2.tick_params(axis='y', labelcolor='#2563eb')
                                # Format event rate as percentage
                                ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda y, _: f'{y*100:.1f}%' if y <= 1 else f'{y:.1f}%'))
                            
                            plt.title('Segment Sizes', fontsize=12, fontweight='bold', pad=15)
                            plt.tight_layout()
                            
                            # Save to BytesIO
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            
                            # Validate image data
                            if img_buffer.getvalue():
                                # Insert image into first cell
                                cell1 = charts_table.rows[0].cells[0]
                                paragraph = cell1.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(img_buffer, width=Inches(3))
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                raise ValueError("Empty image buffer generated")
                            
                            plt.close()
                    except Exception as e:
                        logger.warning(f"Failed to generate Segment Sizes chart: {e}")
                        charts_table.rows[0].cells[0].text = "Segment Sizes chart not available"
                    
                    # Chart 2: Segment Proportions (Pie Chart)
                    try:
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        
                        # Generate Segment Proportions pie chart
                        labels = segment_proportions_chart.get('labels', [])
                        data = segment_proportions_chart.get('data', [])
                        colors = segment_proportions_chart.get('colors', ['#6366f1', '#3b82f6', '#10b981', '#f59e0b', '#ec4899', '#ef4444'])
                        
                        if labels and data:
                            # Convert percentages to decimals if needed (data should be 0-1 for pie chart)
                            data_decimal = [d / 100 if d > 1 else d for d in data]
                            
                            fig, ax = plt.subplots(figsize=(6, 4))
                            # Use legend instead of labels to avoid overlap
                            # When autopct=None, pie() returns only (wedges, texts), not (wedges, texts, autotexts)
                            pie_result = ax.pie(
                                data_decimal, 
                                labels=None,  # No labels on pie
                                autopct=None,  # No percentage text on pie segments
                                colors=colors[:len(data)], 
                                startangle=90
                            )
                            wedges = pie_result[0]
                            # Add legend with labels and percentages
                            legend_labels = [f'{label}: {data[i]:.1f}%' if data[i] > 1 else f'{label}: {data[i]*100:.1f}%' for i, label in enumerate(labels)]
                            ax.legend(wedges, legend_labels, loc='center left', bbox_to_anchor=(1, 0, 0.5, 1), fontsize=9)
                            ax.set_title('Segment Proportions', fontsize=12, fontweight='bold')
                            plt.tight_layout()
                            
                            # Save to BytesIO
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            
                            # Validate image data
                            if img_buffer.getvalue():
                                # Insert image into second cell
                                cell2 = charts_table.rows[0].cells[1]
                                paragraph = cell2.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(img_buffer, width=Inches(3))
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                raise ValueError("Empty image buffer generated")
                            
                            plt.close()
                    except Exception as e:
                        logger.warning(f"Failed to generate Segment Proportions chart: {e}")
                        charts_table.rows[0].cells[1].text = "Segment Proportions chart not available"
                    
                    doc.add_paragraph()  # Add spacing after charts
                
                # Add IV Visualization Charts
                iv_charts = segmentation.get('ivVisualizationCharts')
                iv_report = iv_charts.get('ivReport') if iv_charts else None
                
                if iv_report and iv_report.get('table'):
                    doc.add_paragraph()
                    doc.add_paragraph('IV Visualization Charts', style='Heading 3')
                    
                    try:
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        
                        # Create a grid layout for IV charts (2 columns)
                        iv_charts_list = []
                        
                        # 1. Weight of Evidence by Segment
                        try:
                            fig, ax = plt.subplots(figsize=(5, 3.5))
                            table = iv_report['table']
                            labels = [f'Segment_{r["segment_id"] + 1}' for r in table]
                            woes = [r['woe'] for r in table]
                            colors_woe = ['#22c55e' if w > 0 else '#ef4444' for w in woes]
                            ax.bar(labels, woes, color=colors_woe, alpha=0.8, edgecolor='black', linewidth=1)
                            ax.set_ylabel('Weight of Evidence (WOE)', fontsize=9)
                            ax.set_xlabel('Segment', fontsize=9)
                            ax.set_title('Weight of Evidence by Segment', fontsize=10, fontweight='bold')
                            ax.axhline(y=0, color='black', linestyle='-', linewidth=0.8)
                            plt.xticks(rotation=45, ha='right', fontsize=8)
                            plt.tight_layout()
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            plt.close()
                            iv_charts_list.append(('Weight of Evidence by Segment', img_buffer))
                        except Exception as e:
                            logger.warning(f"Failed to generate WoE chart: {e}")
                        
                        # 2. IV Components by Segment
                        try:
                            fig, ax = plt.subplots(figsize=(5, 3.5))
                            iv_contribs = [r['iv_contribution'] for r in table]
                            ax.bar(labels, iv_contribs, color='#22c55e', alpha=0.8, edgecolor='black', linewidth=1)
                            ax.set_ylabel('IV Component', fontsize=9)
                            ax.set_xlabel('Segment', fontsize=9)
                            ax.set_title('IV Components by Segment', fontsize=10, fontweight='bold')
                            plt.xticks(rotation=45, ha='right', fontsize=8)
                            plt.tight_layout()
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            if not img_buffer.getvalue():
                                raise ValueError("Empty image buffer generated")
                            plt.close()
                            iv_charts_list.append(('IV Components by Segment', img_buffer))
                        except Exception as e:
                            logger.warning(f"Failed to generate IV Components chart: {e}")
                        
                        # 3. Distribution of Good vs Bad by Segment
                        try:
                            fig, ax = plt.subplots(figsize=(5, 3.5))
                            dist_goods = [r['dist_goods'] * 100 for r in table]
                            dist_bads = [r['dist_bads'] * 100 for r in table]
                            x = np.arange(len(labels))
                            width = 0.35
                            ax.bar(x - width/2, dist_goods, width, label='% of Total Good', color='#22c55e', alpha=0.8)
                            ax.bar(x + width/2, dist_bads, width, label='% of Total Bad', color='#ef4444', alpha=0.8)
                            ax.set_ylabel('Percentage of Total', fontsize=9)
                            ax.set_xlabel('Segment', fontsize=9)
                            ax.set_title('Distribution of Good vs Bad by Segment', fontsize=10, fontweight='bold')
                            ax.set_xticks(x)
                            ax.set_xticklabels(labels, rotation=45, ha='right', fontsize=8)
                            ax.legend(fontsize=8)
                            plt.tight_layout()
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            if not img_buffer.getvalue():
                                raise ValueError("Empty image buffer generated")
                            plt.close()
                            iv_charts_list.append(('Distribution of Good vs Bad by Segment', img_buffer))
                        except Exception as e:
                            logger.warning(f"Failed to generate Good vs Bad chart: {e}")
                        
                        # 4. Bad Rate by Segment
                        try:
                            fig, ax = plt.subplots(figsize=(5, 3.5))
                            bad_rates = [r['bad_rate'] * 100 for r in table]
                            ax.bar(labels, bad_rates, color='#fbbf24', alpha=0.8, edgecolor='black', linewidth=1)
                            ax.set_ylabel('Bad Rate (%)', fontsize=9)
                            ax.set_xlabel('Segment', fontsize=9)
                            ax.set_title('Bad Rate by Segment', fontsize=10, fontweight='bold')
                            plt.xticks(rotation=45, ha='right', fontsize=8)
                            plt.tight_layout()
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            if not img_buffer.getvalue():
                                raise ValueError("Empty image buffer generated")
                            plt.close()
                            iv_charts_list.append(('Bad Rate by Segment', img_buffer))
                        except Exception as e:
                            logger.warning(f"Failed to generate Bad Rate chart: {e}")
                        
                        # 5. Population Distribution by Segment (Pie Chart)
                        try:
                            fig, ax = plt.subplots(figsize=(5, 3.5))
                            accounts = [r['accounts'] for r in table]
                            colors_pie = ['#22c55e', '#3b82f6', '#fbbf24', '#ef4444', '#a855f7', '#ec4899']
                            # Use legend instead of labels to avoid overlap
                            # When autopct=None, pie() returns only (wedges, texts), not (wedges, texts, autotexts)
                            pie_result = ax.pie(
                                accounts, 
                                labels=None,  # No labels on pie
                                autopct=None,  # No percentage text on pie segments
                                colors=colors_pie[:len(accounts)], 
                                startangle=90
                            )
                            wedges = pie_result[0]
                            # Calculate percentages for legend
                            total_accounts = sum(accounts)
                            legend_labels = [f'{label}: {acc/total_accounts*100:.1f}%' for label, acc in zip(labels, accounts)]
                            ax.legend(wedges, legend_labels, loc='center left', bbox_to_anchor=(1, 0, 0.5, 1), fontsize=8)
                            ax.set_title('Population Distribution by Segment', fontsize=10, fontweight='bold')
                            plt.tight_layout()
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            if not img_buffer.getvalue():
                                raise ValueError("Empty image buffer generated")
                            plt.close()
                            iv_charts_list.append(('Population Distribution by Segment', img_buffer))
                        except Exception as e:
                            logger.warning(f"Failed to generate Population Distribution chart: {e}")
                        
                        # Add charts in a grid (2 columns)
                        for i in range(0, len(iv_charts_list), 2):
                            row_charts = iv_charts_list[i:i+2]
                            if row_charts:
                                chart_row_table = doc.add_table(rows=1, cols=len(row_charts))
                                chart_row_table.style = 'Normal Table'
                                
                                for idx, (title, img_buffer) in enumerate(row_charts):
                                    cell = chart_row_table.rows[0].cells[idx]
                                    
                                    # Add title
                                    title_para = cell.paragraphs[0]
                                    title_run = title_para.add_run(title)
                                    title_run.font.bold = True
                                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Add image
                                    try:
                                        img_para = cell.add_paragraph()
                                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        run = img_para.add_run()
                                        run.add_picture(img_buffer, width=Inches(2.8))
                                    except Exception as e:
                                        logger.warning(f"Failed to insert IV chart image {title}: {e}")
                                        error_para = cell.add_paragraph()
                                        error_para.add_run("Chart not available")
                                
                                doc.add_paragraph()  # Add spacing between rows
                        
                        # 6. IV Strength Benchmark Chart
                        if iv_charts and iv_charts.get('ivStrength'):
                            try:
                                iv_strength = iv_charts['ivStrength']
                                iv_value = iv_strength.get('value', 0)
                                
                                fig, ax = plt.subplots(figsize=(5, 3.5))
                                labels = ['Not Useful\n(0-0.02)', 'Weak\n(0.02-0.1)', 'Medium\n(0.1-0.3)', 'Strong\n(0.3-0.5)', 'Suspicious\n(>0.5)']
                                data = [0.02, 0.08, 0.2, 0.2, 0.5]
                                colors = ['#ef4444', '#3b82f6', '#fbbf24', '#22c55e', '#7f1d1d']
                                
                                bars = ax.bar(labels, data, color=colors, alpha=0.7, edgecolor='black', linewidth=1)
                                ax.set_ylabel('IV Range', fontsize=9)
                                ax.set_xlabel('IV Strength Categories', fontsize=9)
                                ax.set_title(f'IV Strength: {iv_value:.4f} ({iv_strength.get("label", "Unknown")})', fontsize=10, fontweight='bold')
                                ax.set_ylim(0, 0.6)
                                
                                # Add indicator line for current IV value
                                ax.axhline(y=iv_value, color='blue', linestyle='--', linewidth=2, label=f'Current IV: {iv_value:.4f}')
                                ax.legend(fontsize=8, loc='upper right')
                                
                                plt.xticks(rotation=45, ha='right', fontsize=8)
                                plt.tight_layout()
                                
                                img_buffer = BytesIO()
                                plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                                img_buffer.seek(0)
                                
                                # Validate image data
                                if not img_buffer.getvalue():
                                    raise ValueError("Empty image buffer generated")
                                
                                plt.close()
                                
                                # Add chart in a new row (single column)
                                doc.add_paragraph()
                                chart_cell_table = doc.add_table(rows=1, cols=1)
                                chart_cell_table.style = 'Normal Table'
                                
                                # Remove borders properly
                                from docx.oxml import OxmlElement
                                from docx.oxml.ns import qn
                                for row in chart_cell_table.rows:
                                    for cell in row.cells:
                                        tcPr = cell._element.get_or_add_tcPr()
                                        # Remove existing borders if any
                                        existing_borders = tcPr.find(qn('w:tcBorders'))
                                        if existing_borders is not None:
                                            tcPr.remove(existing_borders)
                                        # Add nil borders
                                        tcBorders = OxmlElement('w:tcBorders')
                                        for border_name in ['top', 'left', 'bottom', 'right']:
                                            border = OxmlElement(f'w:{border_name}')
                                            border.set(qn('w:val'), 'nil')
                                            border.set(qn('w:sz'), '0')
                                            border.set(qn('w:space'), '0')
                                            tcBorders.append(border)
                                        tcPr.append(tcBorders)
                                
                                cell = chart_cell_table.rows[0].cells[0]
                                
                                title_para = cell.paragraphs[0]
                                title_run = title_para.add_run('IV Strength Benchmark')
                                title_run.font.bold = True
                                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                img_para = cell.add_paragraph()
                                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = img_para.add_run()
                                run.add_picture(img_buffer, width=Inches(4))
                            except Exception as e:
                                logger.warning(f"Failed to generate IV Strength chart: {e}")
                    except Exception as e:
                        logger.warning(f"Failed to generate IV Visualization Charts: {e}")
        
        # 6. FEATURE ENGINEERING/TRANSFORMATION section
        feature_engineering = documentation_data.get('featureEngineering', {})
        transformed_variables = feature_engineering.get('transformedVariables', [])
        
        # Check if FEATURE ENGINEERING section has data
        writeup = feature_engineering.get('writeup')
        has_feature_engineering_data = bool(
            (transformed_variables and len(transformed_variables) > 0) or
            (writeup and writeup.get('content'))
        )
        
        doc.add_paragraph()  # Add spacing before new section
        heading = doc.add_heading('6. FEATURE ENGINEERING/TRANSFORMATION', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        if not has_feature_engineering_data:
            # Show default message
            doc.add_paragraph('You didn\'t transform any features. Go to the Feature Engineering page and transform your features.')
        elif transformed_variables and len(transformed_variables) > 0:
            
            # Add "Understand the Transformations" write-up if available (moved before table)
            writeup = feature_engineering.get('writeup')
            if writeup and writeup.get('content'):
                doc.add_paragraph()
                doc.add_paragraph('Understand the Transformations', style='Heading 3')
                writeup_para = doc.add_paragraph(writeup['content'])
                writeup_para.style = 'Body Text'
                doc.add_paragraph()  # Add spacing before table
            
            # Create Excel file instead of table
            doc.add_paragraph()
            
            # Prepare table data with proper headers
            table_data_raw = safe_slice(transformed_variables, feature_engineering.get('rowsToShow', 20))
            # Convert data keys to match headers
            table_data = []
            for var in table_data_raw:
                table_data.append({
                    'New Transformed Variable': var.get('new_variable_name', ''),
                    'Var Type': var.get('var_type', ''),
                    'Variable definition': var.get('variable_definition', ''),
                    'Transformation method': var.get('transformation_methods', '')
                })
            
            headers = ['New Transformed Variable', 'Var Type', 'Variable definition', 'Transformation method']
            excel_filename = 'TransformedVariables.xlsx'
            excel_path = os.path.join(temp_dir, excel_filename)
            
            create_excel_from_table_data(table_data, headers, excel_path)
            excel_files.append(excel_path)
            
            # Embed Excel file as linked object
            embed_excel_as_linked_object(doc, excel_path, 'TransformedVariables')
        
        # MODEL DEVELOPMENT section
        doc.add_paragraph()  # Add spacing before new section
        heading = doc.add_heading('7. MODEL DEVELOPMENT', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        model_design = documentation_data.get('modelDesign', {})
        model_selection = model_design.get('modelSelection', {})
        
        # Check if MODEL DEVELOPMENT section has data
        has_model_development_data = bool(
            model_selection and (
                model_selection.get('metadata') or
                model_selection.get('narrative') or
                model_selection.get('finalVariables') or
                model_selection.get('hyperparameters')
            )
        )
        
        if not has_model_development_data:
            # Show default message
            doc.add_paragraph('You didn\'t train a model yet. Go to the Model Training page and build a model.')
        elif model_selection:
            # 7.1 Model Selection
            doc.add_heading('7.1 Model Selection', 2)
            
            # Algorithm and optimization method
            metadata = model_selection.get('metadata', {})
            if metadata.get('algorithm'):
                algo_para = doc.add_paragraph()
                algo_para.add_run(metadata.get('algorithm', '')).bold = True
            if metadata.get('optimizationMethod'):
                opt_para = doc.add_paragraph(metadata.get('optimizationMethod', ''))
            
            # Narrative
            narrative = model_selection.get('narrative', '')
            if narrative:
                doc.add_paragraph(narrative, style='Body Text')
            
            # Metadata grid
            doc.add_paragraph()  # Add spacing
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Iterations: {metadata.get('iterationCount', 0)}").bold = True
            meta_para.add_run(f" | Models tested: {metadata.get('totalModelsTested', 0)}")
            if metadata.get('scoreMetric') and metadata.get('bestScore') is not None:
                score_metric = metadata.get('scoreMetric', '')
                best_score = metadata.get('bestScore', 0)
                score_text = f"{best_score:.2f}" if score_metric == 'KS' else f"{best_score:.4f}"
                meta_para.add_run(f" | {score_metric}: {score_text}")
            if metadata.get('generatedAt'):
                from datetime import datetime
                try:
                    gen_date = datetime.fromisoformat(metadata['generatedAt'].replace('Z', '+00:00'))
                    meta_para.add_run(f" | Last updated: {gen_date.strftime('%Y-%m-%d')}")
                except:
                    pass
            
            # 7.1.1 Final Set of Variables
            doc.add_heading('7.1.1 Final Set of Variables', 3)
            final_variables = model_selection.get('finalVariables', {})
            total_count = final_variables.get('totalCount', 0)
            
            if total_count > 0:
                count_para = doc.add_paragraph()
                count_para.add_run(f'Total retained: {total_count}')
            
            # Variable Analysis Table - Create Excel file instead of table
            variable_analysis = final_variables.get('variableAnalysis', [])
            if variable_analysis and len(variable_analysis) > 0:
                # Prepare table data with proper formatting
                table_data = []
                for stat in variable_analysis:
                    row = {
                        'Variable Name': str(stat.get('variable', '')),
                        'Correlation': stat.get('correlation') if stat.get('correlation') is not None else 'N/A',
                        'VIF': stat.get('vif') if stat.get('vif') is not None else 'N/A',
                        'IV': stat.get('iv') if stat.get('iv') is not None else 'N/A',
                        'Interpretation': str(stat.get('interpretation', 'Normal'))
                    }
                    table_data.append(row)
                
                headers = ['Variable Name', 'Correlation', 'VIF', 'IV', 'Interpretation']
                excel_filename = 'ModelUsedFeatures.xlsx'
                excel_path = os.path.join(temp_dir, excel_filename)
                
                create_excel_from_table_data(table_data, headers, excel_path)
                excel_files.append(excel_path)
                
                # Embed Excel file as linked object
                embed_excel_as_linked_object(doc, excel_path, 'ModelUsedFeatures')
                
                doc.add_paragraph()  # Add spacing after embedded file
            else:
                # Fallback to categories if variable analysis not available
                categories = final_variables.get('categories', [])
                if categories:
                    for category in categories:
                        cat_para = doc.add_paragraph()
                        cat_para.add_run(f"{category.get('label', '')}: ").bold = True
                        cat_para.add_run(f"{category.get('count', 0)} - {category.get('description', '')}")
                else:
                    doc.add_paragraph('Variable list will be populated when training outputs are available.', style='Body Text')
            
            # Bivariate Analysis Charts (if available)
            bivariate_charts = final_variables.get('bivariateAnalysisCharts', {})
            if bivariate_charts and bivariate_charts.get('charts') and len(bivariate_charts.get('charts', [])) > 0:
                doc.add_heading('Bivariate Analysis', 4)
                charts_list = bivariate_charts.get('charts', [])
                variable_count = bivariate_charts.get('variableCount', 4)
                if variable_count != 'all' and isinstance(variable_count, int):
                    charts_list = charts_list[:variable_count]
                
                try:
                    import matplotlib
                    matplotlib.use('Agg')
                    import matplotlib.pyplot as plt
                    import numpy as np
                    
                    # Generate all chart images first
                    chart_images = []
                    for chart_data in charts_list:
                        try:
                            var_name = chart_data.get('variable_name', 'Unknown')
                            var_type = chart_data.get('variable_type', 'numerical')
                            viz_data = chart_data.get('visualization_data', {})
                            
                            if not viz_data or 'data' not in viz_data:
                                continue
                            
                            chart_info = viz_data.get('data', {})
                            categories = chart_info.get('categories', [])
                            bar_data = chart_info.get('bar_data', {})
                            line_data = chart_info.get('line_data', {})
                            
                            if not categories or not bar_data.get('values') or not line_data.get('values'):
                                continue
                            
                            # Create combination chart (bar + line)
                            fig, ax1 = plt.subplots(figsize=(5, 3.5))
                            
                            # Bar chart for Total
                            bar_values = bar_data.get('values', [])
                            bar_label = bar_data.get('label', 'Total')
                            
                            bars = ax1.bar(categories, bar_values, color='#22c55e', 
                                          edgecolor='#16a34a', linewidth=1, alpha=0.6, label=bar_label)
                            ax1.set_xlabel('Category' if var_type == 'categorical' else 'Bin Range (Decile)', fontsize=9)
                            ax1.set_ylabel('Total', fontsize=9, color='black')
                            ax1.tick_params(axis='y', labelcolor='black', labelsize=8)
                            ax1.tick_params(axis='x', labelsize=8, rotation=45)
                            # Set horizontal alignment for x-axis labels
                            for label in ax1.get_xticklabels():
                                label.set_ha('right')
                            
                            # Format Y-axis for Total
                            ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda y, _: f'{int(y):,}' if y >= 1000 else f'{int(y)}'))
                            
                            # Line chart for Event Rate (on secondary axis)
                            line_values = line_data.get('values', [])
                            line_label = line_data.get('label', 'Event Rate')
                            
                            ax2 = None
                            if line_values and len(line_values) == len(categories):
                                ax2 = ax1.twinx()
                                line = ax2.plot(categories, line_values, color='#3b82f6', 
                                              linewidth=2, marker='o', markersize=3, label=line_label)
                                ax2.set_ylabel('Event Rate', fontsize=9, color='#3b82f6')
                                ax2.tick_params(axis='y', labelcolor='#3b82f6', labelsize=8)
                                # Format Event Rate as percentage
                                ax2.yaxis.set_major_formatter(plt.FuncFormatter(lambda y, _: f'{y*100:.1f}%'))
                            
                            # Title
                            chart_title = viz_data.get('chart_title') or f'{var_name} Analysis'
                            plt.title(chart_title, fontsize=10, fontweight='bold', pad=8)
                            
                            # Legend
                            lines1, labels1 = ax1.get_legend_handles_labels()
                            lines2, labels2 = ax2.get_legend_handles_labels() if ax2 else ([], [])
                            ax1.legend(lines1 + lines2, labels1 + labels2, loc='upper left', fontsize=8)
                            
                            plt.tight_layout()
                            
                            # Save to BytesIO
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            plt.close()
                            
                            if img_buffer.getvalue():
                                chart_images.append((var_name, var_type, img_buffer))
                                logger.info(f"Generated bivariate chart for {var_name}")
                            else:
                                raise ValueError("Empty image buffer generated")
                                
                        except Exception as chart_error:
                            logger.warning(f"Failed to generate bivariate chart for {chart_data.get('variable_name', 'Unknown')}: {chart_error}")
                    
                    # Add charts in a grid layout (2 columns)
                    if chart_images:
                        for i in range(0, len(chart_images), 2):
                            row_charts = chart_images[i:i+2]
                            if row_charts:
                                # Create table for this row (2 columns)
                                chart_row_table = doc.add_table(rows=1, cols=len(row_charts))
                                chart_row_table.style = 'Normal Table'
                                
                                # Remove borders from table
                                from docx.oxml import OxmlElement
                                from docx.oxml.ns import qn
                                for row in chart_row_table.rows:
                                    for cell in row.cells:
                                        tcPr = cell._element.get_or_add_tcPr()
                                        # Remove existing borders if any
                                        existing_borders = tcPr.find(qn('w:tcBorders'))
                                        if existing_borders is not None:
                                            tcPr.remove(existing_borders)
                                        # Add nil borders
                                        tcBorders = OxmlElement('w:tcBorders')
                                        for border_name in ['top', 'left', 'bottom', 'right']:
                                            border = OxmlElement(f'w:{border_name}')
                                            border.set(qn('w:val'), 'nil')
                                            border.set(qn('w:sz'), '0')
                                            border.set(qn('w:space'), '0')
                                            tcBorders.append(border)
                                        tcPr.append(tcBorders)
                                
                                # Add charts to cells
                                for idx, (var_name, var_type, img_buffer) in enumerate(row_charts):
                                    cell = chart_row_table.rows[0].cells[idx]
                                    
                                    # Add title
                                    title_para = cell.paragraphs[0]
                                    title_run = title_para.add_run(f'{var_name} ({var_type})')
                                    title_run.font.bold = True
                                    title_run.font.size = Pt(10)
                                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Add image
                                    img_para = cell.add_paragraph()
                                    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = img_para.add_run()
                                    run.add_picture(img_buffer, width=Inches(2.8))
                                
                                doc.add_paragraph()  # Add spacing after row
                        
                        logger.info(f"Added {len(chart_images)} bivariate analysis charts to documentation in grid layout")
                    else:
                        doc.add_paragraph('No bivariate charts could be generated', style='Body Text')
                        
                except Exception as e:
                    logger.error(f"Error generating bivariate analysis charts: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                    doc.add_paragraph('Bivariate analysis charts could not be generated. Please check the web interface for charts.', style='Body Text')
                
                doc.add_paragraph()  # Add spacing
            
            # 7.1.2 Model parameters
            doc.add_heading('7.1.2 Model parameters', 3)
            hyperparameters = model_selection.get('hyperparameters', {})
            summary_list = hyperparameters.get('summaryList', [])
            
            if summary_list:
                # Create table for hyperparameters
                hp_table = doc.add_table(rows=1, cols=2)
                hp_table.style = 'Light Grid Accent 1'
                
                # Header row
                hp_header_cells = hp_table.rows[0].cells
                hp_header_cells[0].text = 'Parameter'
                hp_header_cells[1].text = 'Value'
                
                # Make header bold
                for cell in hp_header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data rows
                for param in summary_list:
                    hp_row_cells = hp_table.add_row().cells
                    hp_row_cells[0].text = str(param.get('name', ''))
                    hp_row_cells[1].text = str(param.get('value', ''))
                
                # Keep borders on table (Light Grid Accent 1 style already has borders)
                
                doc.add_paragraph()  # Add spacing
            
        doc.add_paragraph()  # Add spacing
        
        # MODEL PERFORMANCE section
        doc.add_paragraph()  # Add spacing before new section
        heading = doc.add_heading('8. MODEL PERFORMANCE', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        model_performance = documentation_data.get('modelPerformance', {})
        
        # Check if MODEL PERFORMANCE section has data
        features_data = model_performance.get('features', {})
        top_features = features_data.get('topFeatures', [])
        used_features = features_data.get('usedFeatures', [])
        models = model_performance.get('models', [])
        has_model_performance_data = bool(
            (top_features and len(top_features) > 0) or
            (used_features and len(used_features) > 0) or
            (models and len(models) > 0) or
            model_performance.get('decileAnalysis') or
            model_performance.get('monotonicity')
        )
        
        if not has_model_performance_data:
            # Show default message
            doc.add_paragraph('You didn\'t train a model yet. Go to the Model Evaluation page and to see how your models performed.')
        else:
            # 8.1 Features
            doc.add_heading('8.1 Features', 2)
        
        # Feature count text
        total_features = features_data.get('totalCount', 0)
        feature_count_para = doc.add_paragraph()
        feature_count_para.add_run(f'There are total ')
        bold_run = feature_count_para.add_run(str(total_features))
        bold_run.bold = True
        bold_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue color
        feature_count_para.add_run(' features being selected in the final model.')
        
        doc.add_paragraph()  # Add spacing
        
        # 8.1.1 Top Features
        doc.add_heading('8.1.1 Top Features', 3)
        
        top_features = features_data.get('topFeatures', [])
        top_n = features_data.get('topN', 20)
        
        if top_features and len(top_features) > 0:
            # Check if any features have descriptions
            has_descriptions = any(f.get('description') for f in top_features)
            
            # Create Excel file instead of table
            # Prepare table data
            table_data = []
            for feature in top_features:
                importance = feature.get('importance', 0)
                row = {
                    'Feature': feature.get('featureName', ''),
                    'Importance': importance
                }
                if has_descriptions:
                    row['Description'] = feature.get('description', '-')
                table_data.append(row)
            
            # Determine headers based on whether descriptions exist
            if has_descriptions:
                headers = ['Feature', 'Importance', 'Description']
            else:
                headers = ['Feature', 'Importance']
            
            excel_filename = 'FeatureImportance.xlsx'
            excel_path = os.path.join(temp_dir, excel_filename)
            
            create_excel_from_table_data(table_data, headers, excel_path)
            excel_files.append(excel_path)
            
            # Embed Excel file as linked object
            embed_excel_as_linked_object(doc, excel_path, 'FeatureImportance')
            
            doc.add_paragraph()  # Add spacing after embedded file
            
            logger.info(f"Created Excel file with {len(top_features)} features")
        else:
            doc.add_paragraph('No feature importance data available', style='Body Text')
        
        doc.add_paragraph()  # Add spacing
        
        # Feature Category Distribution Chart
        category_distribution = features_data.get('categoryDistribution', {})
        
        if category_distribution and len(category_distribution) > 0:
            # Add heading for the chart
            doc.add_heading('Feature Category Distribution', 3)
            
            try:
                import matplotlib
                matplotlib.use('Agg')  # Non-interactive backend
                import matplotlib.pyplot as plt
                import numpy as np
                from io import BytesIO as ChartBytesIO
                
                # Create bar chart
                categories = list(category_distribution.keys())
                counts = list(category_distribution.values())
                
                # Get colors if available
                category_colors_map = features_data.get('categoryColors', {})
                colors = [category_colors_map.get(cat, '#3b82f6') for cat in categories]
                
                # Create figure with appropriate size
                fig, ax = plt.subplots(figsize=(10, 6))
                
                # Create bar chart
                x_pos = np.arange(len(categories))
                bars = ax.bar(x_pos, counts, color=colors, edgecolor='black', linewidth=0.7)
                
                # Set labels and title
                ax.set_xlabel('Variable Category', fontsize=12, fontweight='bold')
                ax.set_ylabel('Number of Features', fontsize=12, fontweight='bold')
                ax.set_title('Feature Category Distribution', fontsize=14, fontweight='bold')
                
                # Set x-axis labels with rotation (slanted) to avoid overlap
                ax.set_xticks(x_pos)
                ax.set_xticklabels(categories, rotation=45, ha='right', fontsize=10)
                
                # Set y-axis ticks properly (integer values from 0 to max+1)
                max_count = max(counts) if counts else 0
                y_ticks = list(range(0, max_count + 2))
                ax.set_yticks(y_ticks)
                ax.set_yticklabels(y_ticks, fontsize=10)
                
                # Add grid for better readability
                ax.grid(axis='y', alpha=0.3, linestyle='--')
                
                # Add value labels on top of bars
                for i, (bar, count) in enumerate(zip(bars, counts)):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{int(count)}',
                           ha='center', va='bottom', fontsize=10, fontweight='bold')
                
                plt.tight_layout()
                
                # Save to BytesIO
                chart_io = ChartBytesIO()
                plt.savefig(chart_io, format='png', dpi=150, bbox_inches='tight')
                chart_io.seek(0)
                plt.close()
                
                # Insert chart into document
                doc.add_picture(chart_io, width=Inches(6.0))
                
                # Center the chart
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                logger.info(f"Added category distribution chart with {len(categories)} categories")
                
            except ImportError as e:
                logger.warning(f"matplotlib not available: {e}")
                # Add text-based representation
                for category, count in category_distribution.items():
                    doc.add_paragraph(f'{category}: {count} features', style='List Bullet')
            except Exception as e:
                logger.error(f"Error creating category distribution chart: {e}")
                import traceback
                logger.error(traceback.format_exc())
                # Add text-based representation
                for category, count in category_distribution.items():
                    doc.add_paragraph(f'{category}: {count} features', style='List Bullet')
        
        # 8.2 ROC Curve Comparison
        roc_curves = model_performance.get('rocCurves')
        if roc_curves:
            doc.add_paragraph()  # Add spacing
            doc.add_heading('8.2 ROC Curve Comparison', 3)
            
            # 8.2.1 Train
            train_roc = roc_curves.get('train', [])
            doc.add_heading('8.2.1 Train', 4)
            if train_roc and len(train_roc) > 0:
                try:
                    import matplotlib
                    matplotlib.use('Agg')
                    import matplotlib.pyplot as plt
                    import numpy as np
                    
                    fig, ax = plt.subplots(figsize=(8, 6))
                    
                    # Plot diagonal reference line
                    ax.plot([0, 1], [0, 1], 'k--', label='Random (AUC=0.5)', linewidth=1.5)
                    
                    # Plot ROC curves for each model
                    for model in train_roc:
                        roc_data = model.get('rocData', {})
                        fpr = roc_data.get('fpr', [])
                        tpr = roc_data.get('tpr', [])
                        auc = roc_data.get('auc', 0)
                        color = model.get('color', '#3b82f6')
                        model_name = model.get('modelName', 'Unknown')
                        
                        if fpr and tpr and len(fpr) == len(tpr):
                            ax.plot(fpr, tpr, label=f'{model_name} (AUC: {auc:.4f})', 
                                   color=color, linewidth=2.5)
                    
                    ax.set_xlabel('False Positive Rate', fontsize=11)
                    ax.set_ylabel('True Positive Rate', fontsize=11)
                    ax.set_title('ROC Curve Comparison (Train)', fontsize=12, fontweight='bold')
                    ax.set_xlim([0, 1])
                    ax.set_ylim([0, 1])
                    ax.legend(loc='lower right', fontsize=9)
                    ax.grid(True, alpha=0.3)
                    
                    img_buffer = BytesIO()
                    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    plt.close()
                    
                    if img_buffer.getvalue():
                        doc.add_picture(img_buffer, width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to generate Train ROC curve: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                    doc.add_paragraph('ROC Curve Comparison (Train) chart not available', style='Body Text')
            else:
                doc.add_paragraph('No train ROC curve data available', style='Body Text')
            
            # 8.2.2 Test
            test_roc = roc_curves.get('test', [])
            doc.add_heading('8.2.2 Test', 4)
            if test_roc and len(test_roc) > 0:
                try:
                    import matplotlib
                    matplotlib.use('Agg')
                    import matplotlib.pyplot as plt
                    import numpy as np
                    
                    fig, ax = plt.subplots(figsize=(8, 6))
                    
                    # Plot diagonal reference line
                    ax.plot([0, 1], [0, 1], 'k--', label='Random (AUC=0.5)', linewidth=1.5)
                    
                    # Plot ROC curves for each model
                    for model in test_roc:
                        roc_data = model.get('rocData', {})
                        fpr = roc_data.get('fpr', [])
                        tpr = roc_data.get('tpr', [])
                        auc = roc_data.get('auc', 0)
                        color = model.get('color', '#3b82f6')
                        model_name = model.get('modelName', 'Unknown')
                        
                        if fpr and tpr and len(fpr) == len(tpr):
                            ax.plot(fpr, tpr, label=f'{model_name} (AUC: {auc:.4f})', 
                                   color=color, linewidth=2.5)
                    
                    ax.set_xlabel('False Positive Rate', fontsize=11)
                    ax.set_ylabel('True Positive Rate', fontsize=11)
                    ax.set_title('ROC Curve Comparison (Test)', fontsize=12, fontweight='bold')
                    ax.set_xlim([0, 1])
                    ax.set_ylim([0, 1])
                    ax.legend(loc='lower right', fontsize=9)
                    ax.grid(True, alpha=0.3)
                    
                    img_buffer = BytesIO()
                    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    plt.close()
                    
                    if img_buffer.getvalue():
                        doc.add_picture(img_buffer, width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to generate Test ROC curve: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                    doc.add_paragraph('ROC Curve Comparison (Test) chart not available', style='Body Text')
            else:
                doc.add_paragraph('No test ROC curve data available', style='Body Text')
        
        # 8.3 Confusion Matrix Comparison
        confusion_matrices = model_performance.get('confusionMatrices', [])
        if confusion_matrices and len(confusion_matrices) > 0:
            doc.add_paragraph()  # Add spacing
            doc.add_heading('8.3 Confusion Matrix Comparison', 3)
            
            for model in confusion_matrices:
                model_name = model.get('modelName', 'Unknown')
                matrix = model.get('matrix', [])
                train_matrix = model.get('trainMatrix')
                color = model.get('color', '#3b82f6')
                accuracy = model.get('accuracy', 0)
                train_accuracy = model.get('trainAccuracy')
                f1_score = model.get('f1Score', 0)
                train_f1_score = model.get('trainF1Score')
                
                if matrix:
                    doc.add_paragraph()  # Add spacing
                    doc.add_paragraph(model_name, style='Heading 4')
                    
                    # Add Accuracy and F1 Score
                    accuracy_text = f"Accuracy: {accuracy:.1%}"
                    if train_accuracy is not None:
                        accuracy_text = f"Accuracy (Train/Test): {train_accuracy:.1%} / {accuracy:.1%}"
                    doc.add_paragraph(accuracy_text, style='Body Text')

                    f1_text = f"F1 Score: {f1_score:.3f}"
                    if train_f1_score is not None:
                        f1_text = f"F1 Score (Train/Test): {train_f1_score:.3f} / {f1_score:.3f}"
                    doc.add_paragraph(f1_text, style='Body Text')

                    # Generate confusion matrix as image
                    try:
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        import matplotlib.colors as mcolors
                        import numpy as np
                        import itertools
                        
                        num_classes = len(matrix)
                        if num_classes > 0:
                            fig, axes = plt.subplots(1, 2 if train_matrix else 1, figsize=(6 * (2 if train_matrix else 1), 5))
                            if not isinstance(axes, np.ndarray):
                                axes = [axes]

                            def plot_confusion_matrix(ax, cm, title, model_color):
                                ax.imshow(cm, interpolation='nearest', cmap=plt.cm.Blues)
                                ax.set_title(title, fontsize=10, fontweight='bold', color=model_color)
                                plt.colorbar(ax.imshow(cm, cmap=plt.cm.Blues), ax=ax, fraction=0.046, pad=0.04)
                                tick_marks = np.arange(num_classes)
                                ax.set_xticks(tick_marks)
                                ax.set_yticks(tick_marks)
                                ax.set_xlabel('Predicted label', fontsize=9)
                                ax.set_ylabel('True label', fontsize=9)
                                ax.tick_params(axis='both', which='major', labelsize=8)

                                # Add text annotations
                                thresh = cm.max() / 2.
                                for i, j in itertools.product(range(cm.shape[0]), range(cm.shape[1])):
                                    ax.text(j, i, format(cm[i, j], 'd'),
                                            horizontalalignment="center",
                                            color="white" if cm[i, j] > thresh else "black",
                                            fontsize=8)
                                # Add percentages
                                total = cm.sum()
                                for i, j in itertools.product(range(cm.shape[0]), range(cm.shape[1])):
                                    percentage = (cm[i, j] / total) * 100 if total > 0 else 0
                                    ax.text(j, i + 0.2, f'({percentage:.1f}%)',
                                            horizontalalignment="center",
                                            color="white" if cm[i, j] > thresh else "black",
                                            fontsize=7)

                            if train_matrix:
                                plot_confusion_matrix(axes[0], np.array(train_matrix), 'Train Confusion Matrix', color)
                                plot_confusion_matrix(axes[1], np.array(matrix), 'Test Confusion Matrix', color)
                            else:
                                plot_confusion_matrix(axes[0], np.array(matrix), 'Confusion Matrix', color)

                            plt.tight_layout()
                            img_buffer = BytesIO()
                            plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                            img_buffer.seek(0)
                            plt.close()

                            if img_buffer.getvalue():
                                doc.add_picture(img_buffer, width=Inches(6.0))
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            doc.add_paragraph(f'Confusion matrix data for {model_name} is empty.', style='Body Text')
                    except Exception as e:
                        logger.warning(f"Failed to generate confusion matrix image for {model_name}: {e}")
                        import traceback
                        logger.error(traceback.format_exc())
                        doc.add_paragraph(f'Confusion matrix for {model_name} not available', style='Body Text')
        
        # 8.4 Performance Radar Chart
        radar_charts = model_performance.get('radarCharts')
        if radar_charts:
            doc.add_paragraph()  # Add spacing
            doc.add_heading('8.4 Performance Radar Chart', 3)
            
            # 8.4.1 Train
            train_radar = radar_charts.get('train', [])
            doc.add_heading('8.4.1 Train', 4)
            if train_radar and len(train_radar) > 0:
                try:
                    import matplotlib
                    matplotlib.use('Agg')
                    import matplotlib.pyplot as plt
                    import numpy as np
                    from math import pi
                    
                    # Prepare data for radar chart
                    metrics = ['Accuracy', 'Precision', 'Recall', 'F1 Score', 'AUC-ROC']
                    num_metrics = len(metrics)
                    angles = [n / float(num_metrics) * 2 * pi for n in range(num_metrics)]
                    angles += angles[:1]  # Complete the circle
                    
                    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
                    
                    for model in train_radar:
                        values = [
                            model.get('accuracy', 0),
                            model.get('precision', 0),
                            model.get('recall', 0),
                            model.get('f1Score', 0),
                            model.get('aucRoc', 0)
                        ]
                        values += values[:1]  # Complete the circle
                        color = model.get('color', '#3b82f6')
                        model_name = model.get('modelName', 'Unknown')
                        
                        ax.plot(angles, values, 'o-', linewidth=2, label=model_name, color=color)
                        ax.fill(angles, values, alpha=0.1, color=color)
                    
                    ax.set_xticks(angles[:-1])
                    ax.set_xticklabels(metrics, fontsize=10)
                    ax.set_ylim(0, 1)
                    ax.set_ylabel('Score', fontsize=10, labelpad=20)
                    ax.set_title('Performance Radar Chart (Train)', fontsize=12, fontweight='bold', pad=20)
                    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=9)
                    ax.grid(True)
                    
                    img_buffer = BytesIO()
                    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    plt.close()
                    
                    if img_buffer.getvalue():
                        doc.add_picture(img_buffer, width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to generate Train Radar chart: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                    doc.add_paragraph('Performance Radar Chart (Train) not available', style='Body Text')
            else:
                doc.add_paragraph('No train radar chart data available', style='Body Text')
            
            # 8.4.2 Test
            test_radar = radar_charts.get('test', [])
            doc.add_heading('8.4.2 Test', 4)
            if test_radar and len(test_radar) > 0:
                try:
                    import matplotlib
                    matplotlib.use('Agg')
                    import matplotlib.pyplot as plt
                    import numpy as np
                    from math import pi
                    
                    # Prepare data for radar chart
                    metrics = ['Accuracy', 'Precision', 'Recall', 'F1 Score', 'AUC-ROC']
                    num_metrics = len(metrics)
                    angles = [n / float(num_metrics) * 2 * pi for n in range(num_metrics)]
                    angles += angles[:1]  # Complete the circle
                    
                    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
                    
                    for model in test_radar:
                        values = [
                            model.get('accuracy', 0),
                            model.get('precision', 0),
                            model.get('recall', 0),
                            model.get('f1Score', 0),
                            model.get('aucRoc', 0)
                        ]
                        values += values[:1]  # Complete the circle
                        color = model.get('color', '#3b82f6')
                        model_name = model.get('modelName', 'Unknown')
                        
                        ax.plot(angles, values, 'o-', linewidth=2, label=model_name, color=color)
                        ax.fill(angles, values, alpha=0.1, color=color)
                    
                    ax.set_xticks(angles[:-1])
                    ax.set_xticklabels(metrics, fontsize=10)
                    ax.set_ylim(0, 1)
                    ax.set_ylabel('Score', fontsize=10, labelpad=20)
                    ax.set_title('Performance Radar Chart (Test)', fontsize=12, fontweight='bold', pad=20)
                    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=9)
                    ax.grid(True)
                    
                    img_buffer = BytesIO()
                    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight')
                    img_buffer.seek(0)
                    plt.close()
                    
                    if img_buffer.getvalue():
                        doc.add_picture(img_buffer, width=Inches(6.0))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    logger.warning(f"Failed to generate Test Radar chart: {e}")
                    import traceback
                    logger.error(traceback.format_exc())
                    doc.add_paragraph('Performance Radar Chart (Test) not available', style='Body Text')
            else:
                doc.add_paragraph('No test radar chart data available', style='Body Text')
        
        # 8.5 Monotonicity
        monotonicity_data = model_performance.get('monotonicity', [])
        if monotonicity_data and len(monotonicity_data) > 0:
            doc.add_paragraph()  # Add spacing
            doc.add_heading('8.5 Monotonicity', 3)
            
            # Summary Table
            summary_table = doc.add_table(rows=len(monotonicity_data) + 1, cols=6)
            summary_table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = summary_table.rows[0].cells
            header_cells[0].text = 'Model'
            header_cells[1].text = 'Monotonicity Score'
            header_cells[2].text = 'KS Statistic'
            header_cells[3].text = 'Lift'
            header_cells[4].text = 'AUC/Gini'
            header_cells[5].text = 'PSI'
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for idx, mono in enumerate(monotonicity_data):
                data_cells = summary_table.rows[idx + 1].cells
                model_name = mono.get('modelName', 'Unknown')
                monotonicity_score = mono.get('monotonicityScore', 0)
                ks_statistic = mono.get('ksStatistic', 0)
                lift_top_decile = mono.get('liftTopDecile')
                auc = mono.get('auc', 0)
                gini = mono.get('gini', 0)
                psi_data = mono.get('psi')
                psi_value = None
                if psi_data:
                    if isinstance(psi_data, dict):
                        psi_value = psi_data.get('value')
                    elif isinstance(psi_data, (int, float)):
                        psi_value = psi_data
                
                data_cells[0].text = model_name
                data_cells[1].text = f'{monotonicity_score:.2f}%'
                data_cells[2].text = f'{ks_statistic:.3f}'
                data_cells[3].text = f'{lift_top_decile:.2f}x' if lift_top_decile is not None else 'N/A'
                data_cells[4].text = f'AUC {auc:.3f} / Gini {gini:.3f}'
                data_cells[5].text = f'{psi_value:.4f}' if psi_value is not None else 'N/A'
            
            # Add borders to summary table
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            for row in summary_table.rows:
                for cell in row.cells:
                    tcPr = cell._element.get_or_add_tcPr()
                    existing_borders = tcPr.find(qn('w:tcBorders'))
                    if existing_borders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    else:
                        tcBorders = existing_borders
                    
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = tcBorders.find(qn(f'w:{border_name}'))
                        if border is None:
                            border = OxmlElement(f'w:{border_name}')
                            border.set(qn('w:val'), 'single')
                            border.set(qn('w:sz'), '4')
                            border.set(qn('w:space'), '0')
                            border.set(qn('w:color'), '000000')
                            tcBorders.append(border)
            
            doc.add_paragraph()  # Add spacing after summary table
            
            # Summary Writeup
            monotonicity_summary = model_performance.get('monotonicitySummary', {})
            summary_writeup = monotonicity_summary.get('writeup', '')
            if summary_writeup:
                doc.add_paragraph(summary_writeup)
                doc.add_paragraph()  # Add spacing after writeup
            
            # Individual model sections
            for idx, mono in enumerate(monotonicity_data):
                model_name = mono.get('modelName', 'Unknown')
                doc.add_heading(f'8.5.{idx + 1} {model_name}', 4)
                
                # Create table for monotonicity metrics
                metrics_table = doc.add_table(rows=2, cols=4)
                metrics_table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = metrics_table.rows[0].cells
                header_cells[0].text = 'Monotonicity Score'
                header_cells[1].text = 'KS Statistic'
                header_cells[2].text = 'Lift (Top Decile)'
                header_cells[3].text = 'AUC / Gini'
                
                # Make header bold
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Data row
                data_cells = metrics_table.rows[1].cells
                
                # Monotonicity Score
                monotonicity_score = mono.get('monotonicityScore', 0)
                data_cells[0].text = f'{monotonicity_score:.2f}%'
                
                # KS Statistic
                ks_statistic = mono.get('ksStatistic', 0)
                ks_threshold = mono.get('ksThreshold', 0)
                data_cells[1].text = f'{ks_statistic:.3f} (Threshold: {ks_threshold:.3f})'
                
                # Lift (Top Decile)
                lift_top_decile = mono.get('liftTopDecile')
                overall_bad_rate = mono.get('overallBadRate', 0)
                if lift_top_decile is not None:
                    data_cells[2].text = f'{lift_top_decile:.2f}x (Overall bad rate: {(overall_bad_rate * 100):.2f}%)'
                else:
                    data_cells[2].text = 'N/A'
                
                # AUC / Gini
                auc = mono.get('auc', 0)
                gini = mono.get('gini', 0)
                data_cells[3].text = f'AUC {auc:.3f} · Gini {gini:.3f}'
                
                # Add borders to table
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                for row in metrics_table.rows:
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        existing_borders = tcPr.find(qn('w:tcBorders'))
                        if existing_borders is None:
                            tcBorders = OxmlElement('w:tcBorders')
                            tcPr.append(tcBorders)
                        else:
                            tcBorders = existing_borders
                        
                        for border_name in ['top', 'left', 'bottom', 'right']:
                            border = tcBorders.find(qn(f'w:{border_name}'))
                            if border is None:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                border.set(qn('w:color'), '000000')
                                tcBorders.append(border)
                
                doc.add_paragraph()  # Add spacing after table
                
                # PSI Section
                psi_data = mono.get('psi')
                if psi_data:
                    doc.add_paragraph('Population Stability Index (PSI):', style='Body Text')
                    psi_table = doc.add_table(rows=2, cols=3)
                    psi_table.style = 'Normal Table'
                    
                    # Header row
                    header_cells = psi_table.rows[0].cells
                    header_cells[0].text = 'PSI Value'
                    header_cells[1].text = 'Status'
                    header_cells[2].text = 'Interpretation'
                    
                    # Make header bold
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data row
                    data_cells = psi_table.rows[1].cells
                    psi_value = psi_data.get('value', 0)
                    psi_status = psi_data.get('status', 'N/A')
                    psi_interpretation = psi_data.get('interpretation', 'N/A')
                    data_cells[0].text = f'{psi_value:.4f}'
                    data_cells[1].text = psi_status
                    data_cells[2].text = psi_interpretation
                    
                    # Add borders to table
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    for row in psi_table.rows:
                        for cell in row.cells:
                            tcPr = cell._element.get_or_add_tcPr()
                            existing_borders = tcPr.find(qn('w:tcBorders'))
                            if existing_borders is not None:
                                tcPr.remove(existing_borders)
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                    
                    doc.add_paragraph()  # Add spacing after table
                
                # CSI Section
                csi_data = mono.get('csi', [])
                if csi_data and len(csi_data) > 0:
                    doc.add_paragraph('Characteristic Stability Index (CSI):', style='Body Text')
                    csi_table = doc.add_table(rows=len(csi_data) + 1, cols=3)
                    csi_table.style = 'Normal Table'
                    
                    # Header row
                    csi_header_cells = csi_table.rows[0].cells
                    csi_header_cells[0].text = 'Variable Name'
                    csi_header_cells[1].text = 'CSI Value'
                    csi_header_cells[2].text = 'Status'
                    
                    # Make header bold
                    for cell in csi_header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data rows
                    for idx, csi_row in enumerate(csi_data):
                        row_cells = csi_table.rows[idx + 1].cells
                        row_cells[0].text = csi_row.get('variable', 'N/A')
                        row_cells[1].text = f'{csi_row.get("csiValue", 0):.4f}'
                        row_cells[2].text = csi_row.get('status', 'N/A')
                    
                    # Add borders to table
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    for row in csi_table.rows:
                        for cell in row.cells:
                            tcPr = cell._element.get_or_add_tcPr()
                            existing_borders = tcPr.find(qn('w:tcBorders'))
                            if existing_borders is not None:
                                tcPr.remove(existing_borders)
                            tcBorders = OxmlElement('w:tcBorders')
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = OxmlElement(f'w:{border_name}')
                                border.set(qn('w:val'), 'single')
                                border.set(qn('w:sz'), '4')
                                border.set(qn('w:space'), '0')
                                tcBorders.append(border)
                            tcPr.append(tcBorders)
                    
                    doc.add_paragraph()  # Add spacing after table
                
                # Understand the Decile Progression
                decile_writeup = mono.get('decileProgressionWriteup')
                if decile_writeup:
                    doc.add_paragraph('Understand the Decile Progression:', style='Body Text')
                    doc.add_paragraph(decile_writeup, style='Body Text')
                
                # Decile Table
                deciles = mono.get('deciles', [])
                if deciles and len(deciles) > 0:
                    doc.add_paragraph()  # Add spacing before table
                    decile_table = doc.add_table(rows=len(deciles) + 1, cols=8)
                    decile_table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = decile_table.rows[0].cells
                    header_cells[0].text = 'Decile'
                    header_cells[1].text = 'Count'
                    header_cells[2].text = 'Bads'
                    header_cells[3].text = 'Goods'
                    header_cells[4].text = 'Bad Rate'
                    header_cells[5].text = 'Avg Score'
                    header_cells[6].text = 'Lift'
                    header_cells[7].text = 'Cum Bad Rate'
                    
                    # Make header bold
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Data rows
                    for idx, decile in enumerate(deciles):
                        data_cells = decile_table.rows[idx + 1].cells
                        decile_num = decile.get('Decile', idx + 1)
                        count = decile.get('Count', 0)
                        bads = decile.get('Bads', 0)
                        goods = decile.get('Goods', 0)
                        bad_rate = decile.get('Bad_Rate', 0)
                        avg_score = decile.get('Avg_Score', 0)
                        lift = decile.get('Lift', 0)
                        cum_bad_rate = decile.get('Cum_Bad_Rate', 0)
                        
                        data_cells[0].text = str(decile_num)
                        data_cells[1].text = str(count)
                        data_cells[2].text = str(bads)
                        data_cells[3].text = str(goods)
                        data_cells[4].text = f'{(bad_rate * 100):.2f}%' if bad_rate is not None else 'N/A'
                        data_cells[5].text = f'{avg_score:.3f}' if avg_score is not None else 'N/A'
                        data_cells[6].text = f'{lift:.2f}x' if lift is not None else 'N/A'
                        data_cells[7].text = f'{(cum_bad_rate * 100):.2f}%' if cum_bad_rate is not None else 'N/A'
                    
                    # Add borders to decile table
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    for row in decile_table.rows:
                        for cell in row.cells:
                            tcPr = cell._element.get_or_add_tcPr()
                            existing_borders = tcPr.find(qn('w:tcBorders'))
                            if existing_borders is None:
                                tcBorders = OxmlElement('w:tcBorders')
                                tcPr.append(tcBorders)
                            else:
                                tcBorders = existing_borders
                            
                            for border_name in ['top', 'left', 'bottom', 'right']:
                                border = tcBorders.find(qn(f'w:{border_name}'))
                                if border is None:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), 'single')
                                    border.set(qn('w:sz'), '4')
                                    border.set(qn('w:space'), '0')
                                    border.set(qn('w:color'), '000000')
                                    tcBorders.append(border)
                
                doc.add_paragraph()  # Add spacing between models
        
        # 8.6 Granular Accuracy
        granular_accuracy = model_performance.get('granularAccuracy')
        if granular_accuracy and granular_accuracy.get('variables') and len(granular_accuracy.get('variables', [])) > 0:
            doc.add_paragraph()  # Add spacing
            doc.add_heading('8.6 Granular Accuracy', 3)
            
            variables = granular_accuracy.get('variables', [])
            variables_to_show = granular_accuracy.get('variablesToShow', 5)
            
            # Limit to variables_to_show
            variables_to_display = variables[:variables_to_show]
            
            for variable in variables_to_display:
                variable_name = variable.get('variableName', 'Unknown')
                segments = variable.get('segments', [])
                
                if segments and len(segments) > 0:
                    doc.add_paragraph(f'Variable: {variable_name}', style='Heading 4')
                    
                    # Create table
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Light Grid Accent 1'
                    
                    # Header row
                    header_cells = table.rows[0].cells
                    header_cells[0].text = 'Segment'
                    header_cells[1].text = 'Accuracy'
                    header_cells[2].text = 'Precision'
                    header_cells[3].text = 'Recall'
                    header_cells[4].text = 'F1 Score'
                    
                    # Make header bold
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Add data rows
                    for segment in segments:
                        row_cells = table.add_row().cells
                        row_cells[0].text = segment.get('segment', 'Unknown')
                        row_cells[1].text = f"{(segment.get('accuracy', 0) * 100):.2f}%"
                        row_cells[2].text = f"{segment.get('precision', 0):.4f}"
                        row_cells[3].text = f"{segment.get('recall', 0):.4f}"
                        row_cells[4].text = f"{segment.get('f1Score', 0):.4f}"
                    
                    doc.add_paragraph()  # Add spacing between variables
        
        # 9. AI EXPLAINABILITY
        model_performance = documentation_data.get('modelPerformance', {})
        explainability = model_performance.get('explainability')
        
        # Check if AI EXPLAINABILITY section has data
        has_ai_explainability_data = False
        if explainability:
            shap_data = explainability.get('shap')
            pdp_data = explainability.get('pdp')
            writeup = explainability.get('writeup')
            has_ai_explainability_data = bool(
                (shap_data and (shap_data.get('beeswarm') or shap_data.get('waterfall'))) or
                (pdp_data and pdp_data.get('data')) or
                (writeup and writeup.get('content'))
            )
        
        doc.add_paragraph()  # Add spacing
        doc.add_heading('9. AI EXPLAINABILITY', 1)
        
        if not has_ai_explainability_data:
            # Show default message
            doc.add_paragraph('You didn\'t generate AI Explainability for your models. Go to AI Explainability page and see the SHAP and PDP values.')
        elif explainability:
            # 9.1 Understand AI Explainability
            explainability_writeup = explainability.get('writeup')
            if explainability_writeup and explainability_writeup.get('content'):
                doc.add_heading('9.1 Understand AI Explainability', 2)
                doc.add_paragraph(explainability_writeup['content'], style='Body Text')
                doc.add_paragraph()  # Add spacing
            
            # 9.2 SHAP
            shap_data = explainability.get('shap')
            if shap_data:
                doc.add_heading('9.2 SHAP', 2)
                
                # 9.2.1 Beeswarm Plot
                beeswarm = shap_data.get('beeswarm')
                if beeswarm and beeswarm.get('data') and len(beeswarm.get('data', [])) > 0:
                    doc.add_heading('9.2.1 Beeswarm Plot', 3)
                    
                    feature_count = beeswarm.get('featureCount', 'all')
                    filtered_data = beeswarm.get('data', [])
                    if feature_count != 'all' and isinstance(feature_count, int):
                        filtered_data = filtered_data[:feature_count]
                    
                    doc.add_paragraph(f'Showing {len(filtered_data)} features (filter: {feature_count if feature_count != "all" else "All"}).', style='Body Text')
                    
                    # Generate Beeswarm plot as image
                    try:
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        import numpy as np
                        from io import BytesIO as ChartBytesIO
                        
                        # Filter out features with no values
                        valid_features = [item for item in filtered_data if item.get('values') and len(item.get('values', [])) > 0]
                        
                        if not valid_features:
                            doc.add_paragraph('No valid SHAP data available for Beeswarm plot.', style='Body Text')
                        else:
                            # Calculate global SHAP value range
                            all_shap_values = []
                            for item in valid_features:
                                item_values = item.get('values', [])
                                if isinstance(item_values, list):
                                    all_shap_values.extend([v for v in item_values if isinstance(v, (int, float))])
                            
                            if all_shap_values:
                                max_abs_shap = max(abs(v) for v in all_shap_values)
                                global_plot_min = -max_abs_shap * 1.1
                                global_plot_max = max_abs_shap * 1.1
                            else:
                                global_plot_min, global_plot_max = -1, 1
                            
                            # Create figure for all beeswarm plots
                            fig_height_per_feature = 0.8
                            fig_height = max(3, len(valid_features) * fig_height_per_feature)
                            fig, axes = plt.subplots(len(valid_features), 1, figsize=(8, fig_height), squeeze=False)
                            fig.suptitle('SHAP Beeswarm Plot', fontsize=14, fontweight='bold', y=1.02)
                            
                            for i, feature_data in enumerate(valid_features):
                                ax = axes[i, 0]
                                values = feature_data.get('values', [])
                                # Ensure values are numeric
                                values = [v for v in values if isinstance(v, (int, float))]
                                
                                if not values:
                                    continue
                                
                                feature_values = feature_data.get('feature_values', [])
                                original_feature_values = feature_data.get('original_feature_values', [])
                                feature_name = feature_data.get('featureName', f'Feature {i}')
                                original_feature_name = feature_data.get('original_feature_name', feature_name)
                                
                                # Use original feature values for coloring if available
                                values_to_use_for_color = original_feature_values if original_feature_values else feature_values
                                # Ensure color values are numeric
                                if values_to_use_for_color:
                                    values_to_use_for_color = [v for v in values_to_use_for_color if isinstance(v, (int, float))]
                                    if len(values_to_use_for_color) != len(values):
                                        values_to_use_for_color = values
                                else:
                                    values_to_use_for_color = values
                                
                                # Determine feature value range for coloring
                                if values_to_use_for_color and len(values_to_use_for_color) > 0:
                                    feat_min = min(values_to_use_for_color)
                                    feat_max = max(values_to_use_for_color)
                                else:
                                    feat_min, feat_max = 0, 1
                                
                                # Create colormap
                                cmap = plt.cm.coolwarm
                                norm = plt.Normalize(vmin=feat_min, vmax=feat_max)
                                
                                # Jitter points vertically
                                np.random.seed(42)  # For reproducibility
                                y_jitter = np.random.rand(len(values)) * 0.8 + 0.1
                                
                                # Plot points
                                scatter = ax.scatter(
                                    values,
                                    y_jitter,
                                    c=values_to_use_for_color,
                                    cmap=cmap,
                                    norm=norm,
                                    s=15,
                                    alpha=0.7,
                                    edgecolors='none',
                                    rasterized=True
                                )
                                
                                # Add vertical line at SHAP value = 0
                                ax.axvline(0, color='black', linestyle='--', linewidth=0.8, alpha=0.7)
                                
                                # Set limits and labels
                                ax.set_xlim(global_plot_min, global_plot_max)
                                ax.set_ylim(0, 1)
                                ax.set_yticks([])
                                ax.set_ylabel(original_feature_name, rotation=0, ha='right', va='center', fontsize=10, fontweight='bold')
                                ax.tick_params(axis='x', labelsize=8)
                                
                                # Remove spines
                                ax.spines['left'].set_visible(False)
                                ax.spines['right'].set_visible(False)
                                ax.spines['top'].set_visible(False)
                                if i < len(valid_features) - 1:
                                    ax.spines['bottom'].set_visible(False)
                                    ax.set_xticks([])
                                else:
                                    ax.set_xlabel('SHAP Value', fontsize=10)
                            
                            plt.tight_layout(rect=[0, 0.03, 1, 0.98])
                            img_buffer = ChartBytesIO()
                            plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
                            img_buffer.seek(0)
                            buffer_size = len(img_buffer.getvalue())
                            plt.close()
                            
                            if buffer_size > 0:
                                doc.add_picture(img_buffer, width=Inches(6.5))
                                last_paragraph = doc.paragraphs[-1]
                                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                doc.add_paragraph('Failed to generate SHAP Beeswarm Plot image (empty buffer).', style='Body Text')
                    except Exception as e:
                        logger.error(f"Error generating SHAP Beeswarm Plot for docx: {str(e)}")
                        import traceback
                        logger.error(traceback.format_exc())
                        doc.add_paragraph(f'Error generating SHAP Beeswarm Plot: {str(e)}', style='Body Text')
                    
                    doc.add_paragraph('Note: SHAP Beeswarm plots show the distribution of SHAP values for each feature. Each dot represents a single prediction. Features are ranked by importance.', style='Body Text')
                
                # 9.2.2 Waterfall
                waterfall = shap_data.get('waterfall')
                if waterfall and waterfall.get('data') and len(waterfall.get('data', [])) > 0:
                    doc.add_heading('9.2.2 Waterfall', 3)
                    
                    feature_count = waterfall.get('featureCount', 'all')
                    filtered_data = waterfall.get('data', [])
                    if feature_count != 'all' and isinstance(feature_count, int):
                        filtered_data = filtered_data[:feature_count]
                    
                    base_value = waterfall.get('baseValue', 0)
                    doc.add_paragraph(f'Base Value: {base_value:.4f}', style='Body Text')
                    doc.add_paragraph(f'Showing {len(filtered_data)} features (filter: {feature_count if feature_count != "all" else "All"}).', style='Body Text')
                    
                    # Generate Waterfall plot as image (matching web UI layout)
                    try:
                        import matplotlib
                        matplotlib.use('Agg')
                        import matplotlib.pyplot as plt
                        import numpy as np
                        from io import BytesIO as ChartBytesIO
                        
                        # Calculate cumulative values
                        cumulative = base_value
                        features_with_cumulative = []
                        for feature in filtered_data:
                            prev_cumulative = cumulative
                            shap_value = feature.get('shap_value', 0)
                            cumulative += shap_value
                            features_with_cumulative.append({
                                'feature': feature.get('feature', 'Unknown'),
                                'feature_value': feature.get('feature_value', 0),
                                'shap_value': shap_value,
                                'prev_cumulative': prev_cumulative,
                                'cumulative': cumulative,
                                'is_positive': shap_value > 0
                            })
                        
                        final_prediction = cumulative
                        all_values = [base_value] + [f['cumulative'] for f in features_with_cumulative] + [final_prediction]
                        min_val = min(all_values)
                        max_val = max(all_values)
                        range_val = max_val - min_val or 0.1
                        padding = range_val * 0.15  # 15% padding like web UI
                        padded_min = min_val - padding
                        padded_max = max_val + padding
                        total_range = padded_max - padded_min
                        
                        # Normalize function
                        def normalize(val):
                            return (val - padded_min) / total_range
                        
                        # Calculate figure size - vertical layout like web UI
                        num_rows = 1 + len(features_with_cumulative) + 1  # Base + features + final
                        row_height = 0.5
                        fig_height = max(4, num_rows * row_height + 1.5)  # Extra space for scale
                        fig, ax = plt.subplots(figsize=(10, fig_height))
                        
                        # Plot base value
                        base_width = normalize(base_value) * 10  # Scale to figure width
                        base_bar = ax.barh(0, base_width, left=0, height=0.4, color='#64748b', alpha=0.8)
                        base_label_x = normalize(base_value) * 10
                        ax.text(base_label_x + 0.3, 0, f'{base_value:.2f}', 
                               ha='left', va='center', fontsize=9, fontweight='bold',
                               bbox=dict(boxstyle='round,pad=0.3', facecolor='#64748b', edgecolor='none', alpha=0.9))
                        ax.text(-1.2, 0, 'Base Value\nE[f(x)]', ha='right', va='center', fontsize=8)
                        
                        # Plot feature contributions
                        y_pos = 1
                        for feat in features_with_cumulative:
                            prev_pos = normalize(feat['prev_cumulative']) * 10
                            curr_pos = normalize(feat['cumulative']) * 10
                            bar_start = min(prev_pos, curr_pos)
                            bar_width = abs(curr_pos - prev_pos)
                            
                            color = '#06b6d4' if feat['is_positive'] else '#f43f5e'
                            feature_bar = ax.barh(y_pos, bar_width, left=bar_start, height=0.4, color=color, alpha=0.8)
                            
                            # Value label at the end of the bar
                            label_x = curr_pos + 0.3
                            shap_sign = '+' if feat['is_positive'] else ''
                            ax.text(label_x, y_pos, f'{shap_sign}{feat["shap_value"]:.3f}', 
                                   ha='left', va='center', fontsize=8, fontweight='bold',
                                   bbox=dict(boxstyle='round,pad=0.3', facecolor=color, edgecolor='none', alpha=0.9))
                            
                            # Feature name on the left
                            feature_display = feat['feature'].replace('_', ' ')[:20]
                            ax.text(-1.2, y_pos, f'{feature_display}\n= {feat["feature_value"]:.2f}', 
                                   ha='right', va='center', fontsize=7)
                            
                            # End position marker
                            ax.axvline(curr_pos, ymin=(y_pos-0.2)/fig_height, ymax=(y_pos+0.2)/fig_height, 
                                      color='gray', linewidth=0.5, alpha=0.5)
                            
                            y_pos += 1
                        
                        # Plot final prediction (with border like web UI)
                        final_width = normalize(final_prediction) * 10
                        final_bar = ax.barh(y_pos, final_width, left=0, height=0.4, color='#3b82f6', alpha=0.8, 
                                           edgecolor='#1e40af', linewidth=2)
                        final_label_x = normalize(final_prediction) * 10
                        ax.text(final_label_x + 0.3, y_pos, f'{final_prediction:.3f}', 
                               ha='left', va='center', fontsize=9, fontweight='bold',
                               bbox=dict(boxstyle='round,pad=0.3', facecolor='#3b82f6', edgecolor='none', alpha=0.9))
                        ax.text(-1.2, y_pos, 'f(x)\nFinal Prediction', ha='right', va='center', fontsize=8, fontweight='bold')
                        
                        # Add separator line before final prediction
                        ax.axhline(y_pos - 0.5, xmin=0, xmax=1, color='gray', linewidth=2, alpha=0.5)
                        
                        # Scale at the bottom (like web UI)
                        scale_y = -0.8
                        scale_min = round(min_val, 1)
                        scale_max = round(max_val, 1)
                        ax.text(0, scale_y, f'{scale_min:.1f}', ha='left', va='center', fontsize=8)
                        ax.text(10, scale_y, 'Prediction Score', ha='center', va='center', fontsize=9, fontweight='bold')
                        ax.text(10, scale_y, f'{scale_max:.1f}', ha='right', va='center', fontsize=8)
                        ax.axhline(scale_y - 0.1, xmin=0, xmax=1, color='gray', linewidth=1)
                        
                        # Set limits and styling
                        ax.set_xlim(-1.5, 11)
                        ax.set_ylim(scale_y - 0.5, y_pos + 0.5)
                        ax.set_yticks([])
                        ax.set_xticks([])
                        ax.spines['left'].set_visible(False)
                        ax.spines['right'].set_visible(False)
                        ax.spines['top'].set_visible(False)
                        ax.spines['bottom'].set_visible(False)
                        
                        plt.tight_layout()
                        img_buffer = ChartBytesIO()
                        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
                        img_buffer.seek(0)
                        buffer_size = len(img_buffer.getvalue())
                        plt.close()
                        
                        if buffer_size > 0:
                            doc.add_picture(img_buffer, width=Inches(6.5))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            doc.add_paragraph('Failed to generate SHAP Waterfall Plot image.', style='Body Text')
                    except Exception as e:
                        logger.error(f"Error generating SHAP Waterfall Plot for docx: {str(e)}")
                        import traceback
                        logger.error(traceback.format_exc())
                        doc.add_paragraph('Error generating SHAP Waterfall Plot.', style='Body Text')
                    
                    doc.add_paragraph('Note: SHAP Waterfall plot explains a single prediction by showing how each feature value pushes the prediction from the base value to the final prediction.', style='Body Text')
            
            # 9.3 PDP/ICE Lines
            pdp_data = explainability.get('pdp')
            if pdp_data and pdp_data.get('data') and len(pdp_data.get('data', [])) > 0:
                doc.add_heading('9.3 PDP/ICE Lines', 2)
                
                feature_count = pdp_data.get('featureCount', 5)
                max_ice_lines = pdp_data.get('maxIceLines', 100)
                filtered_data = pdp_data.get('data', [])
                if feature_count != 'all' and isinstance(feature_count, int):
                    filtered_data = filtered_data[:feature_count]
                
                doc.add_paragraph(f'Showing {len(filtered_data)} features (filter: {feature_count if feature_count != "all" else "All"}), with up to {max_ice_lines} ICE lines per plot.', style='Body Text')
                
                # Generate PDP/ICE plots as images
                try:
                    import matplotlib
                    matplotlib.use('Agg')
                    import matplotlib.pyplot as plt
                    import numpy as np
                    from io import BytesIO as ChartBytesIO
                    
                    for feature_pdp_data in filtered_data:
                        feature_name = feature_pdp_data.get('feature_name', 'Unknown')
                        values = feature_pdp_data.get('values', [])  # Expecting [{x, y}, ...]
                        ice_lines = feature_pdp_data.get('ice_lines', [])
                        
                        if not values:
                            doc.add_paragraph(f'No PDP data for feature: {feature_name}', style='Body Text')
                            continue
                        
                        # Extract x and y values
                        x_values = [p.get('x', 0) if isinstance(p, dict) else 0 for p in values]
                        y_values = [p.get('y', 0) if isinstance(p, dict) else 0 for p in values]
                        
                        # Determine plot limits
                        min_x, max_x = min(x_values), max(x_values)
                        min_y, max_y = min(y_values), max(y_values)
                        
                        # Include ICE lines in y-range calculation
                        all_ice_y_values = []
                        for line in ice_lines[:max_ice_lines]:
                            if isinstance(line, list):
                                all_ice_y_values.extend(line)
                        
                        if all_ice_y_values:
                            min_y = min(min_y, min(all_ice_y_values))
                            max_y = max(max_y, max(all_ice_y_values))
                        
                        # Add padding to y-axis
                        y_range = max_y - min_y
                        if y_range == 0:
                            y_range = abs(min_y) * 0.1 if min_y != 0 else 0.1
                        min_y -= y_range * 0.1
                        max_y += y_range * 0.1
                        
                        fig, ax = plt.subplots(figsize=(8, 4))
                        
                        # Plot ICE lines (limited by maxIceLines)
                        for i, line in enumerate(ice_lines[:max_ice_lines]):
                            if isinstance(line, list) and len(line) == len(x_values):
                                ax.plot(x_values, line, color='gray', alpha=0.1, linewidth=0.8)
                        
                        # Plot PDP line
                        ax.plot(x_values, y_values, color='#06b6d4', linewidth=2, label='Partial Dependence')
                        
                        # Add axes labels
                        display_name = feature_name.replace('_', ' ').title()
                        ax.set_title(f'PDP with ICE Lines for {display_name}', fontsize=12, fontweight='bold')
                        ax.set_xlabel(display_name, fontsize=10)
                        ax.set_ylabel('Prediction Probability', fontsize=10)
                        ax.set_xlim(min_x, max_x)
                        ax.set_ylim(min_y, max_y)
                        ax.grid(True, linestyle='--', alpha=0.6)
                        ax.tick_params(axis='both', which='major', labelsize=8)
                        ax.legend(fontsize=8)
                        
                        plt.tight_layout()
                        img_buffer = ChartBytesIO()
                        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
                        img_buffer.seek(0)
                        plt.close()
                        
                        if img_buffer.getvalue():
                            doc.add_paragraph(feature_name, style='Heading 4')
                            doc.add_picture(img_buffer, width=Inches(6.5))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            doc.add_paragraph(f'Failed to generate PDP/ICE Plot image for {feature_name}.', style='Body Text')
                except Exception as e:
                    logger.error(f"Error generating PDP/ICE Lines Plot for docx: {str(e)}")
                    import traceback
                    logger.error(traceback.format_exc())
                    doc.add_paragraph('Error generating PDP/ICE Lines Plot.', style='Body Text')
                
                doc.add_paragraph('Note: Partial Dependence Plots (PDP) show the marginal effect of a feature on the predicted outcome. ICE (Individual Conditional Expectation) lines show the effect for individual samples.', style='Body Text')
        
        # MODEL OWNER section
        doc.add_paragraph()  # Add spacing before new section
        heading = doc.add_heading('10. MODEL OWNER', 1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        
        model_owner = documentation_data.get('modelOwner', {})
        
        # Check if MODEL OWNER section has data
        approved_by = model_owner.get('approvedBy', '')
        created_by = model_owner.get('createdBy', '')
        created_on = model_owner.get('createdOn', '')
        has_model_owner_data = bool(approved_by or (created_by and created_by != 'Unknown User') or created_on)
        
        if not has_model_owner_data:
            # Show default message
            doc.add_paragraph('You skipped to the documentation, upload the Data first')
        else:
            # Approved By
            doc.add_paragraph('The model will be approved by:', style='Heading 3')
            if approved_by:
                doc.add_paragraph(approved_by, style='Body Text')
            else:
                doc.add_paragraph('(Not specified)', style='Body Text')
            
            doc.add_paragraph()  # Add spacing
            
            # Created By
            doc.add_paragraph('Model created by:', style='Heading 3')
            created_by = model_owner.get('createdBy', 'Unknown User')
            doc.add_paragraph(created_by, style='Body Text')
            
            doc.add_paragraph()  # Add spacing
            
            # Created On
            doc.add_paragraph('Created on:', style='Heading 3')
            created_on = model_owner.get('createdOn', '')
            if created_on:
                try:
                    from datetime import datetime
                    created_date = datetime.fromisoformat(created_on.replace('Z', '+00:00'))
                    formatted_date = created_date.strftime('%B %d, %Y at %I:%M %p')
                    doc.add_paragraph(formatted_date, style='Body Text')
                except Exception as e:
                    logger.warning(f"Error formatting created_on date: {e}")
                    doc.add_paragraph(created_on, style='Body Text')
            else:
                doc.add_paragraph('Not available', style='Body Text')
        
        # Save DOCX to temporary directory (outside the if/else block)
        docx_filename = 'model_documentation.docx'
        docx_path = os.path.join(temp_dir, docx_filename)
        doc.save(docx_path)
        
        logger.info("Documentation .docx file generated successfully")
        
        # Attempt to embed Excel files as OLE objects using COM automation (Windows + Word)
        # This will enhance the document if available, otherwise gracefully skip
        ole_success = post_process_docx_with_ole_objects(docx_path, excel_files)
        if ole_success:
            logger.info("✅ OLE objects successfully embedded in DOCX")
        else:
            logger.info("ℹ️ Using styled references for Excel files (OLE embedding not available)")
        
        # Create ZIP file containing DOCX and all Excel files
        zip_io = BytesIO()
        with zipfile.ZipFile(zip_io, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add DOCX file
            zip_file.write(docx_path, docx_filename)
            logger.info(f"Added {docx_filename} to ZIP")
            
            # Add all Excel files
            for excel_path in excel_files:
                excel_filename = os.path.basename(excel_path)
                zip_file.write(excel_path, excel_filename)
                logger.info(f"Added {excel_filename} to ZIP")
        
        zip_io.seek(0)
        
        logger.info(f"ZIP file created successfully with {len(excel_files)} Excel files and 1 DOCX file")
        
        return StreamingResponse(
            zip_io,
            media_type="application/zip",
            headers={"Content-Disposition": "attachment; filename=model_documentation.zip"}
        )
        
    except Exception as e:
        logger.error(f"Error generating documentation file: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        # Clean up temporary directory
        try:
            shutil.rmtree(temp_dir)
            logger.info(f"Cleaned up temporary directory: {temp_dir}")
        except Exception as cleanup_error:
            logger.warning(f"Failed to clean up temporary directory: {cleanup_error}")
